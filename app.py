# Loads a local .env file (if present) before anything below reads os.environ —
# marketdata/config.py in particular reads its provider credentials at import
# time, so this must run first. Absent .env (e.g. in production, where real env
# vars are set directly), this is a silent no-op.
from dotenv import load_dotenv
load_dotenv()

from flask import Flask, jsonify, request, send_from_directory, Response, session
from flask_cors import CORS
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from functools import wraps
import yfinance as yf
import pandas as pd
import bcrypt
import json
import os
import secrets
import hashlib
import hmac
import time
import threading
import re
import smtplib
import requests
from bs4 import BeautifulSoup
from email.mime.text import MIMEText
try:
    import psycopg2
    import psycopg2.extras
except ImportError:
    psycopg2 = None
from datetime import timedelta
import datetime as _dt
import base64
from werkzeug.utils import secure_filename
from werkzeug.middleware.proxy_fix import ProxyFix

from api.indicators import calculate_all
from api.signals import score_signals
from api.backtest import run_backtest
from api.metrics import calculate_metrics
from api.market_context import enrich_trades_with_sector_context

from marketdata import router as marketdata_router
from marketdata import bus as marketdata_bus
from marketdata import config as marketdata_config

app = Flask(__name__, static_folder="static")
app.secret_key = os.environ.get("SECRET_KEY", "gcg-dev-key-change-in-production")
# Railway (and most PaaS hosts) terminate TLS at an edge proxy and forward requests
# over plain HTTP, so without this Flask sees every request as insecure — which
# breaks Secure-cookie handling (session + remember-me) behind the proxy.
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1)

# Base URL of the realtime/ async streaming service (docs/scaling-plan.md),
# e.g. "https://ai-investor-realtime.up.railway.app" — no trailing slash.
# Blank until that service is actually deployed; every consumer of this
# treats blank as "feature not available yet, fall back to polling".
_REALTIME_BASE_URL = os.environ.get("REALTIME_BASE_URL", "").rstrip("/")

_is_production = os.environ.get("RAILWAY_ENVIRONMENT") or os.environ.get("PRODUCTION")
_allowed_origins_env = os.environ.get("ALLOWED_ORIGINS", "")
_allowed_origins = (
    [o.strip() for o in _allowed_origins_env.split(",") if o.strip()]
    if _allowed_origins_env
    else ["http://localhost:5000", "http://127.0.0.1:5000"]
)

app.config.update(
    MAX_CONTENT_LENGTH=5 * 1024 * 1024,
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=bool(_is_production),
    PERMANENT_SESSION_LIFETIME=timedelta(days=30),
    REMEMBER_COOKIE_DURATION=timedelta(days=30),
    REMEMBER_COOKIE_HTTPONLY=True,
    REMEMBER_COOKIE_SECURE=bool(_is_production),
)
CORS(app, supports_credentials=True, origins=_allowed_origins)

login_manager = LoginManager(app)
login_manager.session_protection = "basic"

USERS_FILE  = os.path.join(os.path.dirname(__file__), "users.json")
DATABASE_URL = os.environ.get("DATABASE_URL", "")
ALLOWED_IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "gif", "webp"}

# ── Alpha content ────────────────────────────────────────────────────────────
ALPHA_ROLES = {"tom", "dave", "gary", "connor"}
ALPHA_CONTENT_KINDS = {"post", "video", "link", "watchlist"}
ALPHA_STANCES = {"bullish", "neutral", "bearish"}
# The audience level a post is written for — shown as a second pill next to
# the topic pill on the public Alpha pages and post page.
ALPHA_LEVELS = {"beginner", "intermediate", "pro"}
ALPHA_CONTENT_FILE = os.path.join(os.path.dirname(__file__), "alpha_content.json")
DATAVIZ_CONTENT_FILE = os.path.join(os.path.dirname(__file__), "dataviz_content.json")
ALPHA_ATTACHMENTS_FILE = os.path.join(os.path.dirname(__file__), "alpha_attachments.json")
ALLOWED_DOC_EXTENSIONS = {"docx", "pdf", "xlsx", "xls"}
MAX_UPLOAD_TEXT_CHARS = 40000  # cap extracted text sent to the normalize step

# Each partner's 3 nominated topics — must match the topic-pill labels
# hardcoded on their public page (static/alpha-<name>.html).
ALPHA_TOPICS = {
    "tom":    ["Stocks", "Metals", "ETFs"],
    "dave":   ["Simple Investing", "Crypto", "Stocks"],
    "gary":   ["Long-Term Investing", "Technical Analysis", "Impact of AI"],
    "connor": ["Macro & Micro", "Quant Statistics", "Finance & Formulae"],
}

# Single source of truth for every /learn/<level>/<slug> lesson: adding a
# lesson is "add one entry here + create the HTML file" — no separate route
# to write, and it shows up in site search automatically (see
# _lesson_search_pages() near api_search below). Registered as Flask routes
# in a loop right after this file's other @app.route definitions.
LESSON_PAGES = [
    {"slug": "start-early", "level": "beginner", "file": "lesson-start-early.html", "title": "Start Early & Compounding"},
    {"slug": "diversify", "level": "beginner", "file": "lesson-diversify.html", "title": "Diversify"},
    {"slug": "invest-consistently", "level": "beginner", "file": "lesson-invest-consistently.html", "title": "Invest Consistently"},
    {"slug": "keep-costs-low", "level": "beginner", "file": "lesson-keep-costs-low.html", "title": "Keep Costs Low"},
    {"slug": "think-in-decades", "level": "beginner", "file": "lesson-think-in-decades.html", "title": "Think in Decades, Not Days"},
    {"slug": "stocks", "level": "beginner", "file": "lesson-stocks.html", "title": "Stocks"},
    {"slug": "etfs", "level": "beginner", "file": "lesson-etfs.html", "title": "ETFs & Index Funds"},
    {"slug": "bonds", "level": "beginner", "file": "lesson-bonds.html", "title": "Bonds"},
    {"slug": "cash", "level": "beginner", "file": "lesson-cash.html", "title": "Cash & Cash Equivalents"},
    {"slug": "alternatives", "level": "beginner", "file": "lesson-alternatives.html", "title": "Alternative Assets"},
    {"slug": "candlesticks", "level": "beginner", "file": "lesson-candlesticks.html", "title": "Candlesticks"},
    {"slug": "trend-lines", "level": "beginner", "file": "lesson-trend-lines.html", "title": "Trend Lines"},
    {"slug": "volume", "level": "beginner", "file": "lesson-volume.html", "title": "Volume"},
    {"slug": "support-resistance", "level": "beginner", "file": "lesson-support-resistance.html", "title": "Support & Resistance"},
    {"slug": "moving-averages", "level": "beginner", "file": "lesson-moving-averages.html", "title": "Moving Averages"},
    {"slug": "chart-patterns", "level": "intermediate", "file": "lesson-chart-patterns.html", "title": "Chart Patterns"},
    {"slug": "rsi", "level": "intermediate", "file": "lesson-rsi.html", "title": "RSI"},
    {"slug": "macd", "level": "intermediate", "file": "lesson-macd.html", "title": "MACD"},
    {"slug": "bollinger-bands", "level": "intermediate", "file": "lesson-bollinger-bands.html", "title": "Bollinger Bands"},
    {"slug": "confluence", "level": "intermediate", "file": "lesson-confluence.html", "title": "Stacking Indicators for Confluence"},
    {"slug": "income-statement", "level": "intermediate", "file": "lesson-income-statement.html", "title": "Reading the Income Statement"},
    {"slug": "balance-sheet", "level": "intermediate", "file": "lesson-balance-sheet.html", "title": "The Balance Sheet"},
    {"slug": "cash-flow", "level": "intermediate", "file": "lesson-cash-flow.html", "title": "Cash Flow"},
]

DATAVIZ_PAGES_FILE = os.path.join(os.path.dirname(__file__), "dataviz_pages.json")


def _db_conn():
    url = DATABASE_URL
    if url.startswith("postgres://"):
        url = url.replace("postgres://", "postgresql://", 1)
    return psycopg2.connect(url)


def _ensure_table() -> None:
    if not DATABASE_URL:
        return
    with _db_conn() as conn, conn.cursor() as cur:
        cur.execute("""
            CREATE TABLE IF NOT EXISTS users (
                username      TEXT PRIMARY KEY,
                password_hash TEXT NOT NULL,
                preferences   JSONB NOT NULL DEFAULT '{}',
                tier          TEXT NOT NULL DEFAULT 'basic'
            )
        """)
        cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS tier TEXT NOT NULL DEFAULT 'basic'")
        cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS profile JSONB NOT NULL DEFAULT '{}'::jsonb")
        cur.execute("UPDATE users SET tier = 'power_user' WHERE tier IN ('basic', 'signal_tester')")
        cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS alpha_role TEXT")
        cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS avatar_file BYTEA")
        cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS avatar_filename TEXT")
        cur.execute("""
            CREATE TABLE IF NOT EXISTS alpha_content (
                id              SERIAL PRIMARY KEY,
                author          TEXT NOT NULL,
                kind            TEXT NOT NULL,
                status          TEXT NOT NULL DEFAULT 'draft',
                topic           TEXT,
                title           TEXT,
                snippet         TEXT,
                body            TEXT,
                stance          TEXT,
                url             TEXT,
                source_kind     TEXT,
                source_filename TEXT,
                source_file     BYTEA,
                source_text     TEXT,
                created_at      TIMESTAMP NOT NULL DEFAULT now(),
                updated_at      TIMESTAMP NOT NULL DEFAULT now(),
                published_at    TIMESTAMP
            )
        """)
        cur.execute("ALTER TABLE alpha_content ADD COLUMN IF NOT EXISTS subtitle TEXT")
        cur.execute("ALTER TABLE alpha_content ADD COLUMN IF NOT EXISTS image_url TEXT")
        cur.execute("ALTER TABLE alpha_content ADD COLUMN IF NOT EXISTS image_filename TEXT")
        cur.execute("ALTER TABLE alpha_content ADD COLUMN IF NOT EXISTS image_file BYTEA")
        cur.execute("ALTER TABLE alpha_content ADD COLUMN IF NOT EXISTS staged_edits JSONB")
        cur.execute("ALTER TABLE alpha_content ADD COLUMN IF NOT EXISTS level TEXT")
        cur.execute("ALTER TABLE alpha_content ADD COLUMN IF NOT EXISTS pinned BOOLEAN NOT NULL DEFAULT false")
        # Cross-link to a /learn lesson this piece of Alpha content relates to —
        # set from the Studio once an author/Gary confirms the two are actually
        # related. related_lesson_slug is one of LESSON_PAGES' slugs (not
        # FK-enforced, since lessons live in code not the DB); related_lesson_note
        # is the one-line "why this is relevant" shown in the lesson's callout box.
        cur.execute("ALTER TABLE alpha_content ADD COLUMN IF NOT EXISTS related_lesson_slug TEXT")
        cur.execute("ALTER TABLE alpha_content ADD COLUMN IF NOT EXISTS related_lesson_note TEXT")
        cur.execute("""
            CREATE TABLE IF NOT EXISTS alpha_content_attachment (
                id          SERIAL PRIMARY KEY,
                content_id  INTEGER NOT NULL,
                filename    TEXT,
                file        BYTEA NOT NULL,
                created_at  TIMESTAMP NOT NULL DEFAULT now()
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS dataviz_content (
                id                 SERIAL PRIMARY KEY,
                author             TEXT NOT NULL,
                page               TEXT,
                status             TEXT NOT NULL DEFAULT 'draft',
                title              TEXT,
                description        TEXT,
                positive_analysis  TEXT,
                warning            TEXT,
                link               TEXT,
                image_filename     TEXT,
                image_file         BYTEA,
                created_at         TIMESTAMP NOT NULL DEFAULT now(),
                updated_at         TIMESTAMP NOT NULL DEFAULT now(),
                published_at       TIMESTAMP
            )
        """)
        cur.execute("ALTER TABLE dataviz_content ADD COLUMN IF NOT EXISTS page TEXT")
        cur.execute("""
            CREATE TABLE IF NOT EXISTS dataviz_pages (
                slug        TEXT PRIMARY KEY,
                label       TEXT NOT NULL,
                author      TEXT,
                created_at  TIMESTAMP NOT NULL DEFAULT now()
            )
        """)
        # has_live_widget marks pages like "market-pulse" that carry a bespoke,
        # code-built visualization (a live map, a chart, a calculator — anything
        # beyond a picture) at the top of the page. Those pages don't need Tom/Gary
        # to attach a hero image to every commentary post underneath — the widget
        # already serves that role — so the publish-time "upload an image first"
        # rule (see api_dataviz_content_item) is skipped for them.
        cur.execute("ALTER TABLE dataviz_pages ADD COLUMN IF NOT EXISTS has_live_widget BOOLEAN NOT NULL DEFAULT FALSE")
        # header_image: an optional page-level hero image, shown above the title
        # (and above any live widget). Separate from dataviz_content.image_file,
        # which is per-post — this is per-page, and the only image field a
        # has_live_widget page has, since its per-post image requirement is
        # waived (see comment above).
        cur.execute("ALTER TABLE dataviz_pages ADD COLUMN IF NOT EXISTS header_image_filename TEXT")
        cur.execute("ALTER TABLE dataviz_pages ADD COLUMN IF NOT EXISTS header_image_file BYTEA")
        # description: the page's own summary text (title/description/image is the
        # whole visualisation now — see the Studio rework below). Replaces the old
        # per-page "posts" model: dataviz_content still exists in the DB but the
        # Studio no longer offers it, so this is the only editable body text a
        # visualisation carries.
        cur.execute("ALTER TABLE dataviz_pages ADD COLUMN IF NOT EXISTS description TEXT")
        # One-time rename: "market-pulse" -> "global-heat-map". Gary decided the
        # page should carry no "Market Pulse" branding anywhere (including the
        # DB) and shouldn't have editorial posts under it — just the live widget
        # — so any old post under the old slug is dropped, then the page row
        # itself is renamed if it's still sitting under the old slug. Must run
        # before the seed-insert below, so that insert's ON CONFLICT DO NOTHING
        # naturally no-ops once this has already put "global-heat-map" in place.
        cur.execute("DELETE FROM dataviz_content WHERE page = 'market-pulse'")
        cur.execute("UPDATE dataviz_pages SET slug = 'global-heat-map', label = 'Global Heat Map' WHERE slug = 'market-pulse'")
        # Seed the 3 pages that existed before pages became self-service, so
        # any content already tagged with these slugs keeps working.
        # "global-heat-map" is the live global heat map (see MARKET_PULSE_INDICES
        # below) — seeded here too so it appears on the hub like any other
        # data-viz page, with the live map rendered above it.
        for slug, label in (("viz-1", "Visualisation 01"), ("viz-2", "Visualisation 02"), ("viz-3", "Visualisation 03"),
                             ("global-heat-map", "Global Heat Map"),
                             ("gold-silver-ratio", "Gold to Silver Ratio")):
            cur.execute("INSERT INTO dataviz_pages (slug, label) VALUES (%s, %s) ON CONFLICT (slug) DO NOTHING", (slug, label))
        # Flip the flag on for global-heat-map/gold-silver-ratio even if the row
        # already existed from before has_live_widget was added (ON CONFLICT DO
        # NOTHING above wouldn't touch it) — these UPDATEs are what actually make
        # existing deployments pick up the new behavior.
        cur.execute("UPDATE dataviz_pages SET has_live_widget = TRUE WHERE slug = 'global-heat-map'")
        cur.execute("UPDATE dataviz_pages SET has_live_widget = TRUE WHERE slug = 'gold-silver-ratio'")

VALID_INTERVALS = {"1m", "2m", "5m", "15m", "30m", "60m", "90m", "1h", "1d", "5d", "1wk", "1mo", "3mo"}

# yfinance has no native 2h/4h/10m/45m/6mo/1y bars — synthesized by fetching a smaller
# native interval and resampling. Kept separate from VALID_INTERVALS since every other
# caller of that set (elsewhere in the codebase, if any) should keep seeing only
# intervals yfinance itself understands; only _fetch_ohlcv needs to know about the
# synthetic ones.
_RESAMPLE_INTERVALS = {"2h": "1h", "4h": "1h", "10m": "5m", "45m": "15m", "6mo": "1mo", "1y": "1mo"}
# pandas resample() rule strings for each synthetic interval above — NOT the same string
# as the interval's own key: pandas 3.x rejects "10m"/"1y" outright ("m" means
# month-end, "y" was removed in favour of "YE"), so this maps each key to the actual
# offset alias pandas expects.
_RESAMPLE_RULES = {"2h": "2h", "4h": "4h", "10m": "10min", "45m": "45min", "6mo": "6MS", "1y": "1YS"}
ALL_VALID_INTERVALS = VALID_INTERVALS | set(_RESAMPLE_INTERVALS)

_INT_CALC_KEYS = {
    "rsi_length": 14, "macd_fast": 12, "macd_slow": 26, "macd_signal": 9,
    "bb_length": 20, "ema_short": 9, "ema_long": 21,
    "stoch_k": 14, "stoch_d": 3, "stoch_smooth": 3,
    "stochrsi_length": 14, "stochrsi_k": 3, "stochrsi_d": 3,
    "cci_length": 20, "willr_length": 14, "adx_length": 14,
    "atr_length": 14, "mfi_length": 14, "aroon_length": 25,
    "supertrend_length": 10, "wma_length": 20, "hma_length": 20, "roc_length": 12,
    "hs_pivot": 3, "hs_lookback": 90,
    # ── Extended set (calculate_all now exposes the same indicators/trigger-modes
    # Backtester does — same keys as _BT_INT_CALC_KEYS below, plus a few that only
    # this endpoint needs since Backtester computes them from `thresholds` instead).
    "rsi_div_lookback": 5,
    "macd_div_lookback": 5, "macd_zscore_length": 100,
    "stochrsi_div_lookback": 5,
    "willr_div_lookback": 5, "willr_confirm_lookback": 5,
    "roc_div_lookback": 5, "roc_momentum_lookback": 3,
    "mfi_div_lookback": 5,
    "tsi_div_lookback": 5,
    "ao_div_lookback": 5, "ao_twin_peaks_lookback": 5,
    "ichimoku_tenkan": 9, "ichimoku_kijun": 26, "ichimoku_senkou": 52,
    "donchian_length": 20,
    "keltner_length": 20, "keltner_atr_length": 10,
    "keltner_walk_min_consecutive": 3, "keltner_squeeze_lookback": 10,
    "stdev_length": 20,
    "chaikin_vol_ema_length": 10, "chaikin_vol_roc_length": 10,
    "hist_vol_length": 20,
    "vwap_length": 20, "vwap_anchored": 0,
    "ad_sma_length": 20, "ad_div_lookback": 5,
    "cmf_length": 20,
    "tsi_long": 25, "tsi_short": 13, "tsi_signal": 13,
    "ao_fast": 5, "ao_slow": 34,
    "obv_sma_length": 20, "obv_div_lookback": 5,
    "vol_profile_lookback": 50, "vol_profile_bins": 24,
    "fib_lookback": 50,
    "hma_slope_lookback": 3, "hma_fast_length": 9,
    "bb_squeeze_lookback": 100, "bb_breakout_window": 10,
    "bb_walk_min_consecutive": 3, "bb_pattern_lookback": 5,
    "ma_short_length": 9, "ma_medium_length": 20, "ma_long_length": 50,
    "psar_gap_lookback": 3, "supertrend_gap_lookback": 3,
    "atr_trend_lookback": 5, "stdev_trend_lookback": 5,
    "chaikin_vol_trend_lookback": 5, "hist_vol_trend_lookback": 5,
}
_FLOAT_CALC_KEYS = {
    "bb_std": 2.0, "supertrend_mult": 3.0,
    "psar_start": 0.02, "psar_inc": 0.02, "psar_max": 0.2,
    # ── Extended set ──────────────────────────────────────────────────────
    "rsi_oversold": 30.0, "rsi_overbought": 70.0,
    "willr_oversold": -80.0, "willr_overbought": -20.0,
    "keltner_mult": 2.0, "keltner_walk_tolerance_pct": 0.5,
    "bb_squeeze_percentile": 20.0, "bb_walk_tolerance_pct": 0.5,
    "vwap_band_pct": 1.0, "fib_tolerance_pct": 0.5,
}


def _extract_calc_params(args) -> dict:
    params = {}
    for key in _INT_CALC_KEYS:
        val = args.get(key)
        if val is not None:
            try:
                params[key] = int(val)
            except ValueError:
                pass
    for key in _FLOAT_CALC_KEYS:
        val = args.get(key)
        if val is not None:
            try:
                params[key] = float(val)
            except ValueError:
                pass
    smoothing = args.get("rsi_smoothing", "").strip().lower()
    if smoothing in ("wilder", "ema", "sma"):
        params["rsi_smoothing"] = smoothing
    ma_type = args.get("ma_type", "").strip().lower()
    if ma_type in ("simple", "smoothed", "exponential", "weighted", "volume_weighted"):
        params["ma_type"] = ma_type
    return params


# Backtester-only calc params (kept separate from _INT_CALC_KEYS/_FLOAT_CALC_KEYS
# above so /api/indicators' calculate_all(**calc_params) never sees an unexpected kwarg).
_BT_INT_CALC_KEYS = {
    "rsi_div_lookback": 5,
    "macd_div_lookback": 5, "macd_zscore_length": 100,
    "stochrsi_div_lookback": 5,
    "willr_div_lookback": 5, "willr_confirm_lookback": 5,
    "roc_div_lookback": 5, "roc_momentum_lookback": 3,
    "mfi_div_lookback": 5,
    "tsi_div_lookback": 5,
    "ao_div_lookback": 5, "ao_twin_peaks_lookback": 5,
    "ichimoku_tenkan": 9, "ichimoku_kijun": 26, "ichimoku_senkou": 52,
    "donchian_length": 20, "donchian_exit_length": 10,
    "keltner_length": 20, "keltner_atr_length": 10,
    "keltner_walk_min_consecutive": 3, "keltner_squeeze_lookback": 10,
    "stdev_length": 20,
    "chaikin_vol_ema_length": 10, "chaikin_vol_roc_length": 10,
    "hist_vol_length": 20,
    "vwap_length": 20, "vwap_anchored": 0,
    "ad_sma_length": 20,
    "cmf_length": 20,
    "tsi_long": 25, "tsi_short": 13, "tsi_signal": 13,
    "ao_fast": 5, "ao_slow": 34,
    "obv_sma_length": 20,
    "vol_profile_lookback": 50, "vol_profile_bins": 24,
    "fib_lookback": 50,
    "hma_slope_lookback": 3, "hma_fast_length": 9,
    "bb_squeeze_lookback": 100, "bb_breakout_window": 10,
    "bb_walk_min_consecutive": 3, "bb_pattern_lookback": 5,
    "ma_short_length": 9, "ma_medium_length": 20, "ma_long_length": 50,
    "obv_div_lookback": 5, "ad_div_lookback": 5,
}
_BT_FLOAT_CALC_KEYS = {
    "keltner_mult": 2.0, "keltner_walk_tolerance_pct": 0.5,
    "bb_squeeze_percentile": 20.0, "bb_walk_tolerance_pct": 0.5,
    "vwap_band_pct": 1.0,
}

_TRIGGER_WHITELISTS = {
    "rsi_trigger":         {"overbought_oversold", "overbought", "oversold", "centerline_cross",
                             "bullish_divergence", "bearish_divergence", "failure_swings"},
    "macd_trigger":        {"signal_cross", "bullish_signal_cross", "bearish_signal_cross", "centerline_cross",
                             "bullish_divergence", "bearish_divergence", "histogram_reversal", "overbought", "oversold"},
    "bb_trigger":          {"percent_b", "upper_touch", "lower_touch", "volatility_breakout",
                             "walking_upper", "walking_lower", "w_bottom", "m_top",
                             "breakout_margin", "pct_below_high", "pct_above_low"},
    "ma_trigger":          {"dual_cross", "price_cross", "two_ma_bull", "two_ma_bear", "three_ma_bull", "three_ma_bear"},
    "adx_trigger":         {"trend_threshold", "bull_di_cross", "bear_di_cross", "above_25", "above_50", "above_75",
                             "strong_di_plus", "strong_di_minus"},
    "psar_trigger":        {"flip", "bull_flip", "bear_flip", "trend_state", "trailing_stop"},
    "ichimoku_trigger":    {"cloud_position", "bullish", "bearish", "tk_cross"},
    "supertrend_trigger":  {"flip", "bull_flip", "bear_flip", "trend_state", "trailing_stop"},
    "donchian_trigger":    {"breakout", "bullish", "bearish", "middle_cross", "two_channel_bull", "two_channel_bear",
                             "resistance_retest", "support_retest"},
    "hma_trigger":         {"slope", "bullish_slope", "bearish_slope", "price_cross", "two_hma_bull", "two_hma_bear"},
    "stoch_trigger":       {"overbought_oversold", "overbought", "oversold", "signal_cross"},
    "stochrsi_trigger":    {"overbought_oversold", "overbought", "oversold", "signal_cross", "bullish_divergence", "bearish_divergence"},
    "cci_trigger":         {"overbought_oversold", "overbought", "oversold", "centerline_cross", "breakout_bull", "breakout_bear"},
    "willr_trigger":       {"overbought_oversold", "overbought", "oversold", "midline_cross",
                             "momentum_failure_bull", "momentum_failure_bear",
                             "trend_confirmation_bull", "trend_confirmation_bear",
                             "bullish_divergence", "bearish_divergence"},
    "roc_trigger":         {"threshold", "bullish", "bearish", "centerline_cross",
                             "bull_momentum", "bear_momentum",
                             "bullish_divergence", "bearish_divergence"},
    "mfi_trigger":         {"overbought_oversold", "overbought", "oversold", "centerline_cross",
                             "bullish_divergence", "bearish_divergence"},
    "tsi_trigger":         {"signal_cross", "bullish", "bearish", "centerline_cross",
                             "overbought", "oversold",
                             "bullish_divergence", "bearish_divergence"},
    "ao_trigger":          {"zero_state", "bullish", "bearish", "zero_cross",
                             "bull_saucer", "bear_saucer",
                             "bull_twin_peaks", "bear_twin_peaks",
                             "bull_divergence", "bear_divergence"},
    "atr_trigger":         {"expansion", "bullish_expansion", "bearish_expansion", "contraction"},
    "keltner_trigger":     {"breakout", "bullish", "bearish", "middle_cross",
                             "bull_band_riding", "bear_band_riding",
                             "bull_mean_reversion", "bear_mean_reversion",
                             "keltner_squeeze"},
    "stdev_trigger":       {"expansion", "bullish_expansion", "bearish_expansion", "contraction"},
    "chaikin_vol_trigger": {"expansion", "bullish_expansion", "bearish_expansion", "contraction"},
    "hist_vol_trigger":    {"expansion", "bullish_expansion", "bearish_expansion", "contraction"},
    "obv_trigger":         {"trend", "bullish", "bearish", "divergence"},
    "vwap_trigger":        {"position", "bullish", "bearish", "band_touch", "pullback_buy", "pullback_sell"},
    "ad_trigger":          {"trend", "bullish", "bearish", "divergence"},
    "cmf_trigger":         {"threshold", "bullish", "bearish", "centerline_cross"},
    "vol_profile_trigger": {"position", "bullish", "bearish", "poc_breakout"},
    "fib_trigger":         {"bounce_reject", "bullish_bounce", "bearish_reject", "any_touch"},
    "inv_hs_trigger":      {"neckline_touch", "neckline_break"},
}


def _extract_backtest_calc_params(args) -> dict:
    params = {}
    for key in _BT_INT_CALC_KEYS:
        val = args.get(key)
        if val is not None:
            try:
                params[key] = int(val)
            except ValueError:
                pass
    for key in _BT_FLOAT_CALC_KEYS:
        val = args.get(key)
        if val is not None:
            try:
                params[key] = float(val)
            except ValueError:
                pass
    ma_type = args.get("ma_type", "").strip().lower()
    if ma_type in ("simple", "smoothed", "exponential", "weighted", "volume_weighted"):
        params["ma_type"] = ma_type
    return params


VALID_PERIODS = {"1d", "5d", "60d", "1mo", "3mo", "6mo", "1y", "2y", "5y", "10y", "ytd", "max"}

TIER_RANKS = {"basic": 0, "signal_tester": 1, "power_user": 2}


# ── User model ───────────────────────────────────────────────────────────────

class User(UserMixin):
    def __init__(self, username: str, tier: str = "basic", alpha_role: str | None = None):
        self.id = username
        self.tier = tier
        self.alpha_role = alpha_role


@login_manager.user_loader
def load_user(username: str):
    users = _load_users()
    if username in users:
        u = users[username]
        return User(username, u.get("tier", "basic"), u.get("alpha_role"))
    return None


@login_manager.unauthorized_handler
def unauthorized():
    return jsonify({"error": "Authentication required"}), 401


def tier_required(min_tier: str):
    def decorator(f):
        @wraps(f)
        def wrapped(*args, **kwargs):
            if not current_user.is_authenticated:
                return jsonify({"error": "Login required", "tier_required": min_tier}), 401
            user_rank = TIER_RANKS.get(getattr(current_user, "tier", "basic"), 0)
            if user_rank < TIER_RANKS.get(min_tier, 0):
                return jsonify({
                    "error": "Subscription upgrade required",
                    "tier_required": min_tier,
                    "current_tier": getattr(current_user, "tier", "basic"),
                }), 403
            return f(*args, **kwargs)
        return wrapped
    return decorator


def alpha_author_required(f):
    """Gates an endpoint to users with an assigned alpha_role (one of the 4 partners)."""
    @wraps(f)
    def wrapped(*args, **kwargs):
        if not current_user.is_authenticated:
            return jsonify({"error": "Login required"}), 401
        if not getattr(current_user, "alpha_role", None):
            return jsonify({"error": "This account has no Alpha author access"}), 403
        return f(*args, **kwargs)
    return wrapped


DATAVIZ_AUTHORS = {"tom", "gary"}


def dataviz_author_required(f):
    """Gates the Data Visualisation studio to authorized accounts (alpha_role in DATAVIZ_AUTHORS)."""
    @wraps(f)
    def wrapped(*args, **kwargs):
        if not current_user.is_authenticated:
            return jsonify({"error": "Login required"}), 401
        if getattr(current_user, "alpha_role", None) not in DATAVIZ_AUTHORS:
            return jsonify({"error": "This account has no Data Visualisation author access"}), 403
        return f(*args, **kwargs)
    return wrapped


# ── User-store helpers ───────────────────────────────────────────────────────

def _load_users() -> dict:
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute("SELECT username, password_hash, preferences, tier, profile, alpha_role FROM users")
            return {
                row["username"]: {
                    "password_hash": row["password_hash"],
                    "preferences":   row["preferences"] or {},
                    "tier":          row.get("tier", "basic"),
                    "profile":       row.get("profile") or {},
                    "alpha_role":    row.get("alpha_role"),
                }
                for row in cur.fetchall()
            }
    if not os.path.exists(USERS_FILE):
        return {}
    with open(USERS_FILE, "r") as f:
        return json.load(f).get("users", {})


def _save_users(users: dict) -> None:
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor() as cur:
            for username, data in users.items():
                cur.execute("""
                    INSERT INTO users (username, password_hash, preferences, tier, profile, alpha_role)
                    VALUES (%s, %s, %s, %s, %s, %s)
                    ON CONFLICT (username) DO UPDATE
                        SET password_hash = EXCLUDED.password_hash,
                            preferences   = EXCLUDED.preferences,
                            tier          = EXCLUDED.tier,
                            profile       = EXCLUDED.profile,
                            alpha_role    = EXCLUDED.alpha_role
                """, (username, data["password_hash"], json.dumps(data.get("preferences", {})), data.get("tier", "basic"), json.dumps(data.get("profile", {})), data.get("alpha_role")))
        return
    with open(USERS_FILE, "w") as f:
        json.dump({"users": users}, f, indent=2)


def get_user_avatar(username: str):
    """Returns (filename, bytes) or (None, None). Kept out of _load_users()'s
    Postgres SELECT so that hot path (hit on every authenticated request)
    never has to move an image blob around."""
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute("SELECT avatar_filename, avatar_file FROM users WHERE username = %s", (username,))
            row = cur.fetchone()
            if not row or row["avatar_file"] is None:
                return None, None
            return row["avatar_filename"], bytes(row["avatar_file"])
    users = _load_users()
    user = users.get(username)
    if not user or not user.get("avatar_file"):
        return None, None
    return user.get("avatar_filename"), base64.b64decode(user["avatar_file"])


def set_user_avatar(username: str, filename: str, file_bytes: bytes) -> None:
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor() as cur:
            cur.execute(
                "UPDATE users SET avatar_file = %s, avatar_filename = %s WHERE username = %s",
                (psycopg2.Binary(file_bytes), filename, username),
            )
        return
    users = _load_users()
    if username not in users:
        return
    users[username]["avatar_file"] = base64.b64encode(file_bytes).decode("ascii")
    users[username]["avatar_filename"] = filename
    _save_users(users)


# ── Alpha content store ──────────────────────────────────────────────────────
# Mirrors the _load_users/_save_users dual-path pattern above: Postgres when
# DATABASE_URL is set, a local JSON file otherwise. source_file (the original
# uploaded document) is never included in list/get results — only
# alpha_content_get_file() fetches it, so listing drafts never has to move a
# binary blob around.

_ALPHA_CONTENT_FIELDS = [
    "author", "kind", "status", "topic", "level", "title", "subtitle", "snippet", "body", "stance", "url",
    "source_kind", "source_filename", "source_text",
    "image_url", "image_filename", "pinned",
    "related_lesson_slug", "related_lesson_note",
    # Pending edits to a *published* item, held back from the live page until the
    # author unpublishes (which folds them in) and re-publishes. Studio-only —
    # the public endpoints whitelist their output via _ALPHA_PUBLIC_FIELDS, so
    # this never leaks. Shape: {title, subtitle, snippet, body, topic, url,
    # stance, image_url} — a subset, only the fields that were edited.
    "staged_edits",
]


def _alpha_row_to_dict(row: dict) -> dict:
    d = {k: row.get(k) for k in ["id", *_ALPHA_CONTENT_FIELDS]}
    for key in ("created_at", "updated_at", "published_at"):
        val = row.get(key)
        d[key] = val.isoformat() if val is not None else None
    return d


def _load_alpha_content_json() -> dict:
    if not os.path.exists(ALPHA_CONTENT_FILE):
        return {"items": [], "next_id": 1}
    with open(ALPHA_CONTENT_FILE, "r") as f:
        return json.load(f)


def _save_alpha_content_json(data: dict) -> None:
    with open(ALPHA_CONTENT_FILE, "w") as f:
        json.dump(data, f, indent=2)


def alpha_content_list(author: str | None = None, status: str | None = None) -> list:
    if DATABASE_URL:
        query = f"SELECT id, {', '.join(_ALPHA_CONTENT_FIELDS)}, created_at, updated_at, published_at FROM alpha_content WHERE 1=1"
        params = []
        if author:
            query += " AND author = %s"
            params.append(author)
        if status:
            query += " AND status = %s"
            params.append(status)
        query += " ORDER BY created_at DESC"
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(query, params)
            return [_alpha_row_to_dict(row) for row in cur.fetchall()]
    data = _load_alpha_content_json()
    items = data["items"]
    if author:
        items = [i for i in items if i.get("author") == author]
    if status:
        items = [i for i in items if i.get("status") == status]
    items = sorted(items, key=lambda i: i.get("created_at") or "", reverse=True)
    return [{**{k: v for k, v in i.items() if k not in ("source_file", "image_file")}, "pinned": bool(i.get("pinned"))} for i in items]


def alpha_content_get(item_id: int) -> dict | None:
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(
                f"SELECT id, {', '.join(_ALPHA_CONTENT_FIELDS)}, created_at, updated_at, published_at "
                f"FROM alpha_content WHERE id = %s", (item_id,)
            )
            row = cur.fetchone()
            return _alpha_row_to_dict(row) if row else None
    data = _load_alpha_content_json()
    for item in data["items"]:
        if item.get("id") == item_id:
            return {**{k: v for k, v in item.items() if k not in ("source_file", "image_file")}, "pinned": bool(item.get("pinned"))}
    return None


def alpha_content_get_file(item_id: int):
    """Returns (filename, bytes) or (None, None) if there's no attached file."""
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute("SELECT source_filename, source_file FROM alpha_content WHERE id = %s", (item_id,))
            row = cur.fetchone()
            if not row or row["source_file"] is None:
                return None, None
            return row["source_filename"], bytes(row["source_file"])
    data = _load_alpha_content_json()
    for item in data["items"]:
        if item.get("id") == item_id:
            b64 = item.get("source_file")
            if not b64:
                return None, None
            return item.get("source_filename"), base64.b64decode(b64)
    return None, None


def alpha_content_get_image(item_id: int):
    """Returns (filename, bytes) or (None, None) if there's no uploaded image file."""
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute("SELECT image_filename, image_file FROM alpha_content WHERE id = %s", (item_id,))
            row = cur.fetchone()
            if not row or row["image_file"] is None:
                return None, None
            return row["image_filename"], bytes(row["image_file"])
    data = _load_alpha_content_json()
    for item in data["items"]:
        if item.get("id") == item_id:
            b64 = item.get("image_file")
            if not b64:
                return None, None
            return item.get("image_filename"), base64.b64decode(b64)
    return None, None


def alpha_content_set_image(item_id: int, filename: str, file_bytes: bytes) -> dict | None:
    """Sets the uploaded image file, clearing image_url to enforce file-vs-URL exclusivity."""
    now = _dt.datetime.utcnow()
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(
                f"UPDATE alpha_content SET image_file = %s, image_filename = %s, image_url = NULL, "
                f"updated_at = %s WHERE id = %s "
                f"RETURNING id, {', '.join(_ALPHA_CONTENT_FIELDS)}, created_at, updated_at, published_at",
                (psycopg2.Binary(file_bytes), filename, now, item_id),
            )
            row = cur.fetchone()
            return _alpha_row_to_dict(row) if row else None
    data = _load_alpha_content_json()
    for item in data["items"]:
        if item.get("id") == item_id:
            item["image_file"] = base64.b64encode(file_bytes).decode("ascii")
            item["image_filename"] = filename
            item["image_url"] = None
            item["updated_at"] = now.isoformat()
            _save_alpha_content_json(data)
            return {k: v for k, v in item.items() if k not in ("source_file", "image_file")}
    return None


# ── Inline post attachments ──────────────────────────────────────────────────
# A post's main image is a single slot on the alpha_content row itself; these
# are additional images a writer drops into the middle of a post's body via
# the Studio's Insert Image button, so a post can hold more than one.

def _load_alpha_attachments_json() -> dict:
    if not os.path.exists(ALPHA_ATTACHMENTS_FILE):
        return {"items": [], "next_id": 1}
    with open(ALPHA_ATTACHMENTS_FILE, "r") as f:
        return json.load(f)


def _save_alpha_attachments_json(data: dict) -> None:
    with open(ALPHA_ATTACHMENTS_FILE, "w") as f:
        json.dump(data, f, indent=2)


def alpha_attachment_create(content_id: int, filename: str, file_bytes: bytes) -> int:
    """Returns the new attachment's id."""
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor() as cur:
            cur.execute(
                "INSERT INTO alpha_content_attachment (content_id, filename, file) VALUES (%s, %s, %s) RETURNING id",
                (content_id, filename, psycopg2.Binary(file_bytes)),
            )
            return cur.fetchone()[0]
    data = _load_alpha_attachments_json()
    new_id = data.get("next_id", 1)
    data["items"].append({
        "id": new_id,
        "content_id": content_id,
        "filename": filename,
        "file": base64.b64encode(file_bytes).decode("ascii"),
    })
    data["next_id"] = new_id + 1
    _save_alpha_attachments_json(data)
    return new_id


def alpha_attachment_get(attachment_id: int):
    """Returns (content_id, filename, bytes) or (None, None, None)."""
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute("SELECT content_id, filename, file FROM alpha_content_attachment WHERE id = %s", (attachment_id,))
            row = cur.fetchone()
            if not row:
                return None, None, None
            return row["content_id"], row["filename"], bytes(row["file"])
    data = _load_alpha_attachments_json()
    for item in data["items"]:
        if item.get("id") == attachment_id:
            return item.get("content_id"), item.get("filename"), base64.b64decode(item["file"])
    return None, None, None


def alpha_content_create(fields: dict, file_bytes: bytes | None = None) -> dict:
    now = _dt.datetime.utcnow()
    if DATABASE_URL:
        cols = [*_ALPHA_CONTENT_FIELDS, "source_file", "created_at", "updated_at"]
        vals = [(bool(fields.get(k)) if k == "pinned" else fields.get(k)) for k in _ALPHA_CONTENT_FIELDS] + [
            psycopg2.Binary(file_bytes) if file_bytes else None, now, now
        ]
        placeholders = ", ".join(["%s"] * len(vals))
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(
                f"INSERT INTO alpha_content ({', '.join(cols)}) VALUES ({placeholders}) "
                f"RETURNING id, {', '.join(_ALPHA_CONTENT_FIELDS)}, created_at, updated_at, published_at",
                vals,
            )
            row = cur.fetchone()
            return _alpha_row_to_dict(row)
    data = _load_alpha_content_json()
    new_id = data.get("next_id", 1)
    item = {k: fields.get(k) for k in _ALPHA_CONTENT_FIELDS}
    item["pinned"] = bool(item.get("pinned"))
    item["id"] = new_id
    item["created_at"] = now.isoformat()
    item["updated_at"] = now.isoformat()
    item["published_at"] = None
    if file_bytes:
        item["source_file"] = base64.b64encode(file_bytes).decode("ascii")
    data["items"].append(item)
    data["next_id"] = new_id + 1
    _save_alpha_content_json(data)
    return {k: v for k, v in item.items() if k not in ("source_file", "image_file")}


def alpha_content_update(item_id: int, updates: dict) -> dict | None:
    now = _dt.datetime.utcnow()
    allowed = set(_ALPHA_CONTENT_FIELDS) | {"published_at", "image_file"}
    updates = {k: v for k, v in updates.items() if k in allowed}
    if DATABASE_URL:
        if not updates:
            return alpha_content_get(item_id)
        set_clauses = [f"{k} = %s" for k in updates] + ["updated_at = %s"]
        # staged_edits is a JSONB column — wrap the dict so psycopg2 adapts it.
        vals = [
            psycopg2.extras.Json(v) if k == "staged_edits" and v is not None else v
            for k, v in updates.items()
        ] + [now, item_id]
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(
                f"UPDATE alpha_content SET {', '.join(set_clauses)} WHERE id = %s "
                f"RETURNING id, {', '.join(_ALPHA_CONTENT_FIELDS)}, created_at, updated_at, published_at",
                vals,
            )
            row = cur.fetchone()
            return _alpha_row_to_dict(row) if row else None
    data = _load_alpha_content_json()
    json_updates = dict(updates)
    if "published_at" in json_updates:
        val = json_updates["published_at"]
        json_updates["published_at"] = val.isoformat() if hasattr(val, "isoformat") else val
    for item in data["items"]:
        if item.get("id") == item_id:
            item.update(json_updates)
            item["updated_at"] = now.isoformat()
            _save_alpha_content_json(data)
            return {k: v for k, v in item.items() if k not in ("source_file", "image_file")}
    return None


def alpha_content_delete(item_id: int) -> bool:
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor() as cur:
            cur.execute("DELETE FROM alpha_content_attachment WHERE content_id = %s", (item_id,))
            cur.execute("DELETE FROM alpha_content WHERE id = %s", (item_id,))
            return cur.rowcount > 0
    attachments = _load_alpha_attachments_json()
    attachments["items"] = [a for a in attachments["items"] if a.get("content_id") != item_id]
    _save_alpha_attachments_json(attachments)
    data = _load_alpha_content_json()
    before = len(data["items"])
    data["items"] = [i for i in data["items"] if i.get("id") != item_id]
    _save_alpha_content_json(data)
    return len(data["items"]) < before


# ── Data Visualisation pages ──────────────────────────────────────────────────
# Self-service pages Tom/Gary spin up from the studio — each is a slug/label
# pair that dataviz_content.page points at. JSON-fallback mirrors the seed
# rows inserted by _ensure_table so dev-mode (no DATABASE_URL) still has the
# original 3 pages available.

_SEED_DATAVIZ_PAGES = [
    {"slug": "viz-1", "label": "Visualisation 01", "author": None, "has_live_widget": False},
    {"slug": "viz-2", "label": "Visualisation 02", "author": None, "has_live_widget": False},
    {"slug": "viz-3", "label": "Visualisation 03", "author": None, "has_live_widget": False},
    {"slug": "global-heat-map", "label": "Global Heat Map", "author": None, "has_live_widget": True},
    {"slug": "gold-silver-ratio", "label": "Gold to Silver Ratio", "author": None, "has_live_widget": True},
]


def _load_dataviz_pages_json() -> dict:
    if not os.path.exists(DATAVIZ_PAGES_FILE):
        return {"pages": list(_SEED_DATAVIZ_PAGES)}
    with open(DATAVIZ_PAGES_FILE, "r") as f:
        return json.load(f)


def _save_dataviz_pages_json(data: dict) -> None:
    with open(DATAVIZ_PAGES_FILE, "w") as f:
        json.dump(data, f, indent=2)


def _slugify(label: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", label.lower()).strip("-")
    return slug or "page"


def dataviz_pages_list() -> list:
    # Pages with a live widget (an actual built visualization, not a placeholder
    # waiting for posts) sort first on the hub page — Market Pulse is the first
    # of these, and it should read as the flagship item, not get buried below
    # the still-empty viz-1/2/3 placeholders just because those were seeded
    # earlier.
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute("SELECT slug, label, description, author, has_live_widget, header_image_filename, created_at FROM dataviz_pages "
                        "ORDER BY has_live_widget DESC, created_at ASC")
            return [
                {"slug": r["slug"], "label": r["label"], "description": r["description"], "author": r["author"],
                 "has_live_widget": bool(r["has_live_widget"]), "has_header_image": bool(r["header_image_filename"]),
                 "created_at": r["created_at"].isoformat() if r["created_at"] else None}
                for r in cur.fetchall()
            ]
    data = _load_dataviz_pages_json()
    pages = sorted(data["pages"], key=lambda p: (0 if p.get("has_live_widget") else 1, p.get("created_at") or ""))
    return [{**p, "has_live_widget": bool(p.get("has_live_widget")), "has_header_image": bool(p.get("header_image_filename"))} for p in pages]


def dataviz_page_get_header_image(slug: str):
    """Returns (filename, bytes) or (None, None)."""
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute("SELECT header_image_filename, header_image_file FROM dataviz_pages WHERE slug = %s", (slug,))
            row = cur.fetchone()
            if not row:
                return None, None
            return row["header_image_filename"], (bytes(row["header_image_file"]) if row["header_image_file"] else None)
    data = _load_dataviz_pages_json()
    for p in data["pages"]:
        if p["slug"] == slug:
            file_hex = p.get("header_image_file")
            return p.get("header_image_filename"), (bytes.fromhex(file_hex) if file_hex else None)
    return None, None


def dataviz_page_set_header_image(slug: str, filename: str | None, file_bytes: bytes | None) -> bool:
    """Pass filename=None to clear the header image."""
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor() as cur:
            cur.execute(
                "UPDATE dataviz_pages SET header_image_filename = %s, header_image_file = %s WHERE slug = %s",
                (filename, psycopg2.Binary(file_bytes) if file_bytes else None, slug),
            )
            return cur.rowcount > 0
    data = _load_dataviz_pages_json()
    for p in data["pages"]:
        if p["slug"] == slug:
            p["header_image_filename"] = filename
            p["header_image_file"] = file_bytes.hex() if file_bytes else None
            _save_dataviz_pages_json(data)
            return True
    return False


def dataviz_page_get(slug: str) -> dict | None:
    for page in dataviz_pages_list():
        if page["slug"] == slug:
            return page
    return None


def dataviz_page_create(label: str, author: str, description: str | None = None, has_live_widget: bool = False) -> dict:
    now = _dt.datetime.utcnow()
    base_slug = _slugify(label)
    existing_slugs = {p["slug"] for p in dataviz_pages_list()}
    slug = base_slug
    n = 2
    while slug in existing_slugs:
        slug = f"{base_slug}-{n}"
        n += 1
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(
                "INSERT INTO dataviz_pages (slug, label, description, author, has_live_widget, created_at) VALUES (%s, %s, %s, %s, %s, %s) "
                "RETURNING slug, label, description, author, has_live_widget, created_at",
                (slug, label, description, author, has_live_widget, now),
            )
            row = cur.fetchone()
            return {"slug": row["slug"], "label": row["label"], "description": row["description"], "author": row["author"],
                     "has_live_widget": bool(row["has_live_widget"]), "created_at": row["created_at"].isoformat()}
    data = _load_dataviz_pages_json()
    page = {"slug": slug, "label": label, "description": description, "author": author, "has_live_widget": has_live_widget, "created_at": now.isoformat()}
    data["pages"].append(page)
    _save_dataviz_pages_json(data)
    return page


def dataviz_page_update(slug: str, updates: dict) -> dict | None:
    """updates may contain 'label' and/or 'description'."""
    if not updates:
        return dataviz_page_get(slug)
    if DATABASE_URL:
        set_clauses = ", ".join(f"{k} = %s" for k in updates)
        with _db_conn() as conn, conn.cursor() as cur:
            cur.execute(f"UPDATE dataviz_pages SET {set_clauses} WHERE slug = %s", (*updates.values(), slug))
            if cur.rowcount == 0:
                return None
        return dataviz_page_get(slug)
    data = _load_dataviz_pages_json()
    for p in data["pages"]:
        if p["slug"] == slug:
            p.update(updates)
            _save_dataviz_pages_json(data)
            return p
    return None


def dataviz_page_delete(slug: str) -> bool:
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor() as cur:
            cur.execute("DELETE FROM dataviz_pages WHERE slug = %s", (slug,))
            return cur.rowcount > 0
    data = _load_dataviz_pages_json()
    before = len(data["pages"])
    data["pages"] = [p for p in data["pages"] if p["slug"] != slug]
    _save_dataviz_pages_json(data)
    return len(data["pages"]) < before


# ── Market Pulse — live global heat map ──────────────────────────────────────
# Powers the "market-pulse" data-viz page: a world map of major indices, colored
# by today's % move. Uses yfinance (already a dependency, no new API key needed)
# rather than Finnhub/etc — one batched download covers every symbol below in a
# single round trip. lat/lon are the rough geographic center of each market, used
# by the frontend to place a dot on an equirectangular (x=lon, y=lat) map.
MARKET_PULSE_INDICES = [
    # A handful of countries have a second index big enough to matter on its own —
    # judgment call, not exhaustive. Each second entry is listed BEFORE its
    # country's primary/broader index: the choropleth and hover tooltip only show
    # one color per country (keyed by iso2), and later entries win ties, so this
    # ordering makes the more широко-recognized index (S&P 500, not Nasdaq; SSE
    # Composite, not Shenzhen; FTSE 100, not FTSE 250) the one that "wins" the
    # country's fill color the more widely-recognized index (S&P 500, not Nasdaq;
    # SSE Composite, not Shenzhen; FTSE 100, not FTSE 250). Both indices still
    # appear individually on the simple dot map, the ticker strip, and the
    # gainers/losers leaderboard.
    {"country": "United States", "index": "Nasdaq Composite", "symbol": "^IXIC",    "lat": 37.4,  "lon": -122.1, "iso2": "us"},
    {"country": "United States", "index": "S&P 500",        "symbol": "^GSPC",    "lat": 39.8,  "lon": -98.6,  "iso2": "us"},
    {"country": "Canada",        "index": "S&P/TSX",         "symbol": "^GSPTSE",  "lat": 56.1,  "lon": -106.3, "iso2": "ca"},
    {"country": "Mexico",        "index": "IPC",              "symbol": "^MXX",     "lat": 23.6,  "lon": -102.5, "iso2": "mx"},
    {"country": "Brazil",        "index": "Bovespa",          "symbol": "^BVSP",    "lat": -14.2, "lon": -51.9,  "iso2": "br"},
    {"country": "Argentina",     "index": "Merval",           "symbol": "^MERV",    "lat": -38.4, "lon": -63.6,  "iso2": "ar"},
    {"country": "United Kingdom","index": "FTSE 250",         "symbol": "^FTMC",    "lat": 52.5,  "lon": -1.9,   "iso2": "gb"},
    {"country": "United Kingdom","index": "FTSE 100",         "symbol": "^FTSE",    "lat": 55.4,  "lon": -3.4,   "iso2": "gb"},
    {"country": "Germany",       "index": "DAX",               "symbol": "^GDAXI",   "lat": 51.2,  "lon": 10.5,   "iso2": "de"},
    {"country": "France",        "index": "CAC 40",            "symbol": "^FCHI",    "lat": 46.6,  "lon": 2.2,    "iso2": "fr"},
    {"country": "Italy",         "index": "FTSE MIB",          "symbol": "FTSEMIB.MI","lat": 41.9, "lon": 12.6,   "iso2": "it"},
    {"country": "Spain",         "index": "IBEX 35",           "symbol": "^IBEX",    "lat": 40.5,  "lon": -3.7,   "iso2": "es"},
    {"country": "Netherlands",   "index": "AEX",               "symbol": "^AEX",     "lat": 52.1,  "lon": 5.3,    "iso2": "nl"},
    {"country": "Switzerland",   "index": "SMI",               "symbol": "^SSMI",    "lat": 46.8,  "lon": 8.2,    "iso2": "ch"},
    {"country": "Sweden",        "index": "OMXS30",            "symbol": "^OMX",     "lat": 60.1,  "lon": 18.6,   "iso2": "se"},
    {"country": "Poland",        "index": "WIG20",             "symbol": "WIG20.WA", "lat": 52.0,  "lon": 19.1,   "iso2": "pl"},
    {"country": "Turkey",        "index": "BIST 100",          "symbol": "XU100.IS", "lat": 38.9,  "lon": 35.2,   "iso2": "tr"},
    {"country": "Russia",        "index": "MOEX",              "symbol": "IMOEX.ME", "lat": 61.5,  "lon": 105.3,  "iso2": "ru"},
    {"country": "South Africa",  "index": "JSE Top 40",        "symbol": "^J203.JO", "lat": -29.0, "lon": 24.0,   "iso2": "za"},
    {"country": "Saudi Arabia",  "index": "TASI",              "symbol": "^TASI.SR", "lat": 24.0,  "lon": 45.0,   "iso2": "sa"},
    # 1306.T (Nomura's TOPIX-tracking ETF) stands in for TOPIX itself — no native
    # ^TOPX-style symbol returns data on free yfinance, tried several variants.
    {"country": "Japan",         "index": "TOPIX (ETF proxy)", "symbol": "1306.T",   "lat": 34.7,  "lon": 135.5,  "iso2": "jp"},
    {"country": "Japan",         "index": "Nikkei 225",        "symbol": "^N225",    "lat": 36.2,  "lon": 138.3,  "iso2": "jp"},
    {"country": "China",         "index": "Shenzhen Component","symbol": "399001.SZ","lat": 22.5,  "lon": 114.1,  "iso2": "cn"},
    {"country": "China",         "index": "SSE Composite",     "symbol": "000001.SS","lat": 31.2,  "lon": 121.5,  "iso2": "cn"},
    {"country": "Hong Kong",     "index": "Hang Seng",         "symbol": "^HSI",     "lat": 22.3,  "lon": 114.2,  "iso2": "hk"},
    {"country": "South Korea",   "index": "KOSDAQ",            "symbol": "^KQ11",    "lat": 37.5,  "lon": 126.7,  "iso2": "kr"},
    {"country": "South Korea",   "index": "KOSPI",             "symbol": "^KS11",    "lat": 36.5,  "lon": 127.9,  "iso2": "kr"},
    {"country": "Taiwan",        "index": "TAIEX",             "symbol": "^TWII",    "lat": 23.7,  "lon": 121.0,  "iso2": "tw"},
    {"country": "India",         "index": "Nifty 50",          "symbol": "^NSEI",    "lat": 19.1,  "lon": 72.9,   "iso2": "in"},
    {"country": "India",         "index": "BSE Sensex",        "symbol": "^BSESN",   "lat": 20.6,  "lon": 78.9,   "iso2": "in"},
    {"country": "Singapore",     "index": "STI",               "symbol": "^STI",     "lat": 1.35,  "lon": 103.8,  "iso2": "sg"},
    {"country": "Indonesia",     "index": "IDX Composite",     "symbol": "^JKSE",    "lat": -0.8,  "lon": 113.9,  "iso2": "id"},
    {"country": "Malaysia",      "index": "FTSE Bursa KLCI",   "symbol": "^KLSE",    "lat": 4.2,   "lon": 101.9,  "iso2": "my"},
    {"country": "Thailand",      "index": "SET",               "symbol": "^SET.BK",  "lat": 15.9,  "lon": 100.9,  "iso2": "th"},
    {"country": "Australia",     "index": "ASX 200",           "symbol": "^AXJO",    "lat": -25.3, "lon": 133.8,  "iso2": "au"},
    {"country": "New Zealand",   "index": "NZX 50",            "symbol": "^NZ50",    "lat": -41.0, "lon": 174.9,  "iso2": "nz"},
    # Added to broaden coverage toward "top ~50 markets by cap" — each symbol below
    # was individually confirmed against yfinance before being added. A lot of
    # smaller markets (Pakistan, Chile, Colombia, Czechia, Hungary, Qatar, Kuwait,
    # Nigeria, Kenya, Ghana, Mauritius, Morocco, most of the Balkans) simply have
    # no reliable free Yahoo Finance symbol — every symbol/suffix variant tried
    # came back empty — so they're left out rather than added as an entry that
    # would silently never show data. Russia (IMOEX.ME, below) is in the same
    # position in practice — the symbol exists but Yahoo has carried no usable
    # data for it since Western sanctions-era delistings — kept in the list on
    # the chance that changes, rather than removed.
    {"country": "Denmark",       "index": "OMX Copenhagen 25", "symbol": "^OMXC25",  "lat": 56.0,  "lon": 10.0,   "iso2": "dk"},
    {"country": "Norway",        "index": "Oslo Børs",         "symbol": "OSEBX.OL", "lat": 60.5,  "lon": 8.5,    "iso2": "no"},
    {"country": "Finland",       "index": "OMX Helsinki 25",   "symbol": "^OMXH25",  "lat": 61.9,  "lon": 25.7,   "iso2": "fi"},
    {"country": "Belgium",       "index": "BEL 20",            "symbol": "^BFX",     "lat": 50.5,  "lon": 4.5,    "iso2": "be"},
    {"country": "Austria",       "index": "ATX",               "symbol": "^ATX",     "lat": 47.5,  "lon": 14.5,   "iso2": "at"},
    {"country": "Ireland",       "index": "ISEQ",              "symbol": "^ISEQ",    "lat": 53.4,  "lon": -8.0,   "iso2": "ie"},
    {"country": "Portugal",      "index": "PSI 20",            "symbol": "PSI20.LS", "lat": 39.4,  "lon": -8.0,   "iso2": "pt"},
    {"country": "Greece",        "index": "Athens General",    "symbol": "GD.AT",    "lat": 39.0,  "lon": 22.0,   "iso2": "gr"},
    {"country": "Israel",        "index": "TA-35",             "symbol": "TA35.TA",  "lat": 31.0,  "lon": 35.0,   "iso2": "il"},
    {"country": "United Arab Emirates", "index": "DFM General", "symbol": "DFMGI.AE","lat": 25.2,  "lon": 55.3,   "iso2": "ae"},
    {"country": "Philippines",   "index": "PSEi Composite",    "symbol": "PSEI.PS",  "lat": 13.0,  "lon": 122.0,  "iso2": "ph"},
    # Vietnam has no native index available on free yfinance (every ^VNINDEX /
    # VNINDEX.VN / VNI variant came back empty) — VNM (the VanEck Vietnam ETF, a
    # US-listed fund of Vietnamese equities) is used instead as the closest
    # available proxy for "how is the Vietnamese market doing today".
    {"country": "Vietnam",       "index": "Vietnam ETF (VNM)", "symbol": "VNM",      "lat": 14.0,  "lon": 108.0,  "iso2": "vn"},
    # Africa: South Africa (JSE, above) is the only individual African market with
    # a working free index symbol — Nigeria, Egypt, Kenya, Morocco and Ghana were
    # all tried under several symbol variants each and none returned data. AFK (the
    # VanEck Africa Index ETF) is added as a broad-continent proxy instead of a
    # 2nd/3rd/4th single-country pick; it spans multiple countries at once, so
    # iso2 is deliberately blank — it shows on the simple dot map and in the
    # leaderboard/ticker like any other market, but isn't eligible to color any
    # one country on the Detailed choropleth (which is keyed strictly by iso2).
    {"country": "Africa (broad)","index": "Africa ETF (AFK)",  "symbol": "AFK",      "lat": 2.0,   "lon": 20.0,   "iso2": ""},
]

_MARKET_PULSE_CACHE_TTL_SECONDS = 300  # 5 min — plenty fresh for a daily % move, keeps Yahoo calls modest
_market_pulse_cache: dict = {"at": 0.0, "data": None}
_market_pulse_lock = threading.Lock()


def _fetch_market_pulse_live() -> dict:
    """Batch-downloads 2 daily bars for every tracked index and computes each one's
    %% move from the prior close. Cached for _MARKET_PULSE_CACHE_TTL_SECONDS so a page
    full of visitors collapses to one Yahoo round trip every 5 minutes rather than one
    per visitor. Any symbol Yahoo doesn't return data for is silently dropped — better
    to show 29 markets than fail the whole map over one bad ticker."""
    now = time.monotonic()
    with _market_pulse_lock:
        cached = _market_pulse_cache["data"]
        if cached is not None and now - _market_pulse_cache["at"] < _MARKET_PULSE_CACHE_TTL_SECONDS:
            return cached

    symbols = [m["symbol"] for m in MARKET_PULSE_INDICES]
    try:
        # 1mo rather than 5d — several thinner-volume markets (Saudi Arabia's TASI,
        # the Philippines' PSEi) don't trade every single day, so a 5-day window
        # sometimes contained fewer than the 2 closes needed to compute a % move
        # and that market would silently vanish from the map for no real reason.
        raw = yf.download(tickers=symbols, period="1mo", interval="1d", group_by="ticker",
                           threads=True, auto_adjust=False, progress=False)
    except Exception:
        raw = None

    markets = []
    for m in MARKET_PULSE_INDICES:
        try:
            df = raw[m["symbol"]] if len(symbols) > 1 else raw
            closes = df["Close"].dropna()
            if len(closes) < 2:
                continue
            # A handful of the thinner markets (confirmed live: Greece's Athens
            # General on Yahoo) have multi-week gaps in their daily bars — the
            # "previous close" ends up weeks stale, so the %% move computed from
            # it isn't a daily move at all, just a misleadingly large one. Skip
            # any pair of closes more than a long holiday weekend apart (a real
            # market closure never runs past ~4-5 calendar days) rather than
            # publish a number that looks like today's move but isn't.
            gap_days = (closes.index[-1] - closes.index[-2]).days
            if gap_days > 5:
                continue
            last, prev = float(closes.iloc[-1]), float(closes.iloc[-2])
            if not prev:
                continue
            pct = (last - prev) / prev * 100
            markets.append({**{k: m[k] for k in ("country", "index", "symbol", "lat", "lon", "iso2")},
                             "price": round(last, 2), "pct_change": round(pct, 2)})
        except Exception:
            continue  # symbol missing from the batch, no rows, or a bad frame — skip it

    up = sum(1 for x in markets if x["pct_change"] > 0)
    down = sum(1 for x in markets if x["pct_change"] < 0)
    unchanged = len(markets) - up - down
    avg_move = round(sum(x["pct_change"] for x in markets) / len(markets), 2) if markets else 0.0

    result = {
        "updated_at": _dt.datetime.utcnow().isoformat() + "Z",
        "markets": markets,
        "summary": {"tracked": len(markets), "up": up, "down": down, "unchanged": unchanged, "avg_move": avg_move},
    }
    with _market_pulse_lock:
        # A transient Yahoo hiccup (rate limit, timeout, empty batch) would otherwise
        # get cached as an empty "markets": [] result and served — grey map, "0
        # markets tracked" — for the full 5-minute TTL. If this attempt came back
        # empty and we have a previous good result, keep serving that stale-but-good
        # data instead. "at" still advances either way, so a sustained outage still
        # only retries once per TTL window rather than hammering Yahoo every request.
        _market_pulse_cache["at"] = now
        if markets or _market_pulse_cache["data"] is None:
            _market_pulse_cache["data"] = result
        else:
            result = _market_pulse_cache["data"]
    return result


# Powers the "gold-silver-ratio" data-viz page: a live gold:silver price ratio
# animated as a tilting balance scale, plus a real historical-average reference
# line (not the ~50:1 folklore number often quoted, but an actual mean computed
# from real daily history). Reuses _fetch_ohlcv's existing XAUUSD=X/XAGUSD=X
# path — OANDA real spot when configured, GC=F/SI=F futures proxy otherwise —
# the same real-spot pipeline every other metals number on this site already
# goes through, rather than standing up a separate feed.
_GOLD_SILVER_LIVE_CACHE_TTL_SECONDS = 12 * 60 * 60  # 12h, per spec — doesn't need to be any fresher
# A visitor can click the scale to force a fresh read (see api_gold_silver_ratio_live's
# ?refresh=1) rather than wait out the full 12h window. That still goes through this
# same cache/TTL check, just with a much shorter threshold — genuinely fresher data on
# a click, while capping how often the upstream feed can actually be hit to once per
# 30s no matter how many visitors click at once (protects Yahoo from a click-burst the
# same way the 12h TTL protects it from a traffic burst).
_GOLD_SILVER_FORCE_REFRESH_TTL_SECONDS = 30
_gold_silver_live_cache: dict = {"at": 0.0, "data": None}
_gold_silver_live_lock = threading.Lock()

# The historical-average leg pulls the *entire* available daily history (decades)
# to compute its mean, so it's cached far longer than the live leg — recomputing
# it every 12h alongside the live ratio would be needless load for a number that
# only drifts by fractions of a point per year.
_GOLD_SILVER_HIST_CACHE_TTL_SECONDS = 24 * 60 * 60  # 24h
_gold_silver_hist_cache: dict = {"at": 0.0, "data": None}
_gold_silver_hist_lock = threading.Lock()


def _fetch_gold_silver_historical_average() -> dict | None:
    """Mean gold:silver ratio across the full daily history both GC=F and SI=F carry
    on Yahoo — a real computed average rather than the commonly-quoted-but-unsourced
    "~50:1" figure. Calls yfinance directly with period="30y" rather than going
    through _fetch_ohlcv/VALID_PERIODS: confirmed live, period="max" and any
    explicit start_date both come back with zero rows for these continuous-futures
    symbols (a Yahoo quirk specific to that pair of request shapes), but an
    oversized relative period like "30y" is accepted and simply clamps to
    whatever's actually available — both symbols' real history turns out to start
    2000-08-30 (confirmed by "30y" and "40y" returning the identical row count),
    so this is genuinely "as far back as the data goes", not an arbitrary window.
    Futures-continuous closes (not the OANDA-spot leg the live number uses) since
    that's the only leg with multi-decade daily history here; the basis/contango
    premium they carry over true spot is a few percent at most and washes out
    across a 26-year average, unlike the live number where it actually matters."""
    now = time.monotonic()
    with _gold_silver_hist_lock:
        cached = _gold_silver_hist_cache["data"]
        if cached is not None and now - _gold_silver_hist_cache["at"] < _GOLD_SILVER_HIST_CACHE_TTL_SECONDS:
            return cached
    try:
        gold = _yf_history_with_retry(yf.Ticker("GC=F"), period="30y", interval="1d")["Close"].dropna()
        silver = _yf_history_with_retry(yf.Ticker("SI=F"), period="30y", interval="1d")["Close"].dropna()
        joined = pd.concat([gold, silver], axis=1, join="inner", keys=["gold", "silver"]).dropna()
        if joined.empty:
            return None
        ratio_series = joined["gold"] / joined["silver"]
        result = {
            "average": round(float(ratio_series.mean()), 1),
            "since": joined.index[0].date().isoformat(),
        }
    except Exception:
        return None
    with _gold_silver_hist_lock:
        _gold_silver_hist_cache["at"] = now
        _gold_silver_hist_cache["data"] = result
    return result


def _fetch_gold_silver_ratio_live(force: bool = False) -> dict:
    """Live gold:silver ratio from real spot XAU/USD and XAG/USD (see _fetch_ohlcv's
    metal-currency synthesis — OANDA real spot when configured, GC=F/SI=F futures
    proxy otherwise), plus the real historical average above. Cached for
    _GOLD_SILVER_LIVE_CACHE_TTL_SECONDS (12h) so a page full of visitors collapses
    to one round trip per window rather than one per visitor — same pattern as
    _fetch_market_pulse_live. A transient fetch failure falls back to the last
    good cached result (if any) rather than serving nulls to every visitor for
    the next 12h.

    force=True (from a click-to-refresh) checks against the much shorter
    _GOLD_SILVER_FORCE_REFRESH_TTL_SECONDS instead — still a real cache check, so a
    burst of clicks from one or many visitors within that window all just get the
    same still-fresh result rather than each triggering their own upstream fetch."""
    ttl = _GOLD_SILVER_FORCE_REFRESH_TTL_SECONDS if force else _GOLD_SILVER_LIVE_CACHE_TTL_SECONDS
    now = time.monotonic()
    with _gold_silver_live_lock:
        cached = _gold_silver_live_cache["data"]
        if cached is not None and now - _gold_silver_live_cache["at"] < ttl:
            return cached

    result = None
    try:
        xau = _fetch_ohlcv("XAUUSD=X", period="5d", interval="1d")["Close"].dropna()
        xag = _fetch_ohlcv("XAGUSD=X", period="5d", interval="1d")["Close"].dropna()
        if len(xau) and len(xag):
            xau_price, xag_price = float(xau.iloc[-1]), float(xag.iloc[-1])
            hist = _fetch_gold_silver_historical_average()
            result = {
                "updated_at": _dt.datetime.utcnow().isoformat() + "Z",
                "xau_usd": round(xau_price, 2),
                "xag_usd": round(xag_price, 2),
                "ratio": round(xau_price / xag_price, 1),
                "historical_average": hist["average"] if hist else None,
                "historical_since": hist["since"] if hist else None,
            }
    except Exception:
        result = None

    with _gold_silver_live_lock:
        _gold_silver_live_cache["at"] = now
        if result is not None:
            _gold_silver_live_cache["data"] = result
        else:
            result = _gold_silver_live_cache["data"]
    return result or {"updated_at": None, "xau_usd": None, "xag_usd": None, "ratio": None,
                       "historical_average": None, "historical_since": None}


# ── Data Visualisation content store ─────────────────────────────────────────
# Same Postgres/JSON dual-path pattern as alpha_content above, trimmed down to
# the 5 fields Tom's upload tool needs: image, description, positive analysis,
# an accuracy/usefulness warning, and an optional further-reading link.

_DATAVIZ_FIELDS = ["author", "page", "status", "title", "description", "positive_analysis", "warning", "link", "image_filename"]


def _dataviz_row_to_dict(row: dict) -> dict:
    d = {k: row.get(k) for k in ["id", *_DATAVIZ_FIELDS]}
    for key in ("created_at", "updated_at", "published_at"):
        val = row.get(key)
        d[key] = val.isoformat() if val is not None else None
    return d


def _load_dataviz_json() -> dict:
    if not os.path.exists(DATAVIZ_CONTENT_FILE):
        return {"items": [], "next_id": 1}
    with open(DATAVIZ_CONTENT_FILE, "r") as f:
        return json.load(f)


def _save_dataviz_json(data: dict) -> None:
    with open(DATAVIZ_CONTENT_FILE, "w") as f:
        json.dump(data, f, indent=2)


def dataviz_content_list(status: str | None = None, page: str | None = None) -> list:
    if DATABASE_URL:
        query = f"SELECT id, {', '.join(_DATAVIZ_FIELDS)}, created_at, updated_at, published_at FROM dataviz_content WHERE 1=1"
        params = []
        if status:
            query += " AND status = %s"
            params.append(status)
        if page:
            query += " AND page = %s"
            params.append(page)
        query += " ORDER BY COALESCE(published_at, created_at) DESC"
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(query, params)
            return [_dataviz_row_to_dict(row) for row in cur.fetchall()]
    data = _load_dataviz_json()
    items = data["items"]
    if status:
        items = [i for i in items if i.get("status") == status]
    if page:
        items = [i for i in items if i.get("page") == page]
    items = sorted(items, key=lambda i: i.get("published_at") or i.get("created_at") or "", reverse=True)
    return [{k: v for k, v in i.items() if k != "image_file"} for i in items]


def dataviz_content_get(item_id: int) -> dict | None:
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(
                f"SELECT id, {', '.join(_DATAVIZ_FIELDS)}, created_at, updated_at, published_at "
                f"FROM dataviz_content WHERE id = %s", (item_id,)
            )
            row = cur.fetchone()
            return _dataviz_row_to_dict(row) if row else None
    data = _load_dataviz_json()
    for item in data["items"]:
        if item.get("id") == item_id:
            return {k: v for k, v in item.items() if k != "image_file"}
    return None


def dataviz_content_get_image(item_id: int):
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute("SELECT image_filename, image_file FROM dataviz_content WHERE id = %s", (item_id,))
            row = cur.fetchone()
            if not row or row["image_file"] is None:
                return None, None
            return row["image_filename"], bytes(row["image_file"])
    data = _load_dataviz_json()
    for item in data["items"]:
        if item.get("id") == item_id:
            b64 = item.get("image_file")
            if not b64:
                return None, None
            return item.get("image_filename"), base64.b64decode(b64)
    return None, None


def dataviz_content_set_image(item_id: int, filename: str, file_bytes: bytes) -> dict | None:
    now = _dt.datetime.utcnow()
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(
                f"UPDATE dataviz_content SET image_file = %s, image_filename = %s, updated_at = %s WHERE id = %s "
                f"RETURNING id, {', '.join(_DATAVIZ_FIELDS)}, created_at, updated_at, published_at",
                (psycopg2.Binary(file_bytes), filename, now, item_id),
            )
            row = cur.fetchone()
            return _dataviz_row_to_dict(row) if row else None
    data = _load_dataviz_json()
    for item in data["items"]:
        if item.get("id") == item_id:
            item["image_file"] = base64.b64encode(file_bytes).decode("ascii")
            item["image_filename"] = filename
            item["updated_at"] = now.isoformat()
            _save_dataviz_json(data)
            return {k: v for k, v in item.items() if k != "image_file"}
    return None


def dataviz_content_create(fields: dict) -> dict:
    now = _dt.datetime.utcnow()
    if DATABASE_URL:
        cols = [*_DATAVIZ_FIELDS, "created_at", "updated_at"]
        vals = [fields.get(k) for k in _DATAVIZ_FIELDS] + [now, now]
        placeholders = ", ".join(["%s"] * len(vals))
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(
                f"INSERT INTO dataviz_content ({', '.join(cols)}) VALUES ({placeholders}) "
                f"RETURNING id, {', '.join(_DATAVIZ_FIELDS)}, created_at, updated_at, published_at",
                vals,
            )
            row = cur.fetchone()
            return _dataviz_row_to_dict(row)
    data = _load_dataviz_json()
    new_id = data.get("next_id", 1)
    item = {k: fields.get(k) for k in _DATAVIZ_FIELDS}
    item["id"] = new_id
    item["created_at"] = now.isoformat()
    item["updated_at"] = now.isoformat()
    item["published_at"] = None
    data["items"].append(item)
    data["next_id"] = new_id + 1
    _save_dataviz_json(data)
    return {k: v for k, v in item.items() if k != "image_file"}


def dataviz_content_update(item_id: int, updates: dict) -> dict | None:
    now = _dt.datetime.utcnow()
    allowed = set(_DATAVIZ_FIELDS) | {"published_at"}
    updates = {k: v for k, v in updates.items() if k in allowed}
    if DATABASE_URL:
        if not updates:
            return dataviz_content_get(item_id)
        set_clauses = [f"{k} = %s" for k in updates] + ["updated_at = %s"]
        vals = list(updates.values()) + [now, item_id]
        with _db_conn() as conn, conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(
                f"UPDATE dataviz_content SET {', '.join(set_clauses)} WHERE id = %s "
                f"RETURNING id, {', '.join(_DATAVIZ_FIELDS)}, created_at, updated_at, published_at",
                vals,
            )
            row = cur.fetchone()
            return _dataviz_row_to_dict(row) if row else None
    data = _load_dataviz_json()
    json_updates = dict(updates)
    if "published_at" in json_updates:
        val = json_updates["published_at"]
        json_updates["published_at"] = val.isoformat() if hasattr(val, "isoformat") else val
    for item in data["items"]:
        if item.get("id") == item_id:
            item.update(json_updates)
            item["updated_at"] = now.isoformat()
            _save_dataviz_json(data)
            return {k: v for k, v in item.items() if k != "image_file"}
    return None


def dataviz_content_delete(item_id: int) -> bool:
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor() as cur:
            cur.execute("DELETE FROM dataviz_content WHERE id = %s", (item_id,))
            return cur.rowcount > 0
    data = _load_dataviz_json()
    before = len(data["items"])
    data["items"] = [i for i in data["items"] if i.get("id") != item_id]
    _save_dataviz_json(data)
    return len(data["items"]) < before


# ── Alpha content: extraction & normalization ────────────────────────────────

def extract_text_from_upload(file_storage) -> str:
    """Extracts plain text from an uploaded .docx/.pdf/.xlsx/.xls FileStorage.
    Raises ValueError with a user-facing message on failure."""
    filename = file_storage.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_DOC_EXTENSIONS:
        raise ValueError(f"Unsupported file type .{ext or '?'} — please upload a Word, PDF or Excel file")
    try:
        if ext == "docx":
            import docx
            from docx.table import Table as DocxTable
            from docx.text.paragraph import Paragraph as DocxParagraph
            document = docx.Document(file_storage)
            # Walk paragraphs and tables in actual document order (python-docx's
            # .paragraphs/.tables only give each kind separately, losing where a
            # table sits relative to the surrounding text) and render each table
            # as a real markdown table (header + `---` separator row) so it's
            # unambiguous downstream — see auto_section_tables(), which wraps a
            # block matching this exact shape into its own Table section.
            parts = []
            for child in document.element.body.iterchildren():
                if child.tag.endswith("}p"):
                    p = DocxParagraph(child, document)
                    if p.text.strip():
                        parts.append(p.text.strip())
                elif child.tag.endswith("}tbl"):
                    table = DocxTable(child, document)
                    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                    rows = [r for r in rows if any(c for c in r)]
                    if not rows:
                        continue
                    md = ["| " + " | ".join(rows[0]) + " |", "|" + "|".join(" --- " for _ in rows[0]) + "|"]
                    for r in rows[1:]:
                        md.append("| " + " | ".join(r) + " |")
                    parts.append("\n".join(md))
            text = "\n\n".join(parts)
        elif ext == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(file_storage)
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        else:  # xlsx / xls
            import openpyxl
            wb = openpyxl.load_workbook(file_storage, data_only=True)
            parts = []
            for sheet in wb.worksheets:
                parts.append(f"[Sheet: {sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts)
    except ValueError:
        raise
    except Exception:
        raise ValueError("Couldn't read that file — it may be corrupted, password-protected, or an unsupported format")
    text = text.strip()
    if not text:
        raise ValueError("No readable text found in that file")
    return text[:MAX_UPLOAD_TEXT_CHARS]


def extract_text_from_url(url: str) -> str:
    """Fetches a URL and strips it down to readable text. Best-effort — many
    sites (paywalls, JS-rendered pages, bot protection) won't work well; the
    'paste text' input mode is the reliable fallback for those."""
    if not url.startswith(("http://", "https://")):
        raise ValueError("Please enter a valid http(s) link")
    try:
        resp = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0 (GCG-AlphaBot)"})
        resp.raise_for_status()
    except Exception:
        raise ValueError("Couldn't fetch that link — try pasting the text directly instead")
    try:
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "header", "footer", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        text = "\n".join(lines)
    except Exception:
        raise ValueError("Couldn't read the content at that link")
    if not text:
        raise ValueError("No readable text found at that link")
    return text[:MAX_UPLOAD_TEXT_CHARS]


def _strip_markdown_links(text: str) -> str:
    """`[label](url)` -> `label`. Title/subtitle are always rendered as plain
    text (see alpha-post.html), so any markdown-link syntax that survives
    into them — e.g. because the source sentence had already been through
    _auto_link_market_xi_assets, or the author pasted already-linked text —
    would otherwise show up literally as brackets instead of rendering or
    being dropped. Images (`![alt](url)`) are left alone; that alt text
    isn't meaningful as standalone title/subtitle text."""
    if not text:
        return text
    return re.sub(r"(?<!!)\[([^\]]+)\]\([^)]*\)", r"\1", text)


def normalize_content(raw_text: str, author: str, topics: list) -> dict:
    """Turns raw extracted text into {title, subtitle, topic, snippet, body} for a draft post.

    STUB — no ANTHROPIC_API_KEY is wired in yet. This does a naive extractive
    pass (first sentence as title, second sentence as a subtitle suggestion,
    next couple of sentences as snippet, raw text reflowed into paragraphs as
    the body, first nominated topic as the default) so the rest of the
    pipeline — upload, draft review, edit, publish, public rendering — is
    fully testable right now.

    To make this real: call the Claude API (see the claude-api skill for the
    current model id and Messages API usage) with a prompt that gives it
    `author`, the exact `topics` list (it must pick one of these three, not
    invent a new one), and `raw_text`, instructing it to return JSON
    {title, subtitle, topic, snippet, body} — allowed to tighten/condense
    wording (per the earlier scoping decision) but not invent facts absent
    from raw_text. Keep the same return shape so no caller needs to change.
    """
    text = " ".join(raw_text.split())
    sentences = re.split(r"(?<=[.!?])\s+", text)
    title = _strip_markdown_links((sentences[0] if sentences else text))[:100].strip() or "Untitled note"
    subtitle_source = sentences[1] if len(sentences) > 1 else (sentences[0] if sentences else text)
    subtitle_source = _strip_markdown_links(subtitle_source)
    subtitle = subtitle_source.strip()[:140]
    if len(subtitle_source) > 140:
        subtitle += "…"
    snippet_source = " ".join(sentences[1:3]) if len(sentences) > 1 else text
    snippet_words = snippet_source.split()
    snippet = " ".join(snippet_words[:40])
    if len(snippet_words) > 40:
        snippet += "…"
    paragraphs = [p.strip() for p in raw_text.split("\n") if p.strip()]
    body_parts = []
    for i, p in enumerate(paragraphs):
        if i > 0:
            # Keep consecutive markdown-table-row lines tight (single \n) so a
            # table survives this reflow intact — the normal blank-line-per-
            # paragraph spacing would otherwise pull each row apart, breaking
            # the table shape auto_section_tables() looks for below.
            prev_is_row = _looks_like_table_row(paragraphs[i - 1])
            this_is_row = _looks_like_table_row(p)
            body_parts.append("\n" if (prev_is_row and this_is_row) else "\n\n")
        body_parts.append(p)
    body = "".join(body_parts) if body_parts else text
    topic = topics[0] if topics else None
    return {"title": title, "subtitle": subtitle, "topic": topic, "snippet": snippet, "body": body}


def _looks_like_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def _is_table_separator_row(line: str) -> bool:
    """`| --- | :--- | ---: |` — a markdown table's header/body divider row.
    Checked cell-by-cell rather than with one big character class, since a
    naive `^\\|[\\s:-]+\\|$` also has to reject the '|' *between* cells."""
    s = line.strip()
    if not (s.startswith("|") and s.endswith("|")):
        return False
    cells = [c.strip() for c in s.strip("|").split("|")]
    return len(cells) >= 1 and all(re.match(r"^:?-{2,}:?$", c) for c in cells)


_LIST_LINE_RE = re.compile(r"^(-\s+|\d+\.\s+)")


def _join_body_lines(lines) -> str:
    """Joins section body lines with a blank line between paragraphs, same
    as before — except consecutive "- "/"1. " list lines are joined tight
    (single newline). The Studio's own line-based renderer closes a list on
    any blank line (see renderPreviewBlocks() in alpha-studio.html), so a
    blank line between what should be one 1-2-3 list would render it as
    three separate one-item lists instead."""
    out = []
    for i, line in enumerate(lines):
        if i > 0:
            out.append("\n" if (_LIST_LINE_RE.match(lines[i - 1]) and _LIST_LINE_RE.match(line)) else "\n\n")
        out.append(line)
    return "".join(out)


# Matches the SECTION_JOIN / type-marker microsyntax the studio's section
# editor uses (static/alpha-studio.html — keep in sync with SECTION_JOIN and
# the type marker patterns there).
_SECTION_JOIN = "<!--section-->"


def auto_section_tables(body: str) -> str:
    """Wraps any markdown-table-shaped block (a `| ... |` row, a `|---|---|`
    separator row, then more `| ... |` rows) in its own Table section, so a
    table detected in an uploaded document lands directly in a Table section
    for editing instead of as garbled pipe-delimited text in a Normal one.
    No-op if the body has no such block.
    """
    lines = body.split("\n")
    out_lines = []
    i = 0
    found_any = False
    while i < len(lines):
        line = lines[i]
        if _looks_like_table_row(line) and i + 1 < len(lines) and _is_table_separator_row(lines[i + 1]):
            # Found a table: header row + separator row, then consume further row lines.
            table_lines = [line, lines[i + 1]]
            j = i + 2
            while j < len(lines) and _looks_like_table_row(lines[j]):
                table_lines.append(lines[j])
                j += 1
            while out_lines and not out_lines[-1].strip():
                out_lines.pop()
            if out_lines:
                out_lines.append("")
                out_lines.append(_SECTION_JOIN)
                out_lines.append("")
            out_lines.append("<!--type:table-->")
            out_lines.extend(table_lines)
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines):
                out_lines.append("")
                out_lines.append(_SECTION_JOIN)
                out_lines.append("")
            found_any = True
            i = j
            continue
        out_lines.append(line)
        i += 1
    return "\n".join(out_lines) if found_any else body


# ── Automatic Word-doc formatting ─────────────────────────────────────────
# Any uploaded .docx is parsed structurally — no special template or
# convention is required from the author. Any built-in Heading style
# ("Heading 1".."Heading 9") starts a new Normal section named after that
# heading; a bracket tag at the start of a heading line — "[TIP] Careful
# with leverage", "[QUOTE] Warren Buffett", "[TABLE] Key Metrics" — still
# switches that section's type for authors who know the shorthand, but it's
# optional, not required. Bold/italic/underline formatting on the words
# themselves is picked up automatically too (see _run_markup below), and any
# inline picture found in the body text is auto-promoted into its own Image
# section (alternating left/right) with the paragraph right after it as the
# caption — no "[IMAGE]" tag needed. Tables are auto-detected from a real
# Word table regardless of any heading — see auto_section_tables above,
# which this mirrors. Keep the type markers in sync with stripTypeMarker()
# in static/alpha-studio.html and static/alpha-post.html.
#
# Two tags don't create a section at all — they fill the draft's own
# Subtitle/Snippet fields instead: "[SUBTITLE] ..." and "[SNIPPET] ...". Put
# the text right on the heading line, or leave the heading bare and write it
# as the paragraph(s) underneath — either works (see flush()'s "meta:" case).
_TYPE_MARKERS = {"tip": "<!--type:tip-->", "quote": "<!--type:quote-->", "table": "<!--type:table-->"}
_HEADING_TAG_RE = re.compile(r"^\[(TIP|QUOTE|TABLE|IMAGE\s+LEFT|IMAGE\s+RIGHT|IMAGE|SUBTITLE|SNIPPET)\]\s*", re.IGNORECASE)
_HEADING_STYLE_RE = re.compile(r"^Heading\s*\d", re.IGNORECASE)
_BLIP_TAG = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
_R_EMBED_ATTR = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
_CODE_FENCE_LINE_RE = re.compile(r"^\s*`{3,}\s*[A-Za-z]*\s*$")


def _strip_code_fence(text: str) -> str:
    """Strips a leading and/or trailing bare ``` fence line, if the whole
    paste was wrapped in one — an easy mistake when copying an AI's reply
    by hand (selecting the text yourself) rather than tapping its dedicated
    "copy code" button, which is the one that omits the fence. Left in
    place, a stray ``` line becomes part of the naive title-guesser's first
    "sentence" (there's no punctuation to stop it) and shows up as its own
    bogus empty section at the top of the post."""
    lines = text.split("\n")
    if lines and _CODE_FENCE_LINE_RE.match(lines[0]):
        lines = lines[1:]
    if lines and _CODE_FENCE_LINE_RE.match(lines[-1]):
        lines = lines[:-1]
    return "\n".join(lines).strip()


def _paragraph_image(paragraph, document):
    """Returns (bytes, extension) for the first inline picture in this
    paragraph's runs, or (None, None) if it has none."""
    for run in paragraph.runs:
        for blip in run._element.iter(_BLIP_TAG):
            rid = blip.get(_R_EMBED_ATTR)
            if rid and rid in document.part.related_parts:
                part = document.part.related_parts[rid]
                ext = (part.content_type.split("/")[-1] or "png").lower()
                if ext == "jpeg":
                    ext = "jpg"
                if ext not in ALLOWED_IMAGE_EXTENSIONS:
                    ext = "png"
                return part.blob, ext
    return None, None


def _run_markup(paragraph) -> str:
    """Renders a paragraph's runs to the Studio's plain-text markup
    (**bold**, _italic_, ++underline++), so formatting an author applied in
    Word survives the trip into the Studio instead of being flattened to
    plain text. Each run is wrapped independently rather than merging
    adjacent same-style runs — slightly more markers in the raw text than a
    human would type by hand, but it renders identically and is far simpler
    than reconstructing run boundaries."""
    parts = []
    for run in paragraph.runs:
        t = run.text
        if not t:
            continue
        if run.underline:
            t = "++" + t + "++"
        if run.italic:
            t = "_" + t + "_"
        if run.bold:
            t = "**" + t + "**"
        parts.append(t)
    return "".join(parts).strip()


def _build_list_classifier(document):
    """Returns classify(paragraph) -> 'bullet' | 'number' | None, so a
    Word-native list survives as the Studio's own "- "/"1. " line prefixes
    instead of silently losing its bullets/numbers (a bare paragraph.text
    never includes them — Word renders list markers from the numbering
    definition, they're not characters in the paragraph).

    Two authoring patterns need covering:
      1. A named "List Bullet"/"List Number" paragraph style (applied from
         the Styles gallery, or by anything scripting a .docx) — cheap
         string check, no XML needed.
      2. Direct formatting from clicking Word's own bullet/numbering toolbar
         buttons, which is the more common real-world case and does *not*
         set a named style — it attaches a bare numId to the paragraph, and
         the bullet-vs-number choice only lives in the numbering part's
         abstractNum definition. Resolved below; wrapped defensively since
         numbering.xml shape varies and this must never break the upload.
    """
    abstract_fmt = {}
    num_to_abstract = {}
    try:
        from docx.oxml.ns import qn
        root = document.part.numbering_part.element
        for absnum in root.findall(qn("w:abstractNum")):
            aid = absnum.get(qn("w:abstractNumId"))
            lvl0 = absnum.find(qn("w:lvl"))
            fmt_el = lvl0.find(qn("w:numFmt")) if lvl0 is not None else None
            abstract_fmt[aid] = fmt_el.get(qn("w:val")) if fmt_el is not None else None
        for numref in root.findall(qn("w:num")):
            nid = numref.get(qn("w:numId"))
            absref = numref.find(qn("w:abstractNumId"))
            if absref is not None:
                num_to_abstract[nid] = absref.get(qn("w:val"))
    except Exception:
        pass  # No numbering part, or an unexpected shape — direct-numPr paragraphs just won't classify below.

    def classify(paragraph):
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("List Number"):
            return "number"
        if style_name.startswith("List Bullet"):
            return "bullet"
        try:
            pPr = paragraph._p.pPr
            num_pr = pPr.numPr if pPr is not None else None
            if num_pr is None or num_pr.numId is None:
                return None
            fmt = abstract_fmt.get(num_to_abstract.get(str(num_pr.numId.val)))
            if fmt == "bullet":
                return "bullet"
            if fmt and fmt != "none":
                return "number"  # decimal, decimalZero, lowerLetter, upperRoman, ...
            return "bullet"  # Direct numPr with an unrecognised/missing format — still a list, default to the safer marker.
        except Exception:
            return None

    return classify


def extract_structured_docx(file_storage):
    """Parses any uploaded .docx into (sections, meta) — see the module
    comment above. Returns None only if the document is entirely empty (no
    text, no headings, no meta tags), in which case the caller falls back to
    the plain extract_text_from_upload() + normalize_content() path.

    sections: list of {type, heading, body_lines: [str, ...],
    rows: [[str,...],...] or None, image: (bytes, ext) or None,
    side: 'left'|'right'}. Images aren't uploaded here (no content row/id
    exists yet at extraction time — see _serialize_structured_sections,
    called after alpha_content_create).

    meta: {"subtitle": str or None, "snippet": str or None} — filled from any
    "[SUBTITLE]"/"[SNIPPET]" tagged block instead of becoming a section.
    """
    import docx
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph
    document = docx.Document(file_storage)

    def blank_section():
        return {"type": "normal", "heading": "", "body_lines": [], "rows": None, "image": None, "side": "left"}

    children = list(document.element.body.iterchildren())
    auto_image_side = ["left"]  # single-item list so the closure below can flip it

    def next_auto_side():
        auto_image_side[0] = "right" if auto_image_side[0] == "left" else "left"
        return auto_image_side[0]

    classify_list = _build_list_classifier(document)
    # Tracks a run of consecutive numbered-list paragraphs so they're
    # renumbered 1, 2, 3... — any paragraph that isn't itself a numbered item
    # (a heading, a bullet, plain prose) breaks the run and the next numbered
    # paragraph restarts at 1, same as a reader would expect two separate
    # lists to look.
    numbered_run = [0]

    sections = []
    meta = {"subtitle": None, "snippet": None}
    cur = blank_section()

    def flush():
        nonlocal cur
        numbered_run[0] = 0  # a section boundary always breaks a numbered-list run
        if cur["type"].startswith("meta:"):
            field = cur["type"].split(":", 1)[1]
            value = (cur["heading"] or "").strip() or "\n".join(cur["body_lines"]).strip()
            if value:
                meta[field] = value
        elif cur["heading"] or cur["body_lines"] or cur["rows"] or cur["image"]:
            sections.append(cur)
        cur = blank_section()

    for child in children:
        if child.tag.endswith("}p"):
            p = DocxParagraph(child, document)
            style_name = p.style.name if p.style else ""
            text = p.text.strip()
            if _HEADING_STYLE_RE.match(style_name):
                flush()
                m = _HEADING_TAG_RE.match(text)
                if m:
                    tag = re.sub(r"\s+", " ", m.group(1).upper())
                    remainder = text[m.end():].strip()
                    if tag == "TIP":
                        cur["type"] = "tip"
                    elif tag == "QUOTE":
                        cur["type"] = "quote"
                    elif tag == "TABLE":
                        cur["type"] = "table"
                    elif tag == "SUBTITLE":
                        cur["type"] = "meta:subtitle"
                    elif tag == "SNIPPET":
                        cur["type"] = "meta:snippet"
                    elif tag.startswith("IMAGE"):
                        cur["type"] = "image"
                        cur["side"] = "right" if "RIGHT" in tag else "left"
                    cur["heading"] = remainder
                else:
                    # A plain heading (no bracket tag) with a picture in the
                    # body — Word "Insert Picture" then a caption is normal
                    # authoring, it doesn't mean the author intended a manual
                    # "[IMAGE]" section — leave this as a Normal section and
                    # let the auto-promotion below carve the picture out.
                    cur["heading"] = text
                continue
            img_bytes, img_ext = _paragraph_image(p, document)
            if img_bytes:
                if cur["type"] == "image" and not cur["image"]:
                    # A manual "[IMAGE]"/"[IMAGE RIGHT]" heading is waiting
                    # for its picture — fill the slot it already opened.
                    cur["image"] = (img_bytes, img_ext)
                    continue
                # No guide, no tag — any picture found in ordinary body text
                # is auto-promoted into its own Image section. Whatever text
                # came before it keeps its own section; the picture starts a
                # fresh one, and the paragraph(s) after it become its caption
                # until the next heading or picture.
                #
                # Common case: "Heading, then picture, then caption" with no
                # body text of its own under that heading — that heading was
                # clearly introducing the picture, not starting a separate
                # empty section, so carry it onto the image section instead
                # of leaving it behind as a stray heading-only Normal one.
                carried_heading = ""
                if cur["type"] == "normal" and cur["heading"] and not cur["body_lines"] and not cur["rows"] and not cur["image"]:
                    carried_heading = cur["heading"]
                    cur = blank_section()
                else:
                    flush()
                cur["type"] = "image"
                cur["heading"] = carried_heading
                cur["side"] = next_auto_side()
                cur["image"] = (img_bytes, img_ext)
                continue
            markup = _run_markup(p)
            if markup:
                kind = classify_list(p)
                if kind == "number":
                    numbered_run[0] += 1
                    markup = f"{numbered_run[0]}. {markup}"
                else:
                    numbered_run[0] = 0
                    if kind == "bullet":
                        markup = f"- {markup}"
                cur["body_lines"].append(markup)
        elif child.tag.endswith("}tbl"):
            table = DocxTable(child, document)
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            rows = [r for r in rows if any(c for c in r)]
            if not rows:
                continue
            if cur["type"] == "table" and cur["rows"] is None:
                cur["rows"] = rows
            else:
                # A table with no preceding "[TABLE]" heading still gets its
                # own section automatically, same rule as auto_section_tables()
                # for a non-template upload. As with the image case above, a
                # heading directly over the table with no body text of its
                # own was introducing the table, not starting an empty
                # section — carry it across instead of leaving it stranded.
                carried_heading = ""
                if cur["type"] == "normal" and cur["heading"] and not cur["body_lines"] and not cur["rows"] and not cur["image"]:
                    carried_heading = cur["heading"]
                    cur = blank_section()
                else:
                    flush()
                cur["type"] = "table"
                cur["heading"] = carried_heading
                cur["rows"] = rows
                flush()
    flush()
    if not sections and not any(meta.values()):
        return None
    return sections, meta


def extract_structured_text(raw_text: str):
    """The plain-text equivalent of extract_structured_docx(), for the
    Paste Text and Link upload paths — there's no Word document to read
    Heading styles from there, so this recognises the same lightweight
    convention directly as literal characters instead: a line starting
    with "## " starts a new section, with an optional bracket tag —
    "[TIP]", "[QUOTE]", "[TABLE]", "[IMAGE]"/"[IMAGE RIGHT]", "[SUBTITLE]",
    "[SNIPPET]" — right after the "## " picking that section's type. This
    is exactly the convention an author's AI is instructed to write in
    (see the "Alpha post" custom-instruction prompt), so pasting its
    output gets real Tip/Quote colour and a genuine Table section instead
    of one flat, undifferentiated block of text.

    A bare markdown table with no "[TABLE]" heading above it still becomes
    its own table section automatically, same rule as auto_section_tables()
    uses for a non-templated upload.

    Returns (sections, meta) in the exact same shape extract_structured_docx()
    returns, so both feed the same _serialize_structured_sections(). Images
    are always None here — plain text can't carry a real picture — but an
    "[IMAGE]" section is still created with its heading and layout side, so
    the author only has to upload the actual picture into it afterward
    rather than build the section from scratch. Returns None if the text
    has no "## " heading anywhere and no bare table (nothing to structure —
    caller falls back to normalize_content() as before).
    """
    lines = raw_text.replace("\r\n", "\n").split("\n")
    if not any(re.match(r"^##\s+\S", ln.strip()) for ln in lines):
        # No deliberate use of the convention at all — leave this text to the
        # existing normalize_content() + auto_section_tables() path exactly
        # as before, rather than routing every ordinary paste through here.
        return None

    def blank_section():
        return {"type": "normal", "heading": "", "body_lines": [], "rows": None, "image": None, "side": "left"}

    sections = []
    meta = {"subtitle": None, "snippet": None}
    cur = blank_section()

    def flush():
        nonlocal cur
        if cur["type"].startswith("meta:"):
            field = cur["type"].split(":", 1)[1]
            value = (cur["heading"] or "").strip() or "\n".join(cur["body_lines"]).strip()
            if value:
                meta[field] = value
        elif cur["heading"] or cur["body_lines"] or cur["rows"]:
            sections.append(cur)
        cur = blank_section()

    i = 0
    while i < len(lines):
        line = lines[i]
        heading_m = re.match(r"^##\s+(.*)$", line.strip())
        if heading_m:
            flush()
            text = heading_m.group(1).strip()
            m = _HEADING_TAG_RE.match(text)
            if m:
                tag = re.sub(r"\s+", " ", m.group(1).upper())
                remainder = text[m.end():].strip()
                if tag == "TIP":
                    cur["type"] = "tip"
                elif tag == "QUOTE":
                    cur["type"] = "quote"
                elif tag == "TABLE":
                    cur["type"] = "table"
                elif tag == "SUBTITLE":
                    cur["type"] = "meta:subtitle"
                elif tag == "SNIPPET":
                    cur["type"] = "meta:snippet"
                elif tag.startswith("IMAGE"):
                    cur["type"] = "image"
                    cur["side"] = "right" if "RIGHT" in tag else "left"
                cur["heading"] = remainder
            else:
                # Same heading-carries-onto-the-picture-or-table idea as the
                # docx path — but here it's simpler: a plain heading with
                # nothing under it yet just stays as-is; the table branch
                # below carries it across the same way if a bare table
                # follows immediately.
                cur["heading"] = text
            i += 1
            continue
        if _looks_like_table_row(line) and i + 1 < len(lines) and _is_table_separator_row(lines[i + 1]):
            table_lines = [line, lines[i + 1]]
            j = i + 2
            while j < len(lines) and _looks_like_table_row(lines[j]):
                table_lines.append(lines[j])
                j += 1
            rows = [[c.strip() for c in ln.strip().strip("|").split("|")] for ln in [table_lines[0]] + table_lines[2:]]
            carried_heading = ""
            if cur["type"] == "table" and cur["rows"] is None:
                cur["rows"] = rows
            else:
                if cur["type"] == "normal" and cur["heading"] and not cur["body_lines"] and not cur["rows"]:
                    carried_heading = cur["heading"]
                    cur = blank_section()
                else:
                    flush()
                cur["type"] = "table"
                cur["heading"] = carried_heading
                cur["rows"] = rows
            i = j
            continue
        if line.strip():
            cur["body_lines"].append(line.rstrip())
        i += 1
    flush()
    if not sections and not any(meta.values()):
        return None
    return sections, meta


def _serialize_structured_sections(sections, upload_image) -> str:
    """Builds the SECTION_JOIN-delimited body string from
    extract_structured_docx() output. `upload_image(bytes, ext) -> url` is
    called once per section that has an extracted image."""
    parts = []
    for sec in sections:
        t = sec["type"]
        heading = (sec.get("heading") or "").strip()
        if t == "table":
            rows = sec.get("rows") or []
            if not rows:
                continue
            md = ["| " + " | ".join(rows[0]) + " |", "|" + "|".join(" --- " for _ in rows[0]) + "|"]
            for r in rows[1:]:
                md.append("| " + " | ".join(r) + " |")
            chunk = "<!--type:table-->"
            if heading:
                chunk += "\n## " + heading
            chunk += "\n" + "\n".join(md)
            parts.append(chunk)
            continue
        if t == "image":
            body = "\n".join(sec.get("body_lines") or [])
            img = sec.get("image")
            url = upload_image(*img) if img else None
            if not heading and not body and not url:
                continue
            chunk = "<!--type:image:" + sec.get("side", "left") + "-->"
            if heading:
                chunk += "\n## " + heading
            if url:
                chunk += "\n![" + heading.replace("[", "").replace("]", "") + "](" + url + ")"
            if body:
                chunk += ("\n\n" if (heading or url) else "\n") + body
            parts.append(chunk)
            continue
        # normal / tip / quote
        body = _join_body_lines(sec.get("body_lines") or [])
        chunk = _TYPE_MARKERS.get(t, "")
        if heading:
            chunk += ("\n" if chunk else "") + "## " + heading
        if body:
            chunk += ("\n\n" if chunk else "") + body
        if chunk.strip():
            parts.append(chunk)
    return ("\n\n" + _SECTION_JOIN + "\n\n").join(parts)


def _send_email(to_addr: str, subject: str, body: str) -> None:
    # Resend's HTTP API is a plain HTTPS POST (port 443), so it works from hosts
    # like Railway that block outbound SMTP (ports 25/465/587) at the network
    # level — raw SMTP there fails immediately with "Network is unreachable"
    # regardless of provider or port. Preferred whenever configured; SMTP below
    # is kept as a fallback for anyone not on a host that blocks it.
    resend_api_key = os.environ.get("RESEND_API_KEY", "")
    if resend_api_key:
        from_addr = os.environ.get("RESEND_FROM", "onboarding@resend.dev")
        try:
            resp = requests.post(
                "https://api.resend.com/emails",
                headers={"Authorization": f"Bearer {resend_api_key}"},
                json={"from": from_addr, "to": [to_addr], "subject": subject, "text": body},
                timeout=10,
            )
            if resp.status_code >= 300:
                print(f"\n  ✉️  [resend send failed] To: {to_addr}  Subject: {subject}  "
                      f"Status: {resp.status_code}  Body: {resp.text}\n")
        except Exception as e:
            print(f"\n  ✉️  [resend send failed] To: {to_addr}  Subject: {subject}  Error: {e}\n")
        return

    smtp_host = os.environ.get("SMTP_HOST", "")
    if not smtp_host:
        # No SMTP configured — log so it's still usable during setup/testing.
        print(f"\n  ✉️  [email not configured] To: {to_addr}  Subject: {subject}\n{body}\n")
        return
    smtp_port = int(os.environ.get("SMTP_PORT", "587"))
    smtp_user = os.environ.get("SMTP_USER", "")
    smtp_pass = os.environ.get("SMTP_PASS", "")
    from_addr = os.environ.get("SMTP_FROM", smtp_user)
    msg = MIMEText(body)
    msg["Subject"] = subject
    msg["From"] = from_addr
    msg["To"] = to_addr
    try:
        with smtplib.SMTP(smtp_host, smtp_port, timeout=10) as server:
            server.starttls()
            if smtp_user:
                server.login(smtp_user, smtp_pass)
            server.sendmail(from_addr, [to_addr], msg.as_string())
    except Exception as e:
        # Never let a broken SMTP config (bad creds, blocked port, timeout) hang or
        # fail the request — forgot-password always returns its generic message
        # regardless of whether the email actually went out, so this just logs it.
        print(f"\n  ✉️  [email send failed] To: {to_addr}  Subject: {subject}  Error: {e}\n")


def _save_preferences(username: str, preferences: dict) -> None:
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor() as cur:
            cur.execute(
                "UPDATE users SET preferences = %s WHERE username = %s",
                (json.dumps(preferences), username)
            )
        return
    users = _load_users()
    users[username]["preferences"] = preferences
    _save_users(users)


def _ensure_default_user() -> None:
    if DATABASE_URL:
        with _db_conn() as conn, conn.cursor() as cur:
            cur.execute("SELECT 1 FROM users LIMIT 1")
            if cur.fetchone():
                return
        password = secrets.token_urlsafe(12)
        pw_hash = bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")
        _save_users({"admin": {"password_hash": pw_hash, "preferences": {}, "tier": "power_user"}})
        print(f"\n  ✓ Created default user  →  username: admin  |  password: {password}  |  tier: power_user\n")
        return
    if os.path.exists(USERS_FILE):
        return
    password = secrets.token_urlsafe(12)
    pw_hash = bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")
    _save_users({"admin": {"password_hash": pw_hash, "preferences": {}, "tier": "power_user"}})
    print(f"\n  ✓ Created default user  →  username: admin  |  password: {password}  |  tier: power_user\n")


# ── OHLCV helper ─────────────────────────────────────────────────────────────

def _resample_ohlcv(df: pd.DataFrame, rule: str) -> pd.DataFrame:
    agg = {"Open": "first", "High": "max", "Low": "min", "Close": "last"}
    if "Volume" in df.columns:
        agg["Volume"] = "sum"
    resampled = df.resample(rule).agg(agg)
    return resampled.dropna(subset=["Open"])


# yfinance/Yahoo occasionally throws transient errors (rate limiting, connection resets,
# brief service hiccups) that have nothing to do with the symbol or params being wrong —
# retrying a couple of times with a short backoff clears most of them before a user ever
# sees a failure. Only exhausted retries bubble up, as a ValueError clearly labeled
# "temporarily unavailable" so callers/UI can tell that apart from "bad symbol/params".
_YF_RETRY_ATTEMPTS = 3
_YF_RETRY_BACKOFF_SECONDS = 0.75

# Short-TTL cache for yfinance's raw history call, keyed by (symbol, period, interval).
# The signal engine and price chart both poll _fetch_ohlcv repeatedly (on cadences down
# to a few seconds for intraday/scalp use — see pollTimer in signal_config.html), and
# every poll used to re-download the full history from Yahoo with zero caching, which
# is what capped how often either loop could safely run without risking rate-limiting.
# gthread workers (see Procfile) share this dict across all threads in a process, so a
# burst of near-simultaneous requests for the same symbol/interval collapses to one
# Yahoo call. Caching only the pre-stitch fetch is safe for "live" freshness: intraday
# staleness of a few seconds on the older, unchanging bars is invisible, and
# _stitch_live_tail still overwrites the trailing bar with a fresh OANDA/Alpaca tick on
# every call regardless of cache hit/miss.
_OHLCV_CACHE_TTL_SECONDS = 4
_ohlcv_cache: dict[tuple, tuple[float, pd.DataFrame]] = {}
_ohlcv_cache_lock = threading.Lock()


def _yf_history_with_retry(ticker: "yf.Ticker", **kwargs) -> pd.DataFrame:
    last_err: Exception | None = None
    for attempt in range(_YF_RETRY_ATTEMPTS):
        try:
            return ticker.history(auto_adjust=False, **kwargs)
        except Exception as e:
            last_err = e
            if attempt < _YF_RETRY_ATTEMPTS - 1:
                time.sleep(_YF_RETRY_BACKOFF_SECONDS * (attempt + 1))
    raise ValueError(f"Yahoo Finance data temporarily unavailable — try again in a moment ({last_err})")


def _fetch_yf_history_cached(ticker: "yf.Ticker", symbol: str, period: str, fetch_interval: str) -> pd.DataFrame:
    """Cached wrapper around the period-based branch of _yf_history_with_retry. Not
    used for the explicit start_date/end_date branch (backtester) — those are one-off
    historical range queries, not something repeatedly polled, so caching them would
    only add staleness risk for no benefit. Returns a copy so callers are free to treat
    the result as theirs alone, even though it's backed by a shared cache entry."""
    key = (symbol.upper(), period, fetch_interval)
    now = time.monotonic()
    with _ohlcv_cache_lock:
        cached = _ohlcv_cache.get(key)
    if cached is not None and now - cached[0] < _OHLCV_CACHE_TTL_SECONDS:
        return cached[1].copy()
    df = _yf_history_with_retry(ticker, period=period, interval=fetch_interval)
    with _ohlcv_cache_lock:
        _ohlcv_cache[key] = (now, df)
    return df.copy()


# Yahoo delisted the real XAUUSD=X/XAUGBP=X/XAGEUR=X-style spot-metal "currency"
# crosses — every one of them now 404s outright ("Quote not found"), confirmed live
# against yfinance for all of USD/GBP/EUR/JPY/CAD/AUD/CHF/INR. The metal futures
# (GC=F, SI=F) and ordinary USD/<currency> FX crosses are both still fine, so
# XAU*/XAG* symbols are synthesized from those two legs instead of fetched directly.
_METAL_USD_FUTURES = {"XAU": "GC=F", "XAG": "SI=F"}
_METAL_FX_CURRENCIES = {"USD", "GBP", "EUR", "JPY", "CAD", "AUD", "CHF", "INR"}


def _parse_metal_currency_symbol(symbol: str) -> tuple[str, str] | None:
    s = symbol.upper()
    if s.endswith("=X"):
        s = s[:-2]
    if len(s) == 6 and s[:3] in _METAL_USD_FUTURES and s[3:] in _METAL_FX_CURRENCIES:
        return s[:3], s[3:]
    return None


def _fetch_oanda_metal_history(symbol: str, period: str, interval: str) -> pd.DataFrame | None:
    """Real spot XAU/XAG history straight from OANDA (whenever it's configured and
    lists this exact pair) — used in preference to _fetch_synthetic_metal_ohlcv's
    GC=F/SI=F futures proxy below. Futures trade at a basis/contango premium to true
    spot (confirmed live: GC=F ran ~$50-60 above real XAU/USD spot on 2026-07-29),
    which showed up on the chart as a misleadingly high daily candle. Returns None
    on any non-coverage/failure so the caller falls straight back to the existing
    futures-synthesis path — same contract as marketdata_router's other lookups.
    Not used for the start_date/end_date branch (backtester) — OANDA's candles
    endpoint takes a "most recent N" count, not an explicit date range, so it isn't
    a fit for a one-off historical query the way it is for the "current" live-chart
    path.

    Cached the same way (and for the same reason) as _fetch_yf_history_cached: the
    signal engine and price chart both poll this on cadences down to a few seconds,
    and every poll would otherwise re-hit OANDA's REST API. Shares _ohlcv_cache with
    the yfinance leg — namespaced with an "oanda" tag so the two never collide on
    the same key even when they'd otherwise fetch the same display symbol."""
    key = ("oanda", symbol.upper(), period, interval)
    now = time.monotonic()
    with _ohlcv_cache_lock:
        cached = _ohlcv_cache.get(key)
    if cached is not None and now - cached[0] < _OHLCV_CACHE_TTL_SECONDS:
        return cached[1].copy()
    fetch_interval = _RESAMPLE_INTERVALS.get(interval, interval)
    df = marketdata_router.get_historical_candles(symbol, fetch_interval, period)
    if df is None or df.empty:
        return None
    if interval in _RESAMPLE_INTERVALS:
        df = _resample_ohlcv(df, _RESAMPLE_RULES[interval])
        if df.empty:
            return None
    with _ohlcv_cache_lock:
        _ohlcv_cache[key] = (now, df)
    return df.copy()


def _fetch_synthetic_metal_ohlcv(
    metal: str, currency: str, period: str, interval: str,
    start_date: str | None, end_date: str | None,
) -> pd.DataFrame:
    usd_df = _fetch_ohlcv(_METAL_USD_FUTURES[metal], period, interval, start_date, end_date)
    if currency == "USD":
        return usd_df
    fx_df = _fetch_ohlcv(f"USD{currency}=X", period, interval, start_date, end_date)
    # Align the FX rate onto every futures bar's timestamp — futures (CME Globex) and
    # FX trade on different session calendars/timezones, so exact timestamp matches
    # aren't guaranteed even though both are tz-aware DatetimeIndexes. ffill covers
    # gaps going forward; bfill covers any futures bars older than the FX history.
    fx_close = fx_df["Close"].reindex(usd_df.index, method="ffill").bfill()
    if fx_close.isna().any():
        raise ValueError(f"No FX rate available to convert {metal} into {currency}")
    # For daily-or-coarser intervals specifically, ffill silently goes stale on the
    # trailing (today's still-forming) bar: Yahoo labels GC=F's daily bars at US-Eastern
    # midnight and USD<ccy>=X's at Europe/London midnight, several hours apart — so right
    # now, FX's own latest bar is labeled "tomorrow" relative to the futures bar even
    # though both sessions are live at this instant, and ffill (correctly, by its own
    # rules) refuses to use a "future" row and falls back a full day. Confirmed live:
    # this drifted XAUGBP off the true rate by ~13pts/0.4%, permanently, until the next
    # calendar rollover masked it again. Overriding just the trailing bar with fx_df's
    # own raw latest close sidesteps the label mismatch entirely; it's a no-op for
    # intraday intervals, where both legs already share real, closely-timestamped bars.
    fx_close.iloc[-1] = fx_df["Close"].iloc[-1]
    converted = usd_df.copy()
    for col in ("Open", "High", "Low", "Close"):
        converted[col] = converted[col] * fx_close
    return converted


def _fetch_ohlcv(
    symbol: str,
    period: str = "3mo",
    interval: str = "1d",
    start_date: str | None = None,
    end_date: str | None = None,
) -> pd.DataFrame:
    if interval not in ALL_VALID_INTERVALS:
        raise ValueError(f"Invalid interval: {interval}")
    # Validated here (rather than only in the plain yfinance branch below) so the new
    # OANDA-first metal path can't silently accept a bogus period — _estimate_candle_count
    # would otherwise just default it to ~90 days instead of surfacing the same 400 a
    # bad period already gets everywhere else.
    if not start_date and period not in VALID_PERIODS:
        raise ValueError(f"Invalid period: {period}")
    metal_ccy = _parse_metal_currency_symbol(symbol)
    if metal_ccy:
        # Used to return directly here, which meant the metal-currency pair itself
        # (e.g. 'XAUGBP=X') never reached _stitch_live_tail below — only the two
        # internal legs (GC=F, USDGBP=X) did, via their own nested _fetch_ohlcv calls,
        # and neither of those is the symbol OANDA is actually asked to watch. Falling
        # through to the same tail logic as the plain path fixes that.
        df = _fetch_oanda_metal_history(symbol, period, interval) if not start_date else None
        if df is None:
            df = _fetch_synthetic_metal_ohlcv(*metal_ccy, period, interval, start_date, end_date)
    else:
        fetch_interval = _RESAMPLE_INTERVALS.get(interval, interval)
        ticker = yf.Ticker(symbol.upper())
        if start_date:
            import datetime as _dt
            try:
                _dt.date.fromisoformat(start_date)
                if end_date:
                    _dt.date.fromisoformat(end_date)
            except ValueError:
                raise ValueError("start_date / end_date must be YYYY-MM-DD")
            # auto_adjust=False: keep raw (split-adjusted only, not dividend-adjusted) closes.
            # yfinance's auto_adjust=True folds dividends into every historical Close, which
            # can badly skew RSI/MACD/etc. for anything with a meaningful dividend/distribution
            # history — and it's also what our own price chart plots, so indicators computed
            # here always line up with what's on screen.
            df = _yf_history_with_retry(ticker, start=start_date, end=end_date or None, interval=fetch_interval)
        else:
            df = _fetch_yf_history_cached(ticker, symbol, period, fetch_interval)
        if df.empty:
            raise ValueError(f"No data returned for symbol: {symbol} — check the ticker is correct")
        if interval in _RESAMPLE_INTERVALS:
            df = _resample_ohlcv(df, _RESAMPLE_RULES[interval])
            if df.empty:
                raise ValueError(f"No data returned for symbol: {symbol}")
    if not start_date:
        df = _stitch_live_tail(df, symbol, interval)
    return df


def _stitch_live_tail(df: pd.DataFrame, symbol: str, interval: str) -> pd.DataFrame:
    """Replaces the trailing rows of `df` with a fresher live tail from
    marketdata.router (OANDA/Alpaca), if one is available for this symbol — never
    raises, and returns `df` unchanged on any failure or if nothing is available
    (not configured, not a covered category, stream stale). Only called for
    period-based "current" queries, not explicit start_date/end_date historical
    ranges (e.g. the backtester), where stitching in "now" wouldn't make sense."""
    try:
        live_tail = marketdata_router.get_live_tail(symbol, interval, tz=df.index.tz)
        if live_tail is None or live_tail.empty:
            return df
        # OANDA/Alpaca timestamps are UTC; match df's own tz so the join point doesn't
        # show a visually-inconsistent offset between historical and live rows (purely
        # cosmetic — pandas compares tz-aware instants correctly either way).
        if df.index.tz is not None:
            live_tail = live_tail.tz_localize("UTC") if live_tail.index.tz is None else live_tail
            live_tail.index = live_tail.index.tz_convert(df.index.tz)
        historical = df[df.index < live_tail.index[0]]
        return pd.concat([historical, live_tail])
    except Exception:
        return df


# ── Static ───────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return send_from_directory("static", "index.html")


@app.route("/login")
def login_page():
    return send_from_directory("static", "login.html")

@app.route("/signal-config")
def signal_config():
    return send_from_directory("static", "signal_config.html")

@app.route("/portfolio-balancer")
def portfolio_balancer():
    return send_from_directory("static", "portfolio-balancer.html")

@app.route("/backtester")
def backtester():
    return send_from_directory("static", "strategy-lab.html")

@app.route("/backtester/forex")
def backtester_forex():
    return send_from_directory("static", "strategy-lab-forex.html")

@app.route("/stories")
def stories():
    return send_from_directory("static", "stories.html")

@app.route("/education-hub")
def education_hub():
    return send_from_directory("static", "EducationHub.html")

@app.route("/stories-podcasts")
def stories_podcasts():
    return send_from_directory("static", "StoriesPodcasts.html")

@app.route("/competitions-partnerships")
def competitions_partnerships():
    return send_from_directory("static", "CompetitionsPartnerships.html")

@app.route("/btc-swing-trade")
def btc_swing_trade():
    return send_from_directory("static", "btc-swing-trade.html")

@app.route("/roadmap")
def roadmap():
    return send_from_directory("static", "roadmap.html")

@app.route("/company")
def company():
    return send_from_directory("static", "company.html")

@app.route("/guides/signals")
def guide_signals():
    return send_from_directory("static", "guide-signals.html")

@app.route("/guides/backtester")
def guide_backtester():
    return send_from_directory("static", "guide-backtester.html")

@app.route("/guides/portfolio")
def guide_portfolio():
    return send_from_directory("static", "guide-portfolio.html")

# ── New nav routes ────────────────────────────────────────────────────────────

@app.route("/learn")
def learn(): return send_from_directory("static", "learn.html")

@app.route("/learn/beginner")
def learn_beginner(): return send_from_directory("static", "learn-beginner.html")

def _make_lesson_view(filename):
    # A closure per lesson so each route serves its own file — add_url_rule
    # needs a distinct function object per endpoint, not just a distinct name.
    def view():
        return send_from_directory("static", filename)
    return view

for _lesson in LESSON_PAGES:
    app.add_url_rule(
        f"/learn/{_lesson['level']}/{_lesson['slug']}",
        endpoint=f"lesson_{_lesson['level']}_{_lesson['slug']}".replace("-", "_"),
        view_func=_make_lesson_view(_lesson["file"]),
    )

@app.route("/learn/intermediate")
def learn_intermediate(): return send_from_directory("static", "learn-intermediate.html")

@app.route("/learn/pro")
def learn_pro(): return send_from_directory("static", "learn-pro.html")

@app.route("/learn/tools")
def learn_tools(): return send_from_directory("static", "learn-tools.html")

@app.route("/tools")
def tools(): return send_from_directory("static", "tools.html")

@app.route("/tools/signals")
def tools_signals(): return send_from_directory("static", "signal_config.html")

@app.route("/tools/portfolio")
def tools_portfolio(): return send_from_directory("static", "portfolio-balancer.html")

@app.route("/tools/calculator")
def tools_calculator(): return send_from_directory("static", "calculator.html")

@app.route("/tools/data-visualisation")
def tools_data_visualisation(): return send_from_directory("static", "data-visualisation.html")

@app.route("/tools/data-visualisation/<slug>")
def tools_data_visualisation_page(slug):
    if not dataviz_page_get(slug):
        return send_from_directory("static", "data-visualisation.html"), 404
    return send_from_directory("static", "data-visualisation-page.html")

@app.route("/arena")
def arena(): return send_from_directory("static", "arena.html")

@app.route("/arena/market-xi")
def arena_market_xi(): return send_from_directory("static", "arena-market-xi.html")

@app.route("/arena/competitions")
def arena_competitions(): return send_from_directory("static", "arena-competitions.html")

@app.route("/arena/predictions")
def arena_predictions(): return send_from_directory("static", "arena-predictions.html")

@app.route("/alpha")
def alpha(): return send_from_directory("static", "alpha.html")

@app.route("/alpha/connor")
def alpha_connor(): return send_from_directory("static", "alpha-connor.html")

@app.route("/alpha/dave")
def alpha_dave(): return send_from_directory("static", "alpha-dave.html")

@app.route("/alpha/gary")
def alpha_gary(): return send_from_directory("static", "alpha-gary.html")

@app.route("/alpha/tom")
def alpha_tom(): return send_from_directory("static", "alpha-tom.html")

@app.route("/alpha/podcast")
def alpha_podcast(): return send_from_directory("static", "alpha-podcast.html")

@app.route("/partners")
def partners(): return send_from_directory("static", "partners.html")

@app.route("/social-post-studio")
def social_post_studio(): return send_from_directory("static", "social-post-studio.html")


# ── Portfolio Balancer — live price feed ──────────────────────────────────────

_PB_TICKERS = {
    # Safe (bonds & cash proxies)
    's_tbills': 'SHY',    # iShares 1-3yr Treasury Bond ETF
    's_gilts':  'IGLT.L', # iShares UK Gilts UCITS ETF
    's_euro':   'IBTE.L', # iShares € Govt Bond 1-3yr UCITS ETF
    's_corp':   'LQD',    # iShares iBoxx IG Corp Bond ETF
    's_cash':   'ERNS.L', # iShares GBP Ultrashort Bond ETF (cash proxy)
    # Hard Assets (property / land / infrastructure proxies)
    'h_ukres':  'IUKP.L', # iShares UK Property UCITS ETF
    'h_comre':  'REM',    # iShares Mortgage Real Estate ETF
    'h_agri':   'MOO',    # VanEck Agribusiness ETF
    'h_infra':  'IGF',    # iShares Global Infrastructure ETF
    'h_reits':  'VNQ',    # Vanguard Real Estate ETF
    # Stocks & Shares
    'k_sp500':  'SPY',    # SPDR S&P 500 ETF Trust
    'k_nas':    'QQQ',    # Invesco QQQ Trust (NASDAQ 100)
    'k_ftse':   'ISF.L',  # iShares Core FTSE 100 UCITS ETF
    'k_em':     'EEM',    # iShares MSCI Emerging Markets ETF
    'k_sc':     'VSS',    # Vanguard FTSE All-World ex-US Small-Cap ETF
    # Metals (ETFs — more reliable than futures via yfinance)
    'm_gold':   'GLD',    # SPDR Gold Shares ETF
    'm_silv':   'SLV',    # iShares Silver Trust ETF
    'm_plat':   'PPLT',   # Aberdeen Physical Platinum ETF
    'm_copp':   'CPER',   # United States Copper Index Fund
    'm_pall':   'PALL',   # Aberdeen Physical Palladium ETF
    # Crypto
    'c_btc':    'BTC-USD',
    'c_eth':    'ETH-USD',
    'c_ada':    'ADA-USD',
    'c_xrp':    'XRP-USD',
    'c_sol':    'SOL-USD',
}


@app.route("/api/portfolio-prices")
def portfolio_prices():
    """Return current prices for all 25 portfolio balancer assets."""
    unique = list(set(_PB_TICKERS.values()))
    price_map = {}
    try:
        raw = yf.download(unique, period='5d', interval='1d',
                          progress=False, auto_adjust=True, threads=True)
        close = raw['Close'] if isinstance(raw.columns, pd.MultiIndex) else raw
        for ticker in unique:
            try:
                s = close[ticker].dropna() if ticker in close.columns else pd.Series(dtype=float)
                if len(s):
                    price_map[ticker] = float(s.iloc[-1])
            except Exception:
                pass
    except Exception:
        pass

    result = {}
    for asset_id, ticker in _PB_TICKERS.items():
        price = price_map.get(ticker)
        result[asset_id] = {
            'ticker': ticker,
            'price':  round(price, 6) if price is not None else None,
            'live':   price is not None,
        }
    return jsonify({'prices': result, 'ts': pd.Timestamp.now().isoformat()})


@app.route("/profile")
def profile_page():
    return send_from_directory("static", "profile.html")


# ── Auth endpoints ───────────────────────────────────────────────────────────

@app.route("/api/register", methods=["POST"])
def api_register():
    data = request.get_json() or {}
    username = data.get("username", "").strip()
    password = data.get("password", "")
    email    = data.get("email", "").strip()

    if not username or not password:
        return jsonify({"error": "Username and password required"}), 400
    if not email:
        return jsonify({"error": "Email address required"}), 400
    if len(password) < 6:
        return jsonify({"error": "Password must be at least 6 characters"}), 400

    users = _load_users()
    if username in users:
        return jsonify({"error": "Username already taken"}), 409

    pw_hash = bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode("utf-8")
    users[username] = {
        "password_hash": pw_hash,
        "preferences": {},
        "tier": "power_user",
        "profile": {
            "email": email,
            "display_name": username,
            "bio": "",
            "investor_type": "beginner",
            "profile_picture": "",
        },
    }
    _save_users(users)

    session.permanent = True
    login_user(User(username, "power_user"), remember=True)
    return jsonify({"success": True, "username": username, "tier": "power_user", "preferences": {}, "landing_page": "/"})


@app.route("/api/login", methods=["POST"])
def api_login():
    data = request.get_json() or {}
    username = data.get("username", "").strip()
    password = data.get("password", "")
    if not username or not password:
        return jsonify({"error": "Username and password required"}), 400

    users = _load_users()
    user_data = users.get(username)
    if not user_data:
        return jsonify({"error": "Invalid credentials"}), 401

    if not bcrypt.checkpw(password.encode("utf-8"), user_data["password_hash"].encode("utf-8")):
        return jsonify({"error": "Invalid credentials"}), 401

    stay_signed_in = user_data.get("preferences", {}).get("stay_signed_in", True)
    session.permanent = bool(stay_signed_in)
    login_user(User(username, user_data.get("tier", "basic")), remember=bool(stay_signed_in))
    return jsonify({
        "success": True,
        "username": username,
        "tier": user_data.get("tier", "basic"),
        "preferences": user_data.get("preferences", {}),
        "landing_page": user_data.get("profile", {}).get("landing_page", "/"),
    })


@app.route("/api/forgot-password", methods=["POST"])
def api_forgot_password():
    data = request.get_json() or {}
    identifier = data.get("identifier", "").strip().lower()
    generic = {"success": True, "message": "If an account matches that username or email, we've sent a reset link."}
    if not identifier:
        return jsonify(generic)

    users = _load_users()
    match_username = None
    for username, user_data in users.items():
        email = (user_data.get("profile", {}) or {}).get("email", "").strip().lower()
        if username.lower() == identifier or (email and email == identifier):
            match_username = username
            break

    if match_username:
        user_data = users[match_username]
        token = secrets.token_urlsafe(32)
        token_hash = hashlib.sha256(token.encode("utf-8")).hexdigest()
        profile = dict(user_data.get("profile", {}) or {})
        profile["reset_token_hash"] = token_hash
        profile["reset_token_expires"] = time.time() + 3600  # 1 hour
        user_data["profile"] = profile
        users[match_username] = user_data
        _save_users(users)

        reset_link = f"{request.host_url.rstrip('/')}/reset-password?u={match_username}&t={token}"
        to_addr = profile.get("email") or match_username
        _send_email(
            to_addr,
            "Reset your Growth Capital Group password",
            f"Someone requested a password reset for the account \"{match_username}\".\n\n"
            f"Reset your password here (valid for 1 hour):\n{reset_link}\n\n"
            f"If you didn't request this, you can safely ignore this email.",
        )

    return jsonify(generic)


@app.route("/api/reset-password", methods=["POST"])
def api_reset_password():
    data = request.get_json() or {}
    username = data.get("username", "").strip()
    token = data.get("token", "")
    new_password = data.get("new_password", "")

    if not username or not token or not new_password:
        return jsonify({"error": "Missing required fields"}), 400
    if len(new_password) < 6:
        return jsonify({"error": "Password must be at least 6 characters"}), 400

    users = _load_users()
    user_data = users.get(username)
    profile = (user_data.get("profile", {}) or {}) if user_data else {}
    stored_hash = profile.get("reset_token_hash")
    expires = profile.get("reset_token_expires", 0)

    if not user_data or not stored_hash or time.time() > expires:
        return jsonify({"error": "This reset link is invalid or has expired. Please request a new one."}), 400

    token_hash = hashlib.sha256(token.encode("utf-8")).hexdigest()
    if not hmac.compare_digest(token_hash, stored_hash):
        return jsonify({"error": "This reset link is invalid or has expired. Please request a new one."}), 400

    user_data["password_hash"] = bcrypt.hashpw(new_password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")
    profile.pop("reset_token_hash", None)
    profile.pop("reset_token_expires", None)
    user_data["profile"] = profile
    users[username] = user_data
    _save_users(users)

    return jsonify({"success": True})


@app.route("/reset-password")
def reset_password_page():
    return send_from_directory("static", "reset-password.html")


@app.route("/api/logout", methods=["POST"])
@login_required
def api_logout():
    logout_user()
    return jsonify({"success": True})


@app.route("/api/me", methods=["GET"])
def api_me():
    if not current_user.is_authenticated:
        return jsonify({
            "authenticated": False, "username": None, "tier": "basic", "alpha_role": None,
            "realtime_base_url": _REALTIME_BASE_URL,
        })
    return jsonify({
        "authenticated": True,
        "username": current_user.id,
        "tier": getattr(current_user, "tier", "basic"),
        "alpha_role": getattr(current_user, "alpha_role", None),
        # Empty until the realtime/ service (docs/scaling-plan.md, Workstream 1/5)
        # is actually deployed — see signal_config.html's engine SSE wiring, which
        # falls back to the existing poll loop whenever this is blank.
        "realtime_base_url": _REALTIME_BASE_URL,
    })


@app.route("/api/admin/set-tier", methods=["POST"])
@login_required
def admin_set_tier():
    if current_user.id != "admin":
        return jsonify({"error": "Admin only"}), 403
    data = request.get_json() or {}
    username = data.get("username", "").strip()
    new_tier = data.get("tier", "").strip()
    if new_tier not in TIER_RANKS:
        return jsonify({"error": f"Invalid tier. Valid options: {list(TIER_RANKS.keys())}"}), 400
    users = _load_users()
    if username not in users:
        return jsonify({"error": "User not found"}), 404
    users[username]["tier"] = new_tier
    _save_users(users)
    return jsonify({"success": True, "username": username, "tier": new_tier})


@app.route("/api/admin/users", methods=["GET"])
@login_required
def admin_list_users():
    if current_user.id != "admin":
        return jsonify({"error": "Admin only"}), 403
    users = _load_users()
    return jsonify({"users": [
        {"username": u, "alpha_role": data.get("alpha_role"), "tier": data.get("tier", "basic")}
        for u, data in sorted(users.items())
    ]})


@app.route("/api/admin/set-alpha-role", methods=["POST"])
@login_required
def admin_set_alpha_role():
    if current_user.id != "admin":
        return jsonify({"error": "Admin only"}), 403
    data = request.get_json() or {}
    username = data.get("username", "").strip()
    role = data.get("alpha_role")
    role = role.strip() if isinstance(role, str) else role
    if role not in ALPHA_ROLES and role not in (None, ""):
        return jsonify({"error": f"Invalid alpha_role. Valid options: {sorted(ALPHA_ROLES)} or null to unassign"}), 400
    users = _load_users()
    if username not in users:
        return jsonify({"error": "User not found"}), 404
    role = role or None
    if role:
        holder = next(
            (u for u, data in users.items() if u != username and data.get("alpha_role") == role),
            None,
        )
        if holder:
            return jsonify({"error": f'Role "{role}" is already assigned to "{holder}". Unassign it there first.'}), 409
    users[username]["alpha_role"] = role
    _save_users(users)
    return jsonify({"success": True, "username": username, "alpha_role": role})


@app.route("/api/save-preferences", methods=["POST"])
@login_required
def save_preferences():
    incoming = request.get_json() or {}
    users = _load_users()
    if current_user.id not in users:
        return jsonify({"error": "User not found"}), 404
    # Shallow-merge onto the existing preferences rather than replacing them outright,
    # so keys this caller doesn't know about (e.g. custom_symbols, saved by the
    # /api/custom-symbols endpoints) survive a save from a page that only manages its
    # own subset of settings.
    existing = users[current_user.id].get("preferences", {}) or {}
    merged = {**existing, **incoming}
    _save_preferences(current_user.id, merged)
    return jsonify({"success": True})


@app.route("/api/load-preferences", methods=["GET"])
@login_required
def load_preferences():
    users = _load_users()
    user_data = users.get(current_user.id, {})
    return jsonify({
        "username": current_user.id,
        "preferences": user_data.get("preferences", {}),
    })


LANDING_PAGE_CHOICES = {
    "/",
    "/learn", "/learn/beginner", "/learn/intermediate", "/learn/pro", "/learn/tools",
    "/tools", "/tools/signals", "/backtester", "/tools/portfolio", "/tools/calculator",
    "/arena", "/arena/market-xi", "/arena/competitions", "/arena/predictions",
    "/alpha", "/alpha/connor", "/alpha/dave", "/alpha/gary", "/alpha/tom", "/alpha/podcast",
    "/partners", "/profile",
}


@app.route("/api/profile", methods=["GET"])
@login_required
def api_get_profile():
    users = _load_users()
    user_data = users.get(current_user.id, {})
    profile = user_data.get("profile", {})
    return jsonify({
        "username":       current_user.id,
        "tier":           user_data.get("tier", "basic"),
        "email":          profile.get("email", ""),
        "display_name":   profile.get("display_name", current_user.id),
        "bio":            profile.get("bio", ""),
        "investor_type":  profile.get("investor_type", "beginner"),
        "profile_picture": profile.get("profile_picture", ""),
        "landing_page":   profile.get("landing_page", "/"),
        "preferences":    user_data.get("preferences", {}),
    })


PRESET_AVATARS = {"preset:beginner", "preset:intermediate", "preset:pro"}


@app.route("/api/profile/update", methods=["POST"])
@login_required
def api_update_profile():
    data = request.get_json() or {}
    users = _load_users()
    if current_user.id not in users:
        return jsonify({"error": "User not found"}), 404
    profile = users[current_user.id].get("profile", {})
    for field in ("email", "display_name", "bio", "investor_type"):
        if field in data:
            profile[field] = str(data[field]).strip()
    if "profile_picture" in data:
        value = str(data["profile_picture"]).strip()
        if value not in PRESET_AVATARS:
            return jsonify({"error": "Invalid avatar selection"}), 400
        profile["profile_picture"] = value
    if "landing_page" in data:
        value = str(data["landing_page"]).strip()
        if value not in LANDING_PAGE_CHOICES:
            return jsonify({"error": "Invalid landing page"}), 400
        profile["landing_page"] = value
    users[current_user.id]["profile"] = profile
    _save_users(users)
    return jsonify({"success": True, "profile": profile})


@app.route("/api/profile/avatar", methods=["POST"])
@login_required
def api_upload_avatar():
    if "avatar" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["avatar"]
    if not file.filename:
        return jsonify({"error": "No file selected"}), 400
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_IMAGE_EXTENSIONS:
        return jsonify({"error": "Allowed types: png, jpg, jpeg, gif, webp"}), 400
    users = _load_users()
    if current_user.id not in users:
        return jsonify({"error": "User not found"}), 404
    filename = secure_filename(f"{current_user.id}.{ext}")
    avatar_url = f"/api/profile/avatar/{current_user.id}"
    profile = users[current_user.id].get("profile", {})
    profile["profile_picture"] = avatar_url
    users[current_user.id]["profile"] = profile
    _save_users(users)
    # set_user_avatar() does its own load/save cycle in the local-JSON-fallback
    # path — must run after the profile_picture save above, not before, or its
    # internal _save_users() call would clobber this one's write and vice versa.
    set_user_avatar(current_user.id, filename, file.read())
    return jsonify({"success": True, "avatar_url": avatar_url})


@app.route("/api/profile/avatar/<username>", methods=["GET"])
def api_get_avatar(username):
    filename, file_bytes = get_user_avatar(username)
    if file_bytes is None:
        return jsonify({"error": "No avatar"}), 404
    import mimetypes
    mimetype = mimetypes.guess_type(filename or "")[0] or "image/jpeg"
    return Response(file_bytes, mimetype=mimetype, headers={"Cache-Control": "public, max-age=3600"})


# ── Alpha content API ─────────────────────────────────────────────────────────

def _alpha_can_touch(item) -> bool:
    return item is not None and item.get("author") == current_user.alpha_role


# ── Market XI asset auto-linking ────────────────────────────────────────────
# Placeholder integration: Market XI (the fantasy-team game at
# https://market-xi-live.vercel.app/) doesn't have per-asset profile pages
# yet, so every match points at the site's root for now, and the asset list
# itself is drawn from GCG's own watchlist data (see
# _market_xi_asset_aliases below) rather than from Market XI directly.
#
# TODO(market-xi-integration): Market XI is a Next.js + Supabase app. It has
# at least one public, unauthenticated JSON route — GET /api/gameweek-timeline
# — confirmed reachable during investigation, so a live asset feed is
# plausible once they add one; no equivalent asset/roster endpoint was found
# (checked common paths like /api/assets, /api/asset-catalogue,
# /api/market-assets — all 404). Their frontend references internal store
# names "assetMarketAssets"/"assetMarketCatalogue", confirming the *feature*
# exists, just not a discoverable public route for it yet.
#
# Once one exists (ideally something like GET /api/asset-catalogue returning
# {"assets": [{"slug", "name", "ticker", "url"}, ...]}), swap this out for a
# fetch from Market XI with a short cache (their roster changes at most a few
# times a week) and a fallback to today's watchlist-derived list if the
# request fails — Market XI being briefly unreachable shouldn't take Studio's
# asset linking down with it. At that point MARKET_XI_URL below becomes a
# per-asset URL (from that "url" field) instead of one constant for
# everything.
MARKET_XI_URL = "https://market-xi-live.vercel.app/"


def _market_xi_asset_aliases() -> dict:
    """Every distinct asset name/ticker mentioned across all 4 partners'
    published watchlist items — the only place "assets" are named anywhere
    on the site today, so it's the whole source of truth for what counts as
    a linkable asset (see the TODO above for the plan to source this from
    Market XI directly instead, once it can be). Returns {lowercased alias:
    display title}. A title like "Cardano (ADA)" contributes three aliases
    (the full title, "Cardano", and "ADA") so a post mentioning any of those
    forms gets linked.
    """
    aliases = {}
    for slug in ALPHA_ROLES:
        for item in alpha_content_list(author=slug, status="published"):
            if item.get("kind") != "watchlist":
                continue
            title = (item.get("title") or "").strip()
            if not title:
                continue
            aliases.setdefault(title.lower(), title)
            m = re.match(r"^(.+?)\s*\(([A-Za-z0-9.]+)\)$", title)
            if m:
                name, ticker = m.group(1).strip(), m.group(2).strip()
                if name:
                    aliases.setdefault(name.lower(), title)
                if ticker:
                    aliases.setdefault(ticker.lower(), title)
    return aliases


def _auto_link_market_xi_assets(body: str) -> str:
    """Wraps the first mention of each known asset (see
    _market_xi_asset_aliases) in a post body with a [text](url) markdown
    link to Market XI — the same [text](url) syntax the Studio's own
    formatting toolbar and renderInlineText() already support, so this
    needs no rendering changes on the public post page.

    Deliberately conservative: whole-word matches only, longest alias first
    (so "Cardano (ADA)" isn't fragmented by its own shorter "ADA" alias
    matching a piece of it first), skips anything already inside a markdown
    link or image so re-running this on an already-linked body is a no-op,
    and links only the first occurrence of each asset per post rather than
    every mention.
    """
    if not body:
        return body
    aliases = _market_xi_asset_aliases()
    if not aliases:
        return body

    # Longest-first so multi-word titles win over a shorter alias they contain.
    ordered = sorted(aliases.items(), key=lambda kv: -len(kv[0]))
    pattern = re.compile(
        r"(?<![\w\-])(" + "|".join(re.escape(a) for a, _ in ordered) + r")(?![\w\-])",
        re.IGNORECASE,
    )
    # Spans already covered by an existing [text](url) or ![alt](url) — never
    # match inside one, whether from a previous auto-link pass or an author's
    # own manual link.
    protected = [m.span() for m in re.finditer(r"!?\[[^\]]*\]\([^)]*\)", body)]

    def is_protected(start, end):
        return any(start < p_end and end > p_start for p_start, p_end in protected)

    linked_lower = set()
    out = []
    last_end = 0
    for m in pattern.finditer(body):
        start, end = m.span()
        matched_lower = m.group(1).lower()
        canonical = aliases.get(matched_lower)
        if not canonical or canonical.lower() in linked_lower or is_protected(start, end):
            continue
        out.append(body[last_end:start])
        out.append(f"[{m.group(1)}]({MARKET_XI_URL})")
        last_end = end
        linked_lower.add(canonical.lower())
    out.append(body[last_end:])
    return "".join(out)


@app.route("/api/alpha/upload", methods=["POST"])
@login_required
@alpha_author_required
def api_alpha_upload():
    author = current_user.alpha_role
    kind = (request.form.get("kind") or "post").strip()
    if kind not in ALPHA_CONTENT_KINDS:
        return jsonify({"error": f"Invalid kind. Valid options: {sorted(ALPHA_CONTENT_KINDS)}"}), 400
    requested_topic = (request.form.get("topic") or "").strip() or None
    requested_level = (request.form.get("level") or "").strip().lower() or None
    if requested_level and requested_level not in ALPHA_LEVELS:
        return jsonify({"error": f"Level must be one of {sorted(ALPHA_LEVELS)}"}), 400

    file_bytes = None
    source_filename = None
    structured_sections = None
    structured_meta = {}
    try:
        if "file" in request.files and request.files["file"].filename:
            f = request.files["file"]
            source_filename = secure_filename(f.filename)
            file_bytes = f.read()
            f.seek(0)
            raw_text = extract_text_from_upload(f)
            source_kind = "file"
            file_ext = source_filename.rsplit(".", 1)[-1].lower() if "." in source_filename else ""
            if file_ext == "docx" and kind == "post":
                f.seek(0)
                try:
                    result = extract_structured_docx(f)
                    if result is not None:
                        structured_sections, structured_meta = result
                except Exception:
                    structured_sections = None  # not a template doc, or unreadable structure — fall back below
        elif (request.form.get("link") or "").strip():
            raw_text = _strip_code_fence(extract_text_from_url(request.form.get("link").strip()))
            source_kind = "link"
        elif (request.form.get("paste") or "").strip():
            raw_text = _strip_code_fence(request.form.get("paste").strip()[:MAX_UPLOAD_TEXT_CHARS])
            source_kind = "paste"
        else:
            return jsonify({"error": "Provide a file, a link, or pasted text"}), 400
        if kind == "post" and source_kind in ("paste", "link"):
            try:
                result = extract_structured_text(raw_text)
                if result is not None:
                    structured_sections, structured_meta = result
            except Exception:
                structured_sections = None  # didn't use the "## " convention, or something unexpected in it — fall back below
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    author_topics = ALPHA_TOPICS.get(author, [])
    normalized = normalize_content(raw_text, author, author_topics)
    topic = requested_topic if requested_topic in author_topics else normalized["topic"]
    body = auto_section_tables(normalized["body"]) if kind == "post" else normalized["body"]
    if kind == "post":
        body = _auto_link_market_xi_assets(body)
    # normalize_content's title/subtitle/snippet are a naive split of the
    # whole flattened document — a structured upload has much better
    # candidates: the first Normal section's own heading for the title, and
    # any explicit "[SUBTITLE]"/"[SNIPPET]" tagged block for those fields.
    if structured_sections and structured_sections[0].get("type") == "normal" and structured_sections[0].get("heading"):
        normalized["title"] = _strip_markdown_links(structured_sections[0]["heading"])[:100]
    if structured_meta.get("subtitle"):
        normalized["subtitle"] = _strip_markdown_links(structured_meta["subtitle"])[:140]
    if structured_meta.get("snippet"):
        normalized["snippet"] = structured_meta["snippet"][:280]

    fields = {
        "author": author, "kind": kind, "status": "draft", "topic": topic,
        "level": requested_level or "beginner",
        "title": normalized["title"], "subtitle": normalized["subtitle"],
        "snippet": normalized["snippet"], "body": body,
        "stance": None, "url": None,
        "source_kind": source_kind, "source_filename": source_filename, "source_text": raw_text,
    }
    item = alpha_content_create(fields, file_bytes=file_bytes)

    if structured_sections:
        # Images inside a structured doc's [IMAGE] sections need the new
        # item's id to attach to, which only exists after the create above —
        # hence building this replacement body as a second step.
        def _upload_extracted_image(img_bytes, img_ext):
            attachment_id = alpha_attachment_create(item["id"], f"template-image.{img_ext}", img_bytes)
            return f"/api/alpha/content/{item['id']}/images/{attachment_id}"
        structured_body = _serialize_structured_sections(structured_sections, _upload_extracted_image)
        if structured_body.strip():
            item = alpha_content_update(item["id"], {"body": _auto_link_market_xi_assets(structured_body)})

    return jsonify({"success": True, "item": item})


@app.route("/api/alpha/content", methods=["GET", "POST"])
@login_required
@alpha_author_required
def api_alpha_content():
    author = current_user.alpha_role
    if request.method == "GET":
        return jsonify({"items": alpha_content_list(author=author)})

    data = request.get_json() or {}
    kind = (data.get("kind") or "").strip()
    if kind not in ALPHA_CONTENT_KINDS:
        return jsonify({"error": f"Invalid kind. Valid options: {sorted(ALPHA_CONTENT_KINDS)}"}), 400
    if kind == "post":
        return jsonify({"error": "Posts are created via the upload endpoint, not this one"}), 400

    topic = (data.get("topic") or "").strip() or None
    if topic and topic not in ALPHA_TOPICS.get(author, []):
        return jsonify({"error": "Topic must be one of your nominated topics"}), 400

    level = (data.get("level") or "").strip().lower() or None
    if level and level not in ALPHA_LEVELS:
        return jsonify({"error": f"Level must be one of {sorted(ALPHA_LEVELS)}"}), 400

    title = (data.get("title") or "").strip()
    if not title:
        return jsonify({"error": "Title is required"}), 400

    fields = {
        "author": author, "kind": kind, "status": "draft", "topic": topic,
        "level": level or "beginner",
        "title": title,
        "snippet": (data.get("snippet") or "").strip() or None,
        "body": (data.get("body") or "").strip() or None,
        "stance": None, "url": None,
        "source_kind": "manual", "source_filename": None, "source_text": None,
    }
    if kind == "watchlist":
        stance = (data.get("stance") or "").strip().lower()
        if stance not in ALPHA_STANCES:
            return jsonify({"error": f"Stance must be one of {sorted(ALPHA_STANCES)}"}), 400
        fields["stance"] = stance
    if kind in ("video", "link"):
        url = (data.get("url") or "").strip()
        if not url:
            return jsonify({"error": f"A URL is required for a {kind}"}), 400
        fields["url"] = url

    item = alpha_content_create(fields)
    return jsonify({"success": True, "item": item})


@app.route("/api/alpha/content/<int:item_id>", methods=["PUT", "DELETE"])
@login_required
@alpha_author_required
def api_alpha_content_item(item_id):
    existing = alpha_content_get(item_id)
    if not _alpha_can_touch(existing):
        return jsonify({"error": "Not found"}), 404

    if request.method == "DELETE":
        alpha_content_delete(item_id)
        return jsonify({"success": True})

    data = request.get_json() or {}

    # Fields a staged edit can hold and fold back into the live item — shared by
    # the "save while published" path below, republish, and the unpublish fold.
    _stageable_fields = ("topic", "level", "title", "subtitle", "snippet", "body", "url", "stance", "image_url", "image_filename")

    if data.get("republish"):
        # Push any pending staged edits live in one step, without a detour
        # through "draft" — an alternative to Unpublish-then-Publish when the
        # author just wants their saved edits to go live.
        if existing.get("status") != "published":
            return jsonify({"error": "Only a published item can be re-published"}), 400
        fold = {"staged_edits": None}
        for k, v in (existing.get("staged_edits") or {}).items():
            if k in _stageable_fields:
                fold[k] = v
        item = alpha_content_update(item_id, fold)
        return jsonify({"success": True, "item": item})

    updates = {}
    for field in ("topic", "title", "subtitle", "snippet", "body", "url"):
        if field in data:
            value = data[field]
            updates[field] = (str(value).strip() or None) if value is not None else None
    if "body" in updates and updates["body"] and existing.get("kind") == "post":
        updates["body"] = _auto_link_market_xi_assets(updates["body"])
    if "level" in data:
        level = (data["level"] or "").strip().lower()
        if level and level not in ALPHA_LEVELS:
            return jsonify({"error": f"Level must be one of {sorted(ALPHA_LEVELS)}"}), 400
        updates["level"] = level or None
    if data.get("clear_image"):
        # Remove the hero image entirely — both a pasted URL and any uploaded file.
        updates["image_url"] = None
        updates["image_file"] = None
        updates["image_filename"] = None
    elif "image_url" in data:
        image_url = (data["image_url"] or "").strip()
        updates["image_url"] = image_url or None
        if image_url:
            # Switching to URL mode clears any previously uploaded file.
            updates["image_file"] = None
            updates["image_filename"] = None
    if "stance" in data:
        stance = (data["stance"] or "").strip().lower()
        if stance and stance not in ALPHA_STANCES:
            return jsonify({"error": f"Stance must be one of {sorted(ALPHA_STANCES)}"}), 400
        updates["stance"] = stance or None
    if updates.get("topic") and updates["topic"] not in ALPHA_TOPICS.get(current_user.alpha_role, []):
        return jsonify({"error": "Topic must be one of your nominated topics"}), 400

    if "pinned" in data:
        if existing.get("kind") != "post":
            return jsonify({"error": "Only posts can be pinned"}), 400
        pinned = bool(data["pinned"])
        if pinned:
            other_pinned = [
                i for i in alpha_content_list(author=existing["author"])
                if i.get("kind") == "post" and i.get("pinned") and i["id"] != item_id
            ]
            if len(other_pinned) >= 4:
                return jsonify({"error": "You can only pin up to 4 posts at a time. Unpin one first."}), 400
        # Applied immediately, independent of the staged-edits path below —
        # `pinned` isn't in `_stageable_fields`, so if it stayed in `updates`
        # here, a pin/unpin sent alongside content fields on a published post
        # (the studio's Pin button always sends gatherFieldUpdates() too — see
        # its onclick — to avoid clobbering unsaved text) would silently vanish:
        # the staged-edits branch returns early after writing only
        # `staged_edits`, never reaching the `alpha_content_update(updates)`
        # call that would have persisted it. That forced an unpublish → pin →
        # republish workaround. Pinning is a curation flag, not content, so it
        # should take effect right away regardless of draft/published/staged
        # state.
        alpha_content_update(item_id, {"pinned": pinned})
        existing["pinned"] = pinned

    if "related_lesson_slug" in data:
        # A curation link, not content — applied immediately like `pinned`
        # above, independent of the draft/published/staged-edits dance, so
        # confirming a match doesn't require an unpublish/republish cycle.
        slug = (data.get("related_lesson_slug") or "").strip() or None
        if slug and slug not in _LESSON_BY_SLUG:
            return jsonify({"error": "Unknown lesson slug"}), 400
        note = (data.get("related_lesson_note") or "").strip() or None
        alpha_content_update(item_id, {"related_lesson_slug": slug, "related_lesson_note": note if slug else None})
        existing["related_lesson_slug"] = slug
        existing["related_lesson_note"] = note if slug else None

    # ── Staged edits ────────────────────────────────────────────────────────
    # Editing a PUBLISHED item (a content change with no status change) saves to
    # a pending `staged_edits` copy and leaves the live page untouched. The edits
    # are folded into the item when it's next unpublished (see below), so the
    # path to make them live is: edit → Save (staged) → Unpublish → Publish.
    _content_edit_keys = ("topic", "level", "title", "subtitle", "snippet", "body", "url", "stance", "image_url", "clear_image")
    if existing.get("status") == "published" and "status" not in data and any(k in data for k in _content_edit_keys):
        staged = dict(existing.get("staged_edits") or {})
        # image_filename is included so `clear_image` can stage removing an
        # uploaded hero image (setting it to None) — cheap, since it never holds
        # new binary data. A brand-new *upload* is blocked before it reaches here
        # (see /image endpoint below): only removal is staged, not replacement.
        for field in _stageable_fields:
            if field in updates:
                staged[field] = updates[field]
        item = alpha_content_update(item_id, {"staged_edits": staged})
        return jsonify({"success": True, "item": item})

    if "status" in data:
        status = data["status"]
        if status not in ("draft", "published"):
            return jsonify({"error": "Status must be 'draft' or 'published'"}), 400
        if status == "published":
            merged = {**existing, **updates}
            if merged.get("kind") == "post":
                missing = []
                if not (merged.get("title") or "").strip():
                    missing.append("title")
                if not (merged.get("subtitle") or "").strip():
                    missing.append("subtitle")
                if not ((merged.get("image_url") or "").strip() or merged.get("image_filename")):
                    missing.append("image")
                if missing:
                    return jsonify({
                        "error": "Missing before publishing: " + ", ".join(missing) + ".",
                        "missing_fields": missing,
                    }), 400
        updates["status"] = status
        updates["published_at"] = _dt.datetime.utcnow() if status == "published" else None
        # Unpublishing folds any staged edits into the (now draft) live fields so
        # they're preserved and ready to go live again on the next publish. A
        # field already present in `updates` (the studio sends the whole form,
        # not a diff) is the author's current on-screen text and wins over an
        # older staged copy of that same field.
        if status == "draft" and existing.get("staged_edits"):
            for k, v in existing["staged_edits"].items():
                if k in _stageable_fields and k not in updates:
                    updates[k] = v
            updates["staged_edits"] = None

    item = alpha_content_update(item_id, updates)
    return jsonify({"success": True, "item": item})


@app.route("/api/alpha/content/<int:item_id>/file", methods=["GET"])
@login_required
@alpha_author_required
def api_alpha_content_file(item_id):
    existing = alpha_content_get(item_id)
    if not _alpha_can_touch(existing):
        return jsonify({"error": "Not found"}), 404
    filename, file_bytes = alpha_content_get_file(item_id)
    if file_bytes is None:
        return jsonify({"error": "No file attached to this item"}), 404
    import mimetypes
    mimetype = mimetypes.guess_type(filename or "")[0] or "application/octet-stream"
    return Response(
        file_bytes, mimetype=mimetype,
        headers={"Content-Disposition": f"attachment; filename={secure_filename(filename or 'download')}"}
    )


@app.route("/api/alpha/content/<int:item_id>/image", methods=["GET"])
def api_alpha_content_image(item_id):
    item = alpha_content_get(item_id)
    if not item:
        return jsonify({"error": "Not found"}), 404
    is_owner = current_user.is_authenticated and getattr(current_user, "alpha_role", None) == item.get("author")
    if item.get("status") != "published" and not is_owner:
        return jsonify({"error": "Not found"}), 404
    filename, file_bytes = alpha_content_get_image(item_id)
    if file_bytes is None:
        return jsonify({"error": "No image attached"}), 404
    import mimetypes
    mimetype = mimetypes.guess_type(filename or "")[0] or "image/jpeg"
    cache = "public, max-age=3600" if item.get("status") == "published" else "private, no-store"
    return Response(file_bytes, mimetype=mimetype, headers={"Cache-Control": cache})


@app.route("/api/alpha/content/<int:item_id>/image", methods=["POST"])
@login_required
@alpha_author_required
def api_alpha_content_image_upload(item_id):
    existing = alpha_content_get(item_id)
    if not _alpha_can_touch(existing):
        return jsonify({"error": "Not found"}), 404
    if existing.get("kind") != "post":
        return jsonify({"error": "Images can only be attached to posts"}), 400
    if existing.get("status") == "published":
        # This is the single hero-image slot, served live at .../image — unlike
        # text/topic/URL edits it isn't staged, so replacing it here would go
        # straight to the website. Unpublish first, or use an Image URL, which
        # does stage.
        return jsonify({"error": "Unpublish this post first to upload a new hero image, or use an Image URL instead."}), 400
    if "image" not in request.files or not request.files["image"].filename:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["image"]
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_IMAGE_EXTENSIONS:
        return jsonify({"error": "Allowed types: png, jpg, jpeg, gif, webp"}), 400
    item = alpha_content_set_image(item_id, secure_filename(file.filename), file.read())
    return jsonify({"success": True, "item": item})


@app.route("/api/alpha/content/<int:item_id>/images", methods=["POST"])
@login_required
@alpha_author_required
def api_alpha_content_attachment_upload(item_id):
    """Uploads an inline image to embed mid-post (distinct from the single main
    image slot on the row itself — a post can have any number of these)."""
    existing = alpha_content_get(item_id)
    if not _alpha_can_touch(existing):
        return jsonify({"error": "Not found"}), 404
    if existing.get("kind") != "post":
        return jsonify({"error": "Images can only be attached to posts"}), 400
    if "image" not in request.files or not request.files["image"].filename:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["image"]
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_IMAGE_EXTENSIONS:
        return jsonify({"error": "Allowed types: png, jpg, jpeg, gif, webp"}), 400
    attachment_id = alpha_attachment_create(item_id, secure_filename(file.filename), file.read())
    return jsonify({"success": True, "url": f"/api/alpha/content/{item_id}/images/{attachment_id}"})


@app.route("/api/alpha/content/<int:item_id>/images/<int:attachment_id>", methods=["GET"])
def api_alpha_content_attachment_get(item_id, attachment_id):
    content_id, filename, file_bytes = alpha_attachment_get(attachment_id)
    if content_id != item_id or file_bytes is None:
        return jsonify({"error": "Not found"}), 404
    item = alpha_content_get(item_id)
    is_owner = current_user.is_authenticated and getattr(current_user, "alpha_role", None) == (item or {}).get("author")
    if not item or (item.get("status") != "published" and not is_owner):
        return jsonify({"error": "Not found"}), 404
    import mimetypes
    mimetype = mimetypes.guess_type(filename or "")[0] or "image/jpeg"
    cache = "public, max-age=3600" if item.get("status") == "published" else "private, no-store"
    return Response(file_bytes, mimetype=mimetype, headers={"Cache-Control": cache})


def _alpha_public_image_url(item: dict) -> str | None:
    if item.get("image_filename"):
        return f"/api/alpha/content/{item['id']}/image"
    return item.get("image_url")


_ALPHA_PUBLIC_FIELDS = ["id", "kind", "topic", "level", "title", "subtitle", "snippet", "body", "stance", "url", "published_at", "pinned"]

_LESSON_BY_SLUG = {l["slug"]: l for l in LESSON_PAGES}


def _alpha_public_related_lesson(item: dict) -> dict | None:
    """The confirmed Learn lesson this post relates to, if any — rendered as a
    "Learn more" link on the public Alpha post page. None once the lesson is
    removed from LESSON_PAGES, so a stale relation just silently stops showing
    rather than 404ing."""
    lesson = _LESSON_BY_SLUG.get(item.get("related_lesson_slug"))
    if not lesson:
        return None
    return {
        "slug": lesson["slug"], "level": lesson["level"], "title": lesson["title"],
        "url": f"/learn/{lesson['level']}/{lesson['slug']}",
        "note": item.get("related_lesson_note"),
    }


@app.route("/api/alpha/<slug>/content", methods=["GET"])
def api_alpha_public_content(slug):
    if slug not in ALPHA_ROLES:
        return jsonify({"error": "Unknown author"}), 404
    topic = (request.args.get("topic") or "").strip() or None
    items = alpha_content_list(author=slug, status="published")

    def to_public(item):
        public_item = {k: item.get(k) for k in _ALPHA_PUBLIC_FIELDS}
        public_item["image_url"] = _alpha_public_image_url(item)
        public_item["related_lesson"] = _alpha_public_related_lesson(item)
        return public_item

    # Pinned posts fill a fixed 4-slot strip above the main grid, scoped to
    # whichever topic pill the reader has selected — a post pinned while
    # categorized as e.g. "ETFs" should only appear pinned on that topic, not
    # on "All" or other topic pills. On "All" (no topic filter), every pinned
    # post across topics is eligible.
    pinned_source = [i for i in items if i.get("topic") == topic] if topic else items
    pinned_posts = [to_public(i) for i in pinned_source if i.get("kind") == "post" and i.get("pinned")][:4]
    pinned_ids = {p["id"] for p in pinned_posts}

    if topic:
        items = [i for i in items if i.get("topic") == topic]
    grouped = {"watchlist": [], "video": [], "post": [], "link": []}
    for item in items:
        if item["kind"] == "post" and item["id"] in pinned_ids:
            continue
        grouped.setdefault(item["kind"], []).append(to_public(item))
    # Posts are the partner's full research history, not a preview strip —
    # capping them (as watchlist/video intentionally are, below) meant a
    # reader clicking "All" or a topic pill couldn't see everything actually
    # published; see the report that led to removing this cap.
    caps = {"watchlist": 3, "video": 4}
    for kind, cap in caps.items():
        grouped[kind] = grouped[kind][:cap]
    grouped["pinned_posts"] = pinned_posts
    return jsonify(grouped)


@app.route("/api/alpha/<slug>/post/<int:post_id>", methods=["GET"])
def api_alpha_public_post(slug, post_id):
    if slug not in ALPHA_ROLES:
        return jsonify({"error": "Unknown author"}), 404
    item = alpha_content_get(post_id)
    if not item or item["author"] != slug or item["kind"] != "post" or item["status"] != "published":
        return jsonify({"error": "Post not found"}), 404
    public_item = {k: item.get(k) for k in _ALPHA_PUBLIC_FIELDS}
    public_item["image_url"] = _alpha_public_image_url(item)
    public_item["related_lesson"] = _alpha_public_related_lesson(item)
    return jsonify(public_item)


@app.route("/api/lessons", methods=["GET"])
def api_lessons():
    """Slug/level/title for every /learn lesson — feeds the Studio's "Related
    lesson" picker so it doesn't need its own hardcoded copy of LESSON_PAGES."""
    return jsonify({"items": [{"slug": l["slug"], "level": l["level"], "title": l["title"]} for l in LESSON_PAGES]})


@app.route("/api/lessons/<slug>/related-alpha", methods=["GET"])
def api_lesson_related_alpha(slug):
    """Published Alpha content that's been confirmed as related to this lesson —
    powers the "From the Alpha desk" callout on the public lesson page. Public,
    read-only, no auth: same trust level as the lesson page itself."""
    if slug not in _LESSON_BY_SLUG:
        return jsonify({"error": "Unknown lesson"}), 404
    items = [
        i for i in alpha_content_list(status="published")
        if i.get("related_lesson_slug") == slug and i.get("kind") == "post"
    ]
    return jsonify({"items": [
        {
            "id": i["id"], "author": i["author"], "title": i.get("title"),
            "subtitle": i.get("subtitle"), "topic": i.get("topic"),
            "note": i.get("related_lesson_note"),
            "url": f"/alpha/{i['author']}/post/{i['id']}",
        }
        for i in items
    ]})


@app.route("/alpha/<slug>/post/<int:post_id>")
def alpha_post_page(slug, post_id):
    return send_from_directory("static", "alpha-post.html")


# ── Site search ───────────────────────────────────────────────────────────
# A hand-maintained index of the fixed marketing/tool pages (there's no CMS
# behind these, so a static list is the whole "database"). Lesson pages are
# NOT listed here — they're generated from LESSON_PAGES below, so adding a
# lesson there is enough to make it searchable too. Alpha posts aren't listed
# here either: api_search() queries alpha_content_list() live, so newly
# published posts show up automatically, no index update needed at all.
# Update this list alongside the "New pages" step in CLAUDE.md whenever a
# non-lesson route is added.
SEARCH_PAGE_INDEX = [
    {"title": "Home", "url": "/", "sub": "Growth Capital Group"},
    {"title": "Roadmap", "url": "/roadmap", "sub": "Progress towards MVP"},
    {"title": "Company", "url": "/company", "sub": "Growth Capital Group structure"},
    {"title": "Signal Tracker — Full Guide", "url": "/guides/signals", "sub": "Illustrated user guide"},
    {"title": "Backtester — Full Guide", "url": "/guides/backtester", "sub": "Illustrated user guide"},
    {"title": "Portfolio Balancer — Full Guide", "url": "/guides/portfolio", "sub": "Illustrated user guide"},
    {"title": "Learn", "url": "/learn", "sub": "Lessons for every level"},
    {"title": "Learn — Beginner", "url": "/learn/beginner", "sub": "Lesson hub"},
    {"title": "Learn — Intermediate", "url": "/learn/intermediate", "sub": "Lesson hub"},
    {"title": "Learn — Pro", "url": "/learn/pro", "sub": "Lesson hub"},
    {"title": "Learn — Tools", "url": "/learn/tools", "sub": "How to use each tool"},
    {"title": "Tools", "url": "/tools", "sub": "Tools hub"},
    {"title": "Signals", "url": "/tools/signals", "sub": "Tool"},
    {"title": "Backtester", "url": "/backtester", "sub": "Tool"},
    {"title": "Portfolio Manager", "url": "/tools/portfolio", "sub": "Tool"},
    {"title": "Calculator", "url": "/tools/calculator", "sub": "Tool"},
    {"title": "Data Visualisation", "url": "/tools/data-visualisation", "sub": "Tool"},
    {"title": "The Arena", "url": "/arena", "sub": "Arena hub"},
    {"title": "Market XI", "url": "/arena/market-xi", "sub": "The Arena"},
    {"title": "Trading Competitions", "url": "/arena/competitions", "sub": "The Arena"},
    {"title": "Predictions Markets", "url": "/arena/predictions", "sub": "The Arena"},
    {"title": "Alpha", "url": "/alpha", "sub": "Alpha hub"},
    {"title": "Connor", "url": "/alpha/connor", "sub": "Alpha partner"},
    {"title": "Dave", "url": "/alpha/dave", "sub": "Alpha partner"},
    {"title": "Gary", "url": "/alpha/gary", "sub": "Alpha partner"},
    {"title": "Tom", "url": "/alpha/tom", "sub": "Alpha partner"},
    {"title": "Alpha Podcast", "url": "/alpha/podcast", "sub": "Alpha"},
    {"title": "Partners", "url": "/partners", "sub": "Growth Capital Group"},
]


def _search_score(query: str, title: str, sub: str) -> int:
    """Higher is better; 0 means no match. Cheap substring ranking — see the
    ALPHA_LEVELS-style comment in the search options discussion: this is
    deliberately simple, not full-text search, since the site's whole
    content set is a few dozen pages plus a handful of Alpha posts."""
    t, s, q = (title or "").lower(), (sub or "").lower(), query.lower()
    if t == q:
        return 100
    if t.startswith(q):
        return 80
    if q in t:
        return 60
    if q in s:
        return 30
    return 0


def _lesson_search_pages():
    return [
        {
            "title": lesson["title"],
            "url": f"/learn/{lesson['level']}/{lesson['slug']}",
            "sub": f"{lesson['level'].capitalize()} lesson",
        }
        for lesson in LESSON_PAGES
    ]


@app.route("/api/search")
def api_search():
    query = (request.args.get("q") or "").strip()
    if len(query) < 2:
        return jsonify({"groups": []})

    page_matches = []
    for page in SEARCH_PAGE_INDEX + _lesson_search_pages():
        score = _search_score(query, page["title"], page.get("sub", ""))
        if score:
            page_matches.append((score, {"title": page["title"], "url": page["url"], "sub": page.get("sub")}))
    page_matches.sort(key=lambda x: -x[0])

    alpha_matches = []
    for slug in ALPHA_ROLES:
        for item in alpha_content_list(author=slug, status="published"):
            if item.get("kind") != "post":
                continue
            title = item.get("title") or ""
            sub = item.get("subtitle") or item.get("snippet") or ""
            score = _search_score(query, title, sub)
            if not score and query.lower() in (item.get("body") or "").lower():
                score = 15
            if score:
                author_label = slug.capitalize()
                alpha_matches.append((score, {
                    "title": title,
                    "url": f"/alpha/{slug}/post/{item['id']}",
                    "sub": f"{author_label} — {sub}" if sub else author_label,
                }))
    alpha_matches.sort(key=lambda x: -x[0])

    return jsonify({"groups": [
        {"label": "Pages", "items": [m[1] for m in page_matches[:6]]},
        {"label": "Alpha", "items": [m[1] for m in alpha_matches[:6]]},
    ]})


@app.route("/dataviz-studio")
def dataviz_studio(): return send_from_directory("static", "dataviz-studio.html")

@app.route("/alpha/studio")
def alpha_studio():
    # Client-side checks /api/me for alpha_role, same convention as /profile —
    # no server-side @login_required here since page routes in this app rely
    # on the JS auth check rather than a redirect-on-401 pattern.
    return send_from_directory("static", "alpha-studio.html")


@app.route("/api/dataviz/content", methods=["GET", "POST"])
@login_required
@dataviz_author_required
def api_dataviz_content():
    if request.method == "GET":
        status = (request.args.get("status") or "").strip() or None
        page = (request.args.get("page") or "").strip() or None
        return jsonify({"items": dataviz_content_list(status=status, page=page)})
    data = request.get_json() or {}
    page = (data.get("page") or "").strip()
    if not dataviz_page_get(page):
        return jsonify({"error": "Unknown page — create it first"}), 400
    fields = {
        "author": current_user.alpha_role,
        "page": page,
        "status": "draft",
        "title": (data.get("title") or "").strip() or None,
        "description": (data.get("description") or "").strip() or None,
        "positive_analysis": (data.get("positive_analysis") or "").strip() or None,
        "warning": (data.get("warning") or "").strip() or None,
        "link": (data.get("link") or "").strip() or None,
        "image_filename": None,
    }
    item = dataviz_content_create(fields)
    return jsonify({"success": True, "item": item})


@app.route("/api/dataviz/content/<int:item_id>", methods=["PUT", "DELETE"])
@login_required
@dataviz_author_required
def api_dataviz_content_item(item_id):
    existing = dataviz_content_get(item_id)
    if not existing or existing.get("author") != current_user.alpha_role:
        return jsonify({"error": "Not found"}), 404
    if request.method == "DELETE":
        dataviz_content_delete(item_id)
        return jsonify({"success": True})
    data = request.get_json() or {}
    updates = {}
    for key in ("title", "description", "positive_analysis", "warning", "link"):
        if key in data:
            updates[key] = (data.get(key) or "").strip() or None
    if "page" in data:
        if not dataviz_page_get(data["page"]):
            return jsonify({"error": "Unknown page — create it first"}), 400
        updates["page"] = data["page"]
    if "status" in data and data["status"] in ("draft", "published"):
        target_page = dataviz_page_get(updates.get("page", existing.get("page")))
        page_has_widget = bool(target_page and target_page.get("has_live_widget"))
        if data["status"] == "published" and not existing.get("image_filename") and not page_has_widget:
            return jsonify({"error": "Upload an image before publishing"}), 400
        updates["status"] = data["status"]
        updates["published_at"] = _dt.datetime.utcnow() if data["status"] == "published" else None
    item = dataviz_content_update(item_id, updates)
    return jsonify({"success": True, "item": item})


@app.route("/api/dataviz/content/<int:item_id>/image", methods=["GET"])
def api_dataviz_content_image(item_id):
    item = dataviz_content_get(item_id)
    if not item:
        return jsonify({"error": "Not found"}), 404
    is_owner = current_user.is_authenticated and getattr(current_user, "alpha_role", None) == item.get("author")
    if item.get("status") != "published" and not is_owner:
        return jsonify({"error": "Not found"}), 404
    filename, file_bytes = dataviz_content_get_image(item_id)
    if file_bytes is None:
        return jsonify({"error": "No image attached"}), 404
    import mimetypes
    mimetype = mimetypes.guess_type(filename or "")[0] or "image/jpeg"
    cache = "public, max-age=3600" if item.get("status") == "published" else "private, no-store"
    return Response(file_bytes, mimetype=mimetype, headers={"Cache-Control": cache})


@app.route("/api/dataviz/content/<int:item_id>/image", methods=["POST"])
@login_required
@dataviz_author_required
def api_dataviz_content_image_upload(item_id):
    existing = dataviz_content_get(item_id)
    if not existing or existing.get("author") != current_user.alpha_role:
        return jsonify({"error": "Not found"}), 404
    if "image" not in request.files or not request.files["image"].filename:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["image"]
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_IMAGE_EXTENSIONS:
        return jsonify({"error": "Allowed types: png, jpg, jpeg, gif, webp"}), 400
    item = dataviz_content_set_image(item_id, secure_filename(file.filename), file.read())
    return jsonify({"success": True, "item": item})


def _dataviz_public_image_url(item: dict) -> str | None:
    if item.get("image_filename"):
        return f"/api/dataviz/content/{item['id']}/image"
    return None


_DATAVIZ_PUBLIC_FIELDS = ["id", "page", "title", "description", "positive_analysis", "warning", "link", "published_at"]


@app.route("/api/dataviz/<slug>/content", methods=["GET"])
def api_dataviz_public_content(slug):
    page = dataviz_page_get(slug)
    if not page:
        return jsonify({"error": "Unknown page"}), 404
    header_image_url = f"/api/dataviz/pages/{slug}/header-image" if page.get("has_header_image") else None
    return jsonify({"page": slug, "label": page["label"], "description": page.get("description"), "header_image_url": header_image_url})


@app.route("/api/dataviz/global-heat-map/live", methods=["GET"])
def api_market_pulse_live():
    return jsonify(_fetch_market_pulse_live())


@app.route("/api/dataviz/gold-silver-ratio/live", methods=["GET"])
def api_gold_silver_ratio_live():
    force = request.args.get("refresh") in ("1", "true", "yes")
    return jsonify(_fetch_gold_silver_ratio_live(force=force))


@app.route("/api/dataviz/pages", methods=["GET", "POST"])
def api_dataviz_pages():
    if request.method == "GET":
        return jsonify({"pages": dataviz_pages_list()})
    if not current_user.is_authenticated or getattr(current_user, "alpha_role", None) not in DATAVIZ_AUTHORS:
        return jsonify({"error": "This account has no Data Visualisation author access"}), 403
    # multipart/form-data: a visualisation is now created in one step — label,
    # optional description, optional image — rather than a page plus separate
    # "posts" underneath it (see the Studio rework this replaced).
    label = (request.form.get("label") or "").strip()
    description = (request.form.get("description") or "").strip() or None
    if not label:
        return jsonify({"error": "Title is required"}), 400
    if len(label) > 80:
        return jsonify({"error": "Title must be 80 characters or fewer"}), 400
    page = dataviz_page_create(label, current_user.alpha_role, description=description)
    image_file = request.files.get("image")
    if image_file and image_file.filename:
        ext = image_file.filename.rsplit(".", 1)[-1].lower() if "." in image_file.filename else ""
        if ext not in ALLOWED_IMAGE_EXTENSIONS:
            return jsonify({"error": "Allowed image types: png, jpg, jpeg, gif, webp"}), 400
        dataviz_page_set_header_image(page["slug"], secure_filename(image_file.filename), image_file.read())
        page["has_header_image"] = True
    return jsonify({"success": True, "page": page})


@app.route("/api/dataviz/pages/<slug>", methods=["PATCH"])
@login_required
@dataviz_author_required
def api_dataviz_page_update(slug):
    if not dataviz_page_get(slug):
        return jsonify({"error": "Not found"}), 404
    data = request.get_json() or {}
    updates = {}
    if "label" in data:
        label = (data["label"] or "").strip()
        if not label:
            return jsonify({"error": "Title can't be empty"}), 400
        if len(label) > 80:
            return jsonify({"error": "Title must be 80 characters or fewer"}), 400
        updates["label"] = label
    if "description" in data:
        updates["description"] = (data["description"] or "").strip() or None
    page = dataviz_page_update(slug, updates)
    return jsonify({"success": True, "page": page})


@app.route("/api/dataviz/pages/<slug>/header-image", methods=["GET"])
def api_dataviz_page_header_image(slug):
    filename, file_bytes = dataviz_page_get_header_image(slug)
    if file_bytes is None:
        return jsonify({"error": "No header image set"}), 404
    import mimetypes
    mimetype = mimetypes.guess_type(filename or "")[0] or "image/jpeg"
    return Response(file_bytes, mimetype=mimetype, headers={"Cache-Control": "public, max-age=3600"})


@app.route("/api/dataviz/pages/<slug>/header-image", methods=["POST", "DELETE"])
@login_required
@dataviz_author_required
def api_dataviz_page_header_image_upload(slug):
    if not dataviz_page_get(slug):
        return jsonify({"error": "Not found"}), 404
    if request.method == "DELETE":
        dataviz_page_set_header_image(slug, None, None)
        return jsonify({"success": True})
    if "image" not in request.files or not request.files["image"].filename:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["image"]
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_IMAGE_EXTENSIONS:
        return jsonify({"error": "Allowed types: png, jpg, jpeg, gif, webp"}), 400
    dataviz_page_set_header_image(slug, secure_filename(file.filename), file.read())
    return jsonify({"success": True})


@app.route("/api/dataviz/pages/<slug>", methods=["DELETE"])
@login_required
@dataviz_author_required
def api_dataviz_page_delete(slug):
    if not dataviz_page_get(slug):
        return jsonify({"error": "Not found"}), 404
    post_count = len(dataviz_content_list(page=slug))
    if post_count:
        return jsonify({"error": f"This page still has {post_count} post(s) — delete or move them first"}), 400
    dataviz_page_delete(slug)
    return jsonify({"success": True})


@app.route("/api/dataviz/pages/overview", methods=["GET"])
def api_dataviz_pages_overview():
    """Hub-page feed: each visualisation is now the page itself — title,
    description, image — not a page plus a separate feed of posts underneath."""
    overview = []
    for page in dataviz_pages_list():
        overview.append({
            "slug": page["slug"],
            "label": page["label"],
            "description": page.get("description"),
            "image_url": f"/api/dataviz/pages/{page['slug']}/header-image" if page.get("has_header_image") else None,
        })
    return jsonify({"pages": overview})


@app.route("/api/subscription/cancel", methods=["POST"])
@login_required
def api_cancel_subscription():
    users = _load_users()
    if current_user.id not in users:
        return jsonify({"error": "User not found"}), 404
    users[current_user.id]["tier"] = "basic"
    _save_users(users)
    return jsonify({"success": True, "message": "Subscription cancelled. You have been moved to the Basic (free) plan."})


@app.route("/api/account/cancel", methods=["POST"])
@login_required
def api_cancel_account():
    username = current_user.id
    if username == "admin":
        return jsonify({"error": "Cannot delete the admin account"}), 403
    users = _load_users()
    if username not in users:
        return jsonify({"error": "User not found"}), 404
    del users[username]
    _save_users(users)
    logout_user()
    return jsonify({"success": True})


# ── Market data endpoints (public) ───────────────────────────────────────────

@app.route("/api/prices", methods=["GET"])
def prices():
    symbol = request.args.get("symbol", "").strip()
    period = request.args.get("period", "3mo")
    interval = request.args.get("interval", "1d")

    if not symbol:
        return jsonify({"error": "symbol parameter is required"}), 400
    try:
        df = _fetch_ohlcv(symbol, period, interval)
        df.index = df.index.astype(str)
        # yfinance leaves Open/High/Low/Close as NaN on the still-forming "today" bar
        # until it closes — normally overwritten by _stitch_live_tail's OANDA/Alpaca
        # data, but if no live tick has arrived yet the NaN survives through to here.
        # jsonify() (Python's json module) happily emits a literal NaN token, which
        # isn't valid per strict JSON and breaks the chart's client-side JSON.parse.
        # ffill carries the last real close forward (a flat "no trades yet" bar);
        # Volume has no meaningful prior value, so it's zeroed instead.
        df[["Open", "High", "Low", "Close"]] = df[["Open", "High", "Low", "Close"]].ffill()
        df["Volume"] = df["Volume"].fillna(0)
        records = df[["Open", "High", "Low", "Close", "Volume"]].reset_index()
        records.rename(columns={"Date": "date", "Datetime": "date"}, inplace=True)
        if "date" not in records.columns:
            records.rename(columns={records.columns[0]: "date"}, inplace=True)
        return jsonify({
            "symbol": symbol.upper(),
            "period": period,
            "interval": interval,
            "count": len(records),
            "data": records.to_dict(orient="records"),
            "is_live": marketdata_router.is_symbol_live(symbol),
        })
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": f"Failed to fetch prices: {str(e)}"}), 500


# SSE stream of the current still-forming bar, read straight from marketdata's
# in-memory LiveBarBuffer — no Yahoo call ever happens in this loop, so 1s cadence
# is free. Only useful for symbols a streamer (OANDA/Alpaca) covers; for anything
# else it just sends heartbeats and the client keeps relying on its normal polling.
# Streams are deliberately short-lived (~55s) and rely on EventSource's automatic
# reconnect: that keeps any one gunicorn thread from being pinned forever and lets
# deploys/restarts drain quickly. Requires threaded workers (see Procfile) — a
# single sync worker would let one open stream block every other request.
@app.route("/api/prices/stream", methods=["GET"])
def prices_stream():
    symbol = request.args.get("symbol", "").strip()
    interval = request.args.get("interval", "1m")

    if not symbol:
        return jsonify({"error": "symbol parameter is required"}), 400

    # Make sure this symbol's provider stream is running (no-op if none covers it).
    marketdata_router.ensure_symbol_watched(symbol)

    def generate():
        last_payload = None
        deadline = time.monotonic() + 55
        yield "retry: 2000\n\n"
        while time.monotonic() < deadline:
            try:
                tail = marketdata_router.get_live_tail(symbol, interval)
                if tail is not None and not tail.empty:
                    row = tail.iloc[-1]
                    payload = json.dumps({
                        "symbol": symbol.upper(),
                        "interval": interval,
                        "date": str(tail.index[-1]),
                        "Open": float(row["Open"]),
                        "High": float(row["High"]),
                        "Low": float(row["Low"]),
                        "Close": float(row["Close"]),
                        "Volume": float(row["Volume"]),
                        "is_live": True,
                    })
                    if payload != last_payload:
                        last_payload = payload
                        yield f"data: {payload}\n\n"
                    else:
                        yield ": hb\n\n"
                else:
                    # Not covered / stream not warm yet — heartbeat keeps the
                    # connection (and any proxy in front of it) from timing out.
                    yield ": idle\n\n"
            except GeneratorExit:
                raise
            except Exception:
                yield ": err\n\n"
            time.sleep(1)

    resp = Response(generate(), mimetype="text/event-stream")
    resp.headers["Cache-Control"] = "no-cache"
    resp.headers["X-Accel-Buffering"] = "no"  # disable proxy buffering (nginx et al.)
    resp.headers["Connection"] = "keep-alive"
    return resp


@app.route("/api/symbol-search", methods=["GET"])
def symbol_search():
    """Ticker/company autocomplete, proxied through yfinance's Search (which itself
    wraps Yahoo's public search endpoint). Best-effort — degrades to an empty result
    list on any failure (missing yfinance.Search in an older pinned version, Yahoo
    throttling, network error, etc.) rather than surfacing a 500 to the picker UI."""
    query = request.args.get("q", "").strip()
    if len(query) < 1:
        return jsonify({"results": []})
    try:
        quotes = yf.Search(query, max_results=10).quotes
    except Exception:
        return jsonify({"results": []})
    results = []
    for q in quotes:
        symbol = q.get("symbol")
        if not symbol:
            continue
        results.append({
            "symbol": symbol,
            "name": q.get("shortname") or q.get("longname") or symbol,
            "exchange": q.get("exchDisp") or q.get("exchange") or "",
            "type": q.get("quoteType") or "",
        })
    return jsonify({"results": results})


_CUSTOM_SYMBOLS_MAX = 50


@app.route("/api/custom-symbols", methods=["GET"])
@login_required
def get_custom_symbols():
    users = _load_users()
    prefs = users.get(current_user.id, {}).get("preferences", {}) or {}
    return jsonify({"symbols": prefs.get("custom_symbols", [])})


@app.route("/api/custom-symbols", methods=["POST"])
@login_required
def add_custom_symbol():
    data = request.get_json() or {}
    symbol = (data.get("symbol") or "").strip().upper()
    if not symbol:
        return jsonify({"error": "symbol is required"}), 400
    label = (data.get("label") or symbol).strip()
    category = (data.get("category") or "stock").strip().lower()
    exchange = (data.get("exchange") or "").strip()

    users = _load_users()
    if current_user.id not in users:
        return jsonify({"error": "User not found"}), 404
    prefs = users[current_user.id].get("preferences", {}) or {}
    custom = prefs.get("custom_symbols", [])

    if not any(s["symbol"] == symbol for s in custom):
        if len(custom) >= _CUSTOM_SYMBOLS_MAX:
            return jsonify({"error": f"Custom symbol list is full (max {_CUSTOM_SYMBOLS_MAX})"}), 400
        try:
            df = _fetch_ohlcv(symbol, period="5d", interval="1d")
        except Exception:
            df = None
        if df is None or df.empty:
            return jsonify({"error": f"No data found for '{symbol}' — check the ticker"}), 400
        custom.append({"symbol": symbol, "label": label, "category": category, "exchange": exchange})
        marketdata_router.ensure_symbol_watched(symbol)

    prefs["custom_symbols"] = custom
    _save_preferences(current_user.id, prefs)
    return jsonify({"success": True, "symbols": custom})


@app.route("/api/custom-symbols", methods=["DELETE"])
@login_required
def remove_custom_symbol():
    symbol = (request.args.get("symbol") or "").strip().upper()
    users = _load_users()
    if current_user.id not in users:
        return jsonify({"error": "User not found"}), 404
    prefs = users[current_user.id].get("preferences", {}) or {}
    custom = [s for s in prefs.get("custom_symbols", []) if s["symbol"] != symbol]
    prefs["custom_symbols"] = custom
    _save_preferences(current_user.id, prefs)
    return jsonify({"success": True, "symbols": custom})


@app.route("/api/indicators", methods=["GET"])
def indicators():
    symbol = request.args.get("symbol", "").strip()
    period = request.args.get("period", "6mo")
    interval = request.args.get("interval", "1d")

    if not symbol:
        return jsonify({"error": "symbol parameter is required"}), 400
    try:
        df = _fetch_ohlcv(symbol, period, interval)
        calc_params = _extract_calc_params(request.args)
        result = calculate_all(df, **calc_params)
        return jsonify({"symbol": symbol.upper(), "period": period, "interval": interval, **result})
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": f"Failed to calculate indicators: {str(e)}"}), 500


# Shared with the engine worker below (docs/scaling-plan.md, Workstream 5) so
# the two never drift apart on which query params count as scoring thresholds.
_SIGNALS_THRESHOLD_FLOAT_KEYS = [
    "rsi_oversold", "rsi_overbought", "volume_surge",
    "macd_threshold", "bb_oversold", "bb_overbought",
    "rsi_on", "macd_on", "bb_on", "ma_on", "vol_on",
    "ema_short", "ema_long", "macd_cross_lookback", "ema_cross_lookback", "ma_cross_lookback",
]


def _extract_signal_thresholds(args) -> dict:
    thresholds = {}
    for key in _SIGNALS_THRESHOLD_FLOAT_KEYS:
        val = args.get(key)
        if val is not None:
            thresholds[key] = float(val)  # ValueError propagates — callers decide how to report it
    return thresholds


@app.route("/api/signals", methods=["GET"])
@login_required
def signals():
    symbol = request.args.get("symbol", "").strip()
    period = request.args.get("period", "6mo")
    interval = request.args.get("interval", "1d")

    if not symbol:
        return jsonify({"error": "symbol parameter is required"}), 400

    try:
        thresholds = _extract_signal_thresholds(request.args)
    except ValueError as e:
        return jsonify({"error": f"Invalid threshold value: {e}"}), 400

    calc_params = _extract_calc_params(request.args)

    try:
        df = _fetch_ohlcv(symbol, period, interval)
        indicator_data = calculate_all(df, **calc_params)
        signal_result = score_signals(indicator_data, thresholds or None)
        return jsonify({
            "symbol": symbol.upper(),
            "period": period,
            "interval": interval,
            "indicators": {k: v for k, v in indicator_data.items() if k != "history"},
            **signal_result,
        })
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": f"Failed to generate signals: {str(e)}"}), 500


# ── Live signal engine worker (docs/scaling-plan.md, Workstream 5) ──────────
# Background thread that keeps calculate_all()/score_signals() results warm
# for every (symbol, period, interval, params) job the realtime/ service's
# /stream/signals endpoint is currently serving at least one viewer for, and
# publishes the result back over Redis instead of making each viewer poll
# /api/signals themselves. Nothing here runs unless REDIS_URL is set (see
# _should_start_background_streams' invocation below); every failure is
# caught and logged, same "pure optimization layer, never load-bearing"
# posture as marketdata/'s own streamers — /api/signals keeps working exactly
# as it does today regardless of whether this worker is running.
#
# Indicator computation (calculate_all, the expensive part — a full OHLCV
# fetch plus every enabled indicator over the whole history) is deduped and
# throttled per (symbol, period, interval, calc_params), since that's all
# calculate_all's output actually depends on. Scoring (score_signals, cheap
# threshold comparisons against already-computed indicator values) still
# runs per job even when several jobs share one indicator computation, so
# per-user threshold customization is never lost to the sharing.
_ENGINE_THROTTLE_SECONDS = 3
_engine_lock = threading.Lock()
_engine_last_computed: dict[str, float] = {}   # indicator cache key -> monotonic time
_engine_indicator_cache: dict[str, dict] = {}  # indicator cache key -> calculate_all() output

# Both this cache and _ohlcv_cache above are keyed by (symbol, period, interval, ...)
# combos with no natural expiry of the key itself — every distinct combo ever
# requested stays in the dict forever, growing RAM usage all day until a restart.
# Swept periodically (see _sweep_stale_caches) rather than on every write, since
# these caches are hit far more often than jobs actually go stale.
_CACHE_SWEEP_INTERVAL_SECONDS = 300   # how often to run the sweep
_CACHE_ENTRY_MAX_AGE_SECONDS = 600    # evict entries idle longer than this
_last_cache_sweep = 0.0


def _sweep_stale_caches(now: float) -> None:
    """Evict cache entries that haven't been refreshed in a while. Cheap TTL caches
    (_ohlcv_cache, 4s TTL) never need entries older than a few minutes — anything
    that old is a symbol/timeframe nobody's actively polling anymore. Same idea for
    the engine's indicator cache: once a job drops out of the desired set, its entry
    just sits there unused."""
    with _ohlcv_cache_lock:
        stale_keys = [k for k, (ts, _df) in _ohlcv_cache.items() if now - ts > _CACHE_ENTRY_MAX_AGE_SECONDS]
        for k in stale_keys:
            del _ohlcv_cache[k]
    with _engine_lock:
        stale_keys = [k for k, ts in _engine_last_computed.items() if now - ts > _CACHE_ENTRY_MAX_AGE_SECONDS]
        for k in stale_keys:
            _engine_last_computed.pop(k, None)
            _engine_indicator_cache.pop(k, None)


class _DictArgs:
    """Adapts a plain dict to the .get(key)-only interface _extract_calc_params
    and _extract_signal_thresholds expect from a werkzeug MultiDict, so the
    engine worker can reuse those functions unchanged against a job's stored
    params instead of a live request.args."""

    def __init__(self, d: dict):
        self._d = d

    def get(self, key, default=None):
        return self._d.get(key, default)


def _engine_indicator_cache_key(symbol: str, period: str, interval: str, calc_params: dict) -> str:
    blob = json.dumps(
        {"symbol": symbol, "period": period, "interval": interval, "calc_params": calc_params},
        sort_keys=True, separators=(",", ":"),
    )
    return hashlib.sha1(blob.encode()).hexdigest()


def _engine_job_channel(job_json: str) -> str:
    # Must match realtime/main.py's channel derivation exactly — see that
    # file's stream_signals() for the other half of this contract.
    return "signals:" + hashlib.sha1(job_json.encode()).hexdigest()


def _engine_worker_tick() -> None:
    job_keys = marketdata_bus.get_set_members("engine:desired")
    if not job_keys:
        return
    now = time.monotonic()
    for job_json in job_keys:
        try:
            job = json.loads(job_json)
            symbol = job["symbol"]
            period = job.get("period", "6mo")
            interval = job.get("interval", "1d")
            args = _DictArgs(job.get("params") or {})
            calc_params = _extract_calc_params(args)
            try:
                thresholds = _extract_signal_thresholds(args)
            except ValueError:
                continue  # a bad threshold value here means a malformed job — skip, don't crash the loop

            cache_key = _engine_indicator_cache_key(symbol, period, interval, calc_params)
            with _engine_lock:
                stale = (now - _engine_last_computed.get(cache_key, 0)) >= _ENGINE_THROTTLE_SECONDS
            if stale:
                df = _fetch_ohlcv(symbol, period, interval)
                indicator_data = calculate_all(df, **calc_params)
                with _engine_lock:
                    _engine_indicator_cache[cache_key] = indicator_data
                    _engine_last_computed[cache_key] = now
            with _engine_lock:
                indicator_data = _engine_indicator_cache.get(cache_key)
            if indicator_data is None:
                continue

            signal_result = score_signals(indicator_data, thresholds or None)
            payload = {
                "symbol": symbol.upper(),
                "period": period,
                "interval": interval,
                "indicators": {k: v for k, v in indicator_data.items() if k != "history"},
                **signal_result,
            }
            marketdata_bus.publish(_engine_job_channel(job_json), payload)
        except Exception:
            app.logger.exception("engine worker failed for job %s", job_json)


def _engine_worker_loop() -> None:
    global _last_cache_sweep
    while True:
        try:
            _engine_worker_tick()
            now = time.monotonic()
            if now - _last_cache_sweep > _CACHE_SWEEP_INTERVAL_SECONDS:
                _sweep_stale_caches(now)
                _last_cache_sweep = now
        except Exception:
            app.logger.exception("engine worker tick failed")
        time.sleep(1)


# ── Backtest endpoint (public) ───────────────────────────────────────────────

@app.route("/api/backtest", methods=["GET"])
def backtest():
    symbol     = request.args.get("symbol", "").strip()
    period     = request.args.get("period", "2y")
    interval   = request.args.get("interval", "1d")
    start_date = request.args.get("start_date", "").strip() or None
    end_date   = request.args.get("end_date", "").strip() or None

    if not symbol:
        return jsonify({"error": "symbol parameter is required"}), 400

    try:
        stop_loss_pct   = float(request.args.get("stop_loss",     100.0))
        take_profit_pct = float(request.args.get("take_profit",   200.0))
        min_confidence  = float(request.args.get("min_confidence", 60.0))
        trailing_stop      = request.args.get("trailing_stop", "0") in ("1", "true", "True")
        trail_distance_pct = float(request.args.get("trail_distance", 1.5))
        capital            = float(request.args.get("capital", 10000.0))
        trade_amount        = float(request.args.get("trade_amount", 100.0))
        trade_amount_mode   = request.args.get("trade_amount_mode", "percent").strip().lower()
        sl_tp_unit          = request.args.get("sl_tp_unit", "percent").strip().lower()
    except (ValueError, TypeError) as e:
        return jsonify({"error": f"Invalid parameter: {e}"}), 400

    if trade_amount_mode not in ("percent", "gbp"):
        return jsonify({"error": "trade_amount_mode must be 'percent' or 'gbp'"}), 400
    if sl_tp_unit not in ("percent", "pips"):
        return jsonify({"error": "sl_tp_unit must be 'percent' or 'pips'"}), 400

    thresholds = {}
    for key in [
        "rsi_oversold", "rsi_overbought", "volume_surge",
        "bb_oversold", "bb_overbought",
        "rsi_on", "macd_on", "bb_on", "ma_on", "vol_on",
        "macd_cross_lookback", "ema_cross_lookback", "ma_cross_lookback",
        # ── Extended indicator set (Backtester) ──────────────────────────
        "adx_on", "adx_trend_threshold",
        "psar_on", "psar_flip_lookback",
        "ichimoku_on",
        "supertrend_on", "supertrend_flip_lookback",
        "donchian_on",
        "hma_on",
        "stoch_on", "stoch_oversold", "stoch_overbought",
        "stochrsi_on", "stochrsi_oversold", "stochrsi_overbought",
        "cci_on", "cci_oversold", "cci_overbought",
        "willr_on", "willr_oversold", "willr_overbought",
        "roc_on", "roc_threshold",
        "mfi_on", "mfi_oversold", "mfi_overbought",
        "tsi_on", "tsi_oversold", "tsi_overbought",
        "ao_on",
        "atr_on", "atr_trend_lookback",
        "keltner_on",
        "stdev_on", "stdev_trend_lookback",
        "chaikin_vol_on", "chaikin_vol_trend_lookback",
        "hist_vol_on", "hist_vol_trend_lookback",
        "obv_on",
        "vwap_on",
        "ad_on",
        "cmf_on", "cmf_threshold",
        "vol_profile_on",
        "fib_on", "fib_tolerance_pct",
        "inv_hs_on", "inv_hs_tolerance_pct",
        "macd_centerline_lookback", "macd_zscore_overbought", "macd_zscore_oversold",
        "ma_trigger_lookback",
        "adx_di_cross_lookback",
        "psar_gap_lookback", "supertrend_gap_lookback",
        "ichimoku_tk_cross_lookback", "donchian_mid_cross_lookback", "hma_price_cross_lookback",
        "hma_two_cross_lookback",
        "stoch_signal_cross_lookback", "stochrsi_signal_cross_lookback",
        "cci_centerline_lookback", "willr_midline_lookback", "roc_centerline_lookback",
        "mfi_centerline_lookback", "tsi_centerline_lookback", "ao_zero_cross_lookback",
        "keltner_mid_cross_lookback", "cmf_centerline_lookback", "vol_profile_breakout_lookback",
        "bb_breakout_margin_pct", "bb_pct_below_high", "bb_pct_above_low",
        "donchian_retest_lookback", "donchian_retest_tolerance_pct",
    ]:
        val = request.args.get(key)
        if val is not None:
            try:
                thresholds[key] = float(val)
            except ValueError:
                return jsonify({"error": f"Invalid value for '{key}'"}), 400

    # Each indicator's trigger can now have several modes active at once — the
    # frontend sends one repeated query key per active mode (?rsi_trigger=a&rsi_trigger=b),
    # which Flask collects with getlist(). Falls back to DEFAULT_THRESHOLDS' single-mode
    # default (in score_signals) if the caller sends nothing for a given indicator.
    for trig_key, allowed in _TRIGGER_WHITELISTS.items():
        vals = request.args.getlist(trig_key)
        if vals:
            cleaned = []
            for v in vals:
                if v not in allowed:
                    return jsonify({"error": f"Invalid value for '{trig_key}'"}), 400
                if v not in cleaned:
                    cleaned.append(v)
            thresholds[trig_key] = cleaned

    calc_params = _extract_calc_params(request.args)
    calc_params.update(_extract_backtest_calc_params(request.args))

    try:
        df = _fetch_ohlcv(symbol, period, interval, start_date=start_date, end_date=end_date)

        # Forex pip size: derived from the actual traded price rather than
        # assuming only JPY-named pairs quote to 2 decimals. Any pair whose
        # price sits in JPY-like ranges (some exotics quote similarly) gets
        # priced the same way, instead of silently defaulting to 0.0001 and
        # mis-sizing pip-based stop-loss/take-profit for it.
        if len(df.index) > 0:
            last_price = float(df["Close"].iloc[-1])
            pip_size = 0.01 if last_price >= 20 else 0.0001
        else:
            pip_size = 0.01 if "JPY" in symbol.upper() else 0.0001

        trades, equity_curve, bah_curve = run_backtest(
            df,
            thresholds=thresholds or None,
            calc_params=calc_params or None,
            stop_loss_pct=stop_loss_pct,
            take_profit_pct=take_profit_pct,
            min_confidence=min_confidence,
            trailing_stop=trailing_stop,
            trail_distance_pct=trail_distance_pct,
            capital=capital,
            trade_amount_mode=trade_amount_mode,
            trade_amount=trade_amount,
            sl_tp_unit=sl_tp_unit,
            pip_size=pip_size,
        )
        metrics = calculate_metrics(trades, equity_curve)

        # Best-effort: sector/peer context per trade (which sector the symbol
        # belongs to, how that sector and its well-known peers moved on each
        # trade's entry/exit day). Never fails the backtest itself.
        if len(df.index) > 0:
            try:
                enrich_trades_with_sector_context(
                    trades, symbol,
                    str(df.index[0].date()), str(df.index[-1].date()),
                )
            except Exception:
                pass

        period_label = f"{start_date} → {end_date or 'today'}" if start_date else period
        return jsonify({
            "symbol":       symbol.upper(),
            "period":       period_label,
            "interval":     interval,
            "metrics":      metrics,
            "trades":       trades,
            "equity_curve": equity_curve,
            "bah_curve":    bah_curve,
        })
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": f"Backtest failed: {str(e)}"}), 500


# ── Social Post Studio — generate platform-ready posts via Claude ────────────

_SOCIAL_PLATFORMS = {
    "x": {
        "label": "X",
        "guidance": (
            "X (Twitter) post. HARD LIMIT: body + cta combined must be under 270 "
            "characters. Punchy, one clear idea, strong hook. No hashtags. "
            "No title needed (set title to empty string)."
        ),
    },
    "instagram": {
        "label": "Instagram",
        "guidance": (
            "Instagram caption. First line must be a scroll-stopping hook (it gets "
            "truncated). Short paragraphs with line breaks, tasteful emojis. "
            "No hashtags. No title (empty string)."
        ),
    },
    "facebook": {
        "label": "Facebook",
        "guidance": (
            "Facebook post. Conversational, 2-3 short paragraphs, invites "
            "comments/shares. No hashtags. No title (empty string)."
        ),
    },
    "substack": {
        "label": "Substack",
        "guidance": (
            "Substack note/post. Include a compelling title. Body 150-300 words, "
            "written like a mini-essay with a personal, direct voice. No hashtags."
        ),
    },
    "email": {
        "label": "Email newsletter",
        "guidance": (
            "Email newsletter section. 'title' = the subject line (under 60 chars, "
            "curiosity-driven). Body 100-200 words, scannable, warm. CTA should "
            "read like button text + one supporting line. No hashtags."
        ),
    },
}

_SOCIAL_POST_SCHEMA = {
    "type": "object",
    "properties": {
        "title": {"type": "string"},
        "body": {"type": "string"},
        "cta": {"type": "string"},
        "image_prompt": {"type": "string"},
    },
    "required": ["title", "body", "cta", "image_prompt"],
    "additionalProperties": False,
}


@app.route("/api/social-posts", methods=["POST"])
@alpha_author_required
def social_posts():
    if not os.environ.get("ANTHROPIC_API_KEY"):
        return jsonify({"error": "No Anthropic API key is configured on the server "
                                 "(set the ANTHROPIC_API_KEY environment variable)."}), 503

    data = request.get_json(silent=True) or {}
    platform = _SOCIAL_PLATFORMS.get(data.get("platform"))
    if not platform:
        return jsonify({"error": "Unknown platform"}), 400
    ideas = (data.get("ideas") or "").strip()
    if not ideas:
        return jsonify({"error": "Add some ideas or bullet points first."}), 400

    prompt_parts = [
        "You are an expert social media copywriter.",
        "",
        f"Create a {platform['label']} post from the author's raw notes below.",
        "",
        "AUTHOR'S IDEAS / BULLET POINTS:",
        ideas[:8000],
    ]
    prompt_parts.append(
        "TARGET AUDIENCE: people aged 18-50 who are trying to get into investing "
        "as a hobby, a career, or as a necessity"
    )
    cta_goal = (data.get("ctaGoal") or "").strip()
    if cta_goal:
        prompt_parts.append(f"GOAL OF THE CALL TO ACTION: {cta_goal[:500]}")
    source_url = (data.get("sourceUrl") or "").strip()
    if source_url:
        prompt_parts.append(
            f"URL TO INCLUDE: {source_url[:500]} — place this link at the very "
            "end of the post (after the CTA text), on its own line."
        )
    prompt_parts.append("BRAND VOICE NOTES: authoritative and friendly")
    prompt_parts += [
        f"TONE: {(data.get('tone') or 'Friendly')[:50]}",
        "",
        f"PLATFORM RULES: {platform['guidance']}",
        "",
        'Also write "image_prompt": a detailed prompt (40-70 words) the author can '
        "paste into an AI image generator (Midjourney, DALL-E, etc.) to create a "
        "matching visual. Describe subject, style, mood, colours, composition. "
        "No text in the image.",
    ]

    try:
        import anthropic
        client = anthropic.Anthropic()
        response = client.messages.create(
            model="claude-opus-4-8",
            max_tokens=2000,
            output_config={"format": {"type": "json_schema", "schema": _SOCIAL_POST_SCHEMA}},
            messages=[{"role": "user", "content": "\n".join(prompt_parts)}],
        )
        if response.stop_reason == "refusal":
            return jsonify({"error": "The model declined to write this post."}), 502
        text = next((b.text for b in response.content if b.type == "text"), "")
        return jsonify(json.loads(text))
    except Exception as e:
        return jsonify({"error": f"Generation failed: {e}"}), 502


# ── Startup ───────────────────────────────────────────────────────────────────

_ensure_table()
_ensure_default_user()


def _should_start_background_streams() -> bool:
    if _is_production:
        return True  # gunicorn / non-debug run: no reloader involved, start once
    # Local `python app.py` with debug=True: Werkzeug's reloader re-execs the process
    # with WERKZEUG_RUN_MAIN=true once it's the real serving process — app.debug isn't
    # usable for this check here, since app.run(debug=...) below hasn't executed yet
    # at module-load time, so only WERKZEUG_RUN_MAIN reliably tells the two passes
    # apart. Only start in the real serving process, not the one about to be replaced
    # by that re-exec.
    return os.environ.get("WERKZEUG_RUN_MAIN") == "true"


if _should_start_background_streams():
    try:
        _seed_users = _load_users()
        _seed_symbols = [
            s["symbol"]
            for u in _seed_users.values()
            for s in (u.get("preferences", {}) or {}).get("custom_symbols", [])
        ]
        marketdata_router.start_background_streams(seed_symbols=_seed_symbols)
    except Exception:
        pass  # marketdata is a pure optimization layer — never block startup on it

    if marketdata_config.REDIS_URL:
        threading.Thread(target=_engine_worker_loop, name="engine-worker", daemon=True).start()

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(debug=not _is_production, port=port)
